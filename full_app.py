# Unified Legal Assistant Backend - Combines RAG Q&A and Document Analysis
# This version merges both backends to run on a single port
# RAG portion modified to behave like the second backend

from fastapi import FastAPI, Request, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
import os
import json
import requests
import re
import logging
import uuid
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Tuple, Set, Any
from dataclasses import dataclass
from enum import Enum
import io
import tempfile
import sys
import traceback

# Third-party library imports for RAG
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
import spacy
from sentence_transformers import SentenceTransformer
import numpy as np
from chromadb.config import Settings

# AI imports
try:
    import aiohttp
    AIOHTTP_AVAILABLE = True
except ImportError:
    AIOHTTP_AVAILABLE = False
    print("⚠️ aiohttp not available - AI features disabled. Install with: pip install aiohttp")

# Import open-source NLP models support
OPEN_SOURCE_NLP_AVAILABLE = False
try:
    import torch
    from transformers import (
        pipeline, 
        AutoTokenizer, 
        AutoModelForSequenceClassification,
        AutoModelForTokenClassification,
        AutoModelForQuestionAnswering,
        T5ForConditionalGeneration,
        T5Tokenizer
    )
    OPEN_SOURCE_NLP_AVAILABLE = True
    print("✅ Open-source NLP models available (transformers + torch)")
except ImportError:
    print("⚠️ Open-source NLP not available. Install with: pip install transformers torch")

# Set up logging FIRST
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Configuration ---
CHROMA_PATH = os.path.abspath(os.path.join(os.getcwd(), "chromadb-database"))
logger.info(f"Using CHROMA_PATH: {CHROMA_PATH}")

CHROMA_CLIENT_SETTINGS = Settings(
    persist_directory=CHROMA_PATH,
    anonymized_telemetry=False,
    allow_reset=True,
    is_persistent=True
)

# OpenRouter configuration
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_API_KEY = os.environ.get("OPENAI_API_KEY", "")
OPENAI_API_BASE = os.environ.get("OPENAI_API_BASE", "https://openrouter.ai/api/v1")
AI_ENABLED = AIOHTTP_AVAILABLE and bool(OPENROUTER_API_KEY)

if AI_ENABLED:
    print("✅ AI features enabled with OpenRouter/DeepSeek")
else:
    print("⚠️ AI features disabled - using fact extraction only")

# In-memory conversation storage
conversations: Dict[str, Dict] = {}

# Load NLP Models
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("spaCy model loaded successfully")
except Exception as e:
    logger.error(f"Failed to load spaCy model: {e}")
    nlp = None

try:
    sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
    logger.info("Sentence transformer loaded successfully")
except Exception as e:
    logger.error(f"Failed to load Sentence Transformer model: {e}")
    sentence_model = None

# Import document processing libraries
print("Starting document processing imports...")
PYMUPDF_AVAILABLE = False
PDFPLUMBER_AVAILABLE = False

# Import PyPDF2 (always needed as fallback)
try:
    import PyPDF2
    print("✅ PyPDF2 imported successfully")
except ImportError as e:
    print(f"❌ CRITICAL: PyPDF2 import failed: {e}")
    print("Install with: pip install PyPDF2")
    sys.exit(1)

# Import python-docx
try:
    import docx
    print("✅ python-docx imported successfully")
except ImportError as e:
    print(f"❌ CRITICAL: python-docx import failed: {e}")
    print("Install with: pip install python-docx")
    sys.exit(1)

# Try to import PyMuPDF
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    print("✅ PyMuPDF available - using high-quality PDF processing")
except ImportError as e:
    print(f"⚠️ PyMuPDF not available: {e}")
    print("Install with: pip install PyMuPDF")

# Try to import pdfplumber
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    print("✅ pdfplumber available - using enhanced PDF extraction")
except ImportError as e:
    print(f"⚠️ pdfplumber not available: {e}")
    print("Install with: pip install pdfplumber")

print(f"PDF processing status: PyMuPDF={PYMUPDF_AVAILABLE}, pdfplumber={PDFPLUMBER_AVAILABLE}")

# Create FastAPI app
app = FastAPI(
    title="Unified Legal Assistant API",
    description="Combined RAG Q&A and Document Analysis System",
    version="6.0.0-Unified"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Pydantic Models ---

# RAG Models
class Query(BaseModel):
    question: str
    session_id: Optional[str] = None
    response_style: Optional[str] = "balanced"  # "concise", "balanced", "detailed"

class QueryResponse(BaseModel):
    response: Optional[str] = None
    error: Optional[str] = None
    context_found: bool = False
    sources: Optional[list] = None
    session_id: str
    confidence_score: float = 0.0
    expand_available: bool = False

class ConversationHistory(BaseModel):
    session_id: str
    messages: List[Dict]
    created_at: str
    last_updated: str

# Document Analysis Models
class EnhancedAnalysisResponse(BaseModel):
    response: Optional[str] = None
    summary: Optional[str] = None
    factual_summary: Optional[str] = None
    ai_analysis: Optional[str] = None
    extraction_results: Optional[Dict[str, Any]] = None
    analysis_type: str
    confidence_score: float
    processing_info: Optional[Dict[str, Any]] = None
    verification_status: str
    status: str = "completed"
    success: bool = True
    warnings: List[str] = []
    session_id: str
    timestamp: str
    model_used: str = "deepseek-chat" if AI_ENABLED else "fact-extraction-only"

# --- RAG System Functions (Modified to behave like second backend) ---

def cleanup_expired_conversations():
    """Remove conversations older than 1 hour"""
    now = datetime.utcnow()
    expired_sessions = [
        session_id for session_id, data in conversations.items()
        if now - data['last_accessed'] > timedelta(hours=1)
    ]
    for session_id in expired_sessions:
        del conversations[session_id]

def load_database():
    """Load the Chroma database"""
    try:
        embedding_function = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        db = Chroma(
            collection_name="default",
            embedding_function=embedding_function,
            persist_directory=CHROMA_PATH,
            client_settings=CHROMA_CLIENT_SETTINGS
        )
        return db
    except Exception as e:
        logger.error(f"Failed to load database: {e}")
        raise

def get_conversation_context(session_id: str, max_messages: int = 12) -> str:
    """Get recent conversation history as context"""
    if session_id not in conversations:
        return ""
    messages = conversations[session_id]['messages'][-max_messages:]
    context_parts = []
    for msg in messages:
        role = msg['role'].upper()
        content = msg['content']
        # For conversation history, keep more content but still truncate very long messages
        if len(content) > 800:
            content = content[:800] + "..."
        context_parts.append(f"{role}: {content}")
    
    # If we have conversation history, make it clear this is ongoing context
    if context_parts:
        return "Previous conversation:\n" + "\n".join(context_parts)
    return ""

def add_to_conversation(session_id: str, role: str, content: str, sources: Optional[List] = None):
    """Add a message to the conversation history"""
    if session_id not in conversations:
        conversations[session_id] = {
            'messages': [],
            'created_at': datetime.utcnow(),
            'last_accessed': datetime.utcnow()
        }
    message = {
        'role': role,
        'content': content,
        'timestamp': datetime.utcnow().isoformat(),
        'sources': sources or []
    }
    conversations[session_id]['messages'].append(message)
    conversations[session_id]['last_accessed'] = datetime.utcnow()
    # Keep only last 30 messages to prevent memory issues but maintain more history
    if len(conversations[session_id]['messages']) > 30:
        conversations[session_id]['messages'] = conversations[session_id]['messages'][-30:]

def parse_multiple_questions(query_text: str) -> List[str]:
    """Parse multiple questions from input"""
    questions = []
    query_text = query_text.strip()
    
    if '?' in query_text and not query_text.endswith('?'):
        parts = query_text.split('?')
        for part in parts:
            part = part.strip()
            if part:
                questions.append(part + '?')
    else:
        final_question = query_text
        if not final_question.endswith('?') and '?' not in final_question:
            final_question += '?'
        questions = [final_question]
    
    return questions

def enhanced_retrieval_v2(db, query_text: str, conversation_history_context: str, k: int = 8) -> Tuple[List, Any]:
    """
    Performs retrieval, potentially using conversation history to refine the query.
    Modified to be less strict like the second backend.
    """
    # First, try the original query
    logger.info(f"[RETRIEVAL] Searching for query: '{query_text}'")
    
    try:
        # Perform similarity search with scores using original query
        results_with_scores = db.similarity_search_with_relevance_scores(query_text, k=k)
        logger.info(f"[RETRIEVAL] Found {len(results_with_scores)} raw results")
        
        # If we have conversation history and few results, try combined query
        if conversation_history_context and len(results_with_scores) < 3:
            combined_query = f"{conversation_history_context}\n\n{query_text}".strip()
            logger.info(f"[RETRIEVAL] Trying combined query: '{combined_query}'")
            combined_results = db.similarity_search_with_relevance_scores(combined_query, k=k)
            # Use combined results if they're better
            if len(combined_results) > len(results_with_scores):
                results_with_scores = combined_results
        
        # Filter based on a minimum relevance score (LOWERED from 0.4 to 0.2)
        filtered_results = [(doc, score) for doc, score in results_with_scores if score > 0.2]
        logger.info(f"[RETRIEVAL] Filtered to {len(filtered_results)} results with score > 0.2")
        
        # If we have filtered results, use them; otherwise use all results
        final_results = filtered_results if filtered_results else results_with_scores
        
        docs, scores = zip(*final_results) if final_results else ([], [])
        return list(docs), {"query_used": query_text, "scores": list(scores)}
        
    except Exception as e:
        logger.error(f"[RETRIEVAL] Search failed: {e}")
        return [], {"error": str(e)}

def expand_legal_query(query: str) -> str:
    """Expand query with legal synonyms and terms"""
    legal_expansions = {
        "asylum": "asylum refugee protection persecution",
        "immigration": "immigration visa deportation removal",
        "contract": "contract agreement terms conditions breach",
        "criminal": "criminal penal prosecution sentence conviction",
        "civil": "civil tort liability damages compensation",
        "court": "court tribunal judge judicial proceeding",
        "law": "law statute regulation rule code",
        "rights": "rights protections constitutional fundamental"
    }
    
    expanded_terms = []
    query_lower = query.lower()
    
    for term, expansion in legal_expansions.items():
        if term in query_lower:
            expanded_terms.extend(expansion.split())
    
    if expanded_terms:
        return f"{query} {' '.join(set(expanded_terms))}"
    return query

def extract_sub_queries(query: str) -> List[str]:
    """Extract sub-questions from complex queries"""
    sub_queries = []
    
    conjunctions = [" and ", " or ", " also ", " plus ", " additionally "]
    current_query = query
    
    for conj in conjunctions:
        if conj in current_query.lower():
            parts = current_query.split(conj)
            sub_queries.extend([part.strip() for part in parts if len(part.strip()) > 10])
            break
    
    legal_patterns = [
        r"what (?:is|are) ([^?]+)\??",
        r"how (?:does|do) ([^?]+)\??",
        r"when (?:is|are) ([^?]+)\??",
        r"(?:explain|describe) ([^?]+)\??"
    ]
    
    for pattern in legal_patterns:
        matches = re.findall(pattern, query, re.IGNORECASE)
        for match in matches:
            if len(match.strip()) > 5:
                sub_queries.append(f"what is {match.strip()}?")
    
    return sub_queries[:3]

def remove_duplicate_documents(results_with_scores: List[Tuple]) -> List[Tuple]:
    """Remove duplicate documents based on content similarity"""
    if not results_with_scores:
        return []
    
    unique_results = []
    seen_content = set()
    
    for doc, score in results_with_scores:
        content_hash = hash(doc.page_content[:100])
        if content_hash not in seen_content:
            seen_content.add(content_hash)
            unique_results.append((doc, score))
    
    return sorted(unique_results, key=lambda x: x[1], reverse=True)

def calculate_dynamic_threshold(query: str, results: List[Tuple]) -> float:
    """Calculate dynamic threshold based on query complexity and result distribution"""
    if not results:
        return 0.3
    
    scores = [score for _, score in results]
    
    is_complex = len(query.split()) > 8 or '?' in query[:-1]
    has_legal_terms = any(term in query.lower() for term in ['law', 'legal', 'court', 'statute', 'regulation'])
    
    if len(scores) > 1:
        avg_score = np.mean(scores)
        std_score = np.std(scores)
        
        if is_complex:
            return max(0.25, avg_score - std_score)
        elif has_legal_terms:
            return max(0.35, avg_score - 0.5 * std_score)
        else:
            return max(0.4, avg_score - 0.3 * std_score)
    
    return 0.3

def format_response_by_style(content: str, sources: List[Dict], style: str = "balanced") -> Tuple[str, bool]:
    """Format response based on user's preferred style"""
    
    if style == "concise":
        concise_response = create_concise_response(content, sources)
        return concise_response, True
    
    elif style == "detailed":
        detailed_response = create_detailed_response(content, sources)
        return detailed_response, False
    
    else:  # balanced
        balanced_response = create_balanced_response(content, sources)
        return balanced_response, True

def create_concise_response(content: str, sources: List[Dict]) -> str:
    """Create a concise, bullet-point response"""
    lines = content.split('\n')
    key_points = []
    
    for line in lines[:5]:
        if line.strip() and not line.startswith('#'):
            key_points.append(f"• {line.strip()}")
    
    concise = f"""**Quick Answer:**
{chr(10).join(key_points)}

💡 *Need more details? Ask me to expand on any point above.*"""
    
    return concise

def create_balanced_response(content: str, sources: List[Dict]) -> str:
    """Create a balanced response with clear sections"""
    if len(content) > 800:
        preview = content[:600] + "..."
        balanced = f"""{preview}

📖 **Want the complete analysis?** Ask me to provide the full detailed response.
🔍 **Have specific questions?** Ask about any particular aspect mentioned above."""
    else:
        balanced = content
    
    return balanced

def create_detailed_response(content: str, sources: List[Dict]) -> str:
    """Return the full detailed response"""
    return content

def calculate_confidence_score(results: List, search_result: Dict, response_length: int) -> float:
    """Calculate confidence score based on multiple factors"""
    if not results:
        return 0.1
    
    scores = search_result.get("scores", [])
    if not scores:
        return 0.2
    
    avg_relevance = np.mean(scores)
    doc_factor = min(1.0, len(results) / 5.0)
    
    if len(scores) > 1:
        score_std = np.std(scores)
        consistency_factor = max(0.5, 1.0 - score_std)
    else:
        consistency_factor = 0.7
    
    completeness_factor = min(1.0, response_length / 500.0)
    
    confidence = (
        avg_relevance * 0.4 +
        doc_factor * 0.3 +
        consistency_factor * 0.2 +
        completeness_factor * 0.1
    )
    
    return min(1.0, max(0.0, confidence))

def create_enhanced_context(results: List, search_result: Dict, questions: List[str]) -> Tuple[str, List[Dict]]:
    """Enhanced context creation with better relevance filtering"""
    if not results:
        return "", []
    
    context_parts = []
    source_info = []
    seen_sources = set()
    
    # LOWERED threshold from 0.4 to match second backend behavior
    MIN_RELEVANCE_FOR_CONTEXT = 0.25
    
    for i, (doc, score) in enumerate(zip(results, search_result.get("scores", [0.0]*len(results)))):
        # Skip low-relevance documents
        if score < MIN_RELEVANCE_FOR_CONTEXT:
            continue
            
        content = doc.page_content.strip()
        if not content:
            continue
        
        source_path = doc.metadata.get('source', 'Unknown Source')
        page = doc.metadata.get('page', None)
        
        source_id = (source_path, page)
        if source_id in seen_sources:
            continue
        seen_sources.add(source_id)
        
        display_source = os.path.basename(source_path)
        page_info = f" (Page {page})" if page is not None else ""
        
        # Format exactly like second backend
        context_part = f"[{display_source}{page_info}] (Relevance: {score:.2f}): {content}"
        context_parts.append(context_part)
        
        source_info.append({
            'id': i+1,
            'file_name': display_source,
            'page': page,
            'relevance': score,
            'full_path': source_path
        })
    
    context_text = "\n\n".join(context_parts)
    return context_text, source_info

def call_openrouter_api(prompt: str, api_key: str, api_base: str = "https://openrouter.ai/api/v1") -> str:
    """Call OpenRouter API with fallback models - matching second backend exactly"""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8000",
        "X-Title": "Legal Assistant"
    }
    
    # Use same model list and order as second backend
    models_to_try = [
        "deepseek/deepseek-chat",
        "microsoft/phi-3-mini-128k-instruct:free",
        "meta-llama/llama-3.2-3b-instruct:free",
        "google/gemma-2-9b-it:free",
        "mistralai/mistral-7b-instruct:free",
        "openchat/openchat-7b:free"
    ]
    
    logger.info(f"Trying {len(models_to_try)} models...")
    last_exception = None
    
    for i, model in enumerate(models_to_try):
        try:
            payload = {
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,  # Changed from 0.5 to 0.7 to match second backend
                "max_tokens": 3000,  # Changed from 2000 to 3000 to match second backend
                "top_p": 0.95       # Changed from 0.9 to 0.95 to match second backend
            }
            logger.info(f"Attempting model {i+1}/{len(models_to_try)}: {model}")
            
            response = requests.post(
                f"{api_base}/chat/completions",
                headers=headers,
                json=payload,
                timeout=60
            )
            
            # Log response details for debugging
            logger.info(f"Response status: {response.status_code}")
            logger.info(f"Response headers: {dict(response.headers)}")
            
            # Check if response is empty
            if not response.content:
                logger.error(f"Empty response from {model}")
                last_exception = Exception("Empty response from API")
                continue
                
            # Log raw response content for debugging
            response_text = response.content.decode('utf-8')
            logger.info(f"Raw response content (first 200 chars): {response_text[:200]}")
            
            # Check HTTP status
            if response.status_code != 200:
                logger.error(f"HTTP {response.status_code} error from {model}: {response_text}")
                last_exception = requests.exceptions.HTTPError(f"HTTP {response.status_code}: {response_text}")
                continue
            
            # Try to parse JSON
            try:
                result = response.json()
            except json.JSONDecodeError as json_err:
                logger.error(f"JSON decode error from {model}: {json_err}")
                logger.error(f"Response content: {response_text}")
                last_exception = json_err
                continue
            
            # Check if response has expected structure
            if 'choices' not in result or not result['choices']:
                logger.error(f"Invalid response structure from {model}: {result}")
                last_exception = Exception(f"Invalid response structure: {result}")
                continue
                
            if 'message' not in result['choices'][0] or 'content' not in result['choices'][0]['message']:
                logger.error(f"Missing message content from {model}: {result}")
                last_exception = Exception(f"Missing message content: {result}")
                continue
            
            response_content = result['choices'][0]['message']['content']
            if not response_content or not response_content.strip():
                logger.error(f"Empty content from {model}")
                last_exception = Exception("Empty content in response")
                continue
                
            logger.info(f"Success with model {model}! Response length: {len(response_content)}")
            return response_content.strip()
            
        except requests.exceptions.Timeout:
            logger.error(f"Timeout error for model {model}")
            last_exception = Exception("Request timeout")
            continue
        except requests.exceptions.ConnectionError as conn_err:
            logger.error(f"Connection error for model {model}: {conn_err}")
            last_exception = conn_err
            continue
        except requests.exceptions.RequestException as req_err:
            logger.error(f"Request error for model {model}: {req_err}")
            last_exception = req_err
            continue
        except Exception as e:
            logger.error(f"Unexpected error for model {model}: {e}")
            last_exception = e
            continue
    
    # If all models failed
    error_msg = f"All models failed. Last error: {str(last_exception)}"
    logger.error(error_msg)
    
    # Return a fallback response instead of raising an exception
    fallback_response = (
        "I apologize, but I'm currently experiencing technical difficulties with the AI service. "
        "This could be due to API limitations, network issues, or temporary service unavailability. "
        "Please try again in a few moments, or contact support if the issue persists."
    )
    
    return fallback_response

# --- Document Analysis System (from second backend) ---

# AI Analysis prompts
ANALYSIS_PROMPTS = {
    'summarize': """You are a legal document analyst. Analyze this document and provide:
1. A clear summary of the document's purpose and type
2. The main parties involved (with their roles)
3. Key terms and conditions
4. Important dates and deadlines
5. Financial obligations or amounts
6. Any notable risks or concerns

Document text:
{document_text}

Provide a structured summary in plain English while maintaining legal accuracy.""",

    'extract-clauses': """You are a legal document analyst. Extract and categorize the following types of clauses from this document:
1. Termination clauses
2. Indemnification provisions
3. Liability limitations
4. Governing law and jurisdiction
5. Confidentiality/NDA provisions
6. Payment terms
7. Dispute resolution mechanisms

For each clause found, provide:
- Clause type
- Summary of the provision
- Exact location/section reference if available
- Any unusual or concerning aspects

Document text:
{document_text}""",

    'missing-clauses': """You are a legal document analyst. Review this contract and identify commonly expected clauses that appear to be missing or inadequately addressed:

Consider standard clauses such as:
- Force majeure
- Limitation of liability
- Indemnification
- Dispute resolution/arbitration
- Confidentiality
- Termination provisions
- Assignment restrictions
- Severability
- Entire agreement
- Notice provisions
- Governing law

Document text:
{document_text}

For each missing clause, explain why it's typically important and the risks of its absence.""",

    'risk-flagging': """You are a legal risk analyst. Identify and assess legal risks in this document:

Look for:
1. Unilateral termination rights
2. Broad indemnification requirements
3. Unlimited liability exposure
4. Vague or ambiguous obligations
5. Unfavorable payment terms
6. Lack of protection clauses
7. Unusual warranty provisions
8. Problematic intellectual property terms

Document text:
{document_text}

For each risk, provide:
- Risk description
- Severity (High/Medium/Low)
- Potential impact
- Suggested mitigation""",

    'timeline-extraction': """You are a legal document analyst. Extract all time-related information:

Find and list:
1. Contract start and end dates
2. Payment deadlines
3. Notice periods
4. Renewal dates and terms
5. Termination notice requirements
6. Performance deadlines
7. Warranty periods
8. Any other time-sensitive obligations

Document text:
{document_text}

Present as a chronological timeline with clear labels.""",

    'obligations': """You are a legal document analyst. List all obligations and requirements for each party:

Identify:
1. What each party must do
2. When they must do it
3. Conditions or prerequisites
4. Consequences of non-compliance
5. Reporting or notification requirements

Document text:
{document_text}

Organize by party and priority/timeline."""
}

async def perform_ai_analysis(document_text: str, analysis_type: str) -> Tuple[str, float]:
    """Perform AI analysis using DeepSeek via OpenRouter or open-source models as fallback"""
    
    # First try DeepSeek if available
    if AI_ENABLED:
        prompt_template = ANALYSIS_PROMPTS.get(analysis_type, ANALYSIS_PROMPTS['summarize'])
        prompt = prompt_template.format(document_text=document_text[:15000])
        
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:8000",
            "X-Title": "Legal Document Analyzer"
        }
        
        payload = {
            "model": "deepseek/deepseek-chat",
            "messages": [
                {
                    "role": "system",
                    "content": "You are an expert legal document analyst. Provide thorough, accurate analysis while clearly marking any uncertainties. Always include relevant disclaimers about seeking professional legal advice."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,
            "max_tokens": 2000
        }
        
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(OPENROUTER_API_URL, json=payload, headers=headers) as response:
                    if response.status == 200:
                        data = await response.json()
                        ai_response = data['choices'][0]['message']['content']
                        confidence = 0.85
                        return ai_response, confidence
                    else:
                        error_text = await response.text()
                        logger.error(f"OpenRouter API error: {response.status} - {error_text}")
        except Exception as e:
            logger.error(f"AI analysis exception: {e}")
    
    # If DeepSeek fails or is not available, try open-source models
    if open_source_analyzer and open_source_analyzer.models_loaded:
        logger.info(f"Falling back to open-source models for {analysis_type}")
        try:
            return await open_source_analyzer.analyze_document(document_text, analysis_type)
        except Exception as e:
            logger.error(f"Open-source analysis failed: {e}")
    
    # If both fail, return error message
    return "AI analysis not available. Please set OPENAI_API_KEY environment variable or install transformers for open-source models.", 0.0

class VerifiableExtractor:
    """Extract only verifiable information from documents with source locations"""
    
    def __init__(self):
        self.extraction_patterns = {
            'dates': {
                'patterns': [
                    r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
                    r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
                    r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
                ],
                'context_required': ['due', 'deadline', 'effective', 'expires', 'terminates', 'begins', 'starts', 'ends']
            },
            'monetary_amounts': {
                'patterns': [
                    r'\$\s*[\d,]+(?:\.\d{2})?(?:\s*(?:million|Million|billion|Billion|M|B))?\b',
                    r'\$\s*[\d,]+(?:\.\d{2})?(?:\s*(?:per|each|monthly|annually|yearly))?',
                    r'\b[\d,]+(?:\.\d{2})?\s*(?:USD|usd|dollars?|Dollars?)(?:\s*(?:per|each|monthly|annually|yearly))?\b',
                    r'\b(?:USD|usd)\s*[\d,]+(?:\.\d{2})?\b',
                    r'\b[\d,]+\s*cents?\b'
                ],
                'context_required': ['pay', 'payment', 'fee', 'cost', 'price', 'rent', 'salary', 'compensation', 'amount', 'value', 'worth', 'charge', 'expense']
            },
            'percentages': {
                'patterns': [
                    r'\b\d+(?:\.\d+)?%',
                    r'\b\d+(?:\.\d+)?\s*(?:percent|per cent)\b'
                ],
                'context_required': ['rate', 'interest', 'fee', 'penalty', 'commission', 'tax', 'discount', 'margin', 'share', 'ownership', 'equity', 'royalty']
            },
            'party_names': {
                'patterns': [
                    r'([A-Z][A-Za-z\s&.,\'"-]+?(?:Corporation|Corp\.|Company|Co\.|LLC|L\.L\.C\.|Inc\.|Incorporated|Limited|Ltd\.|LLP|LP|plc|PLC))',
                    r'[Bb]y:\s*/?s?/?\s*([A-Z][A-Za-z\s\.-]+?)(?:\n|,|\s{2,})',
                    r'Name:\s*([A-Z][A-Za-z\s\.-]+?)(?:\n|,|\s{2,})',
                    r'between\s+([A-Z][A-Za-z\s&.,\'"-]+?)\s+(?:and|AND)\s+([A-Z][A-Za-z\s&.,\'"-]+?)(?:\s|,|\()',
                    r'\((?:the\s+)?"?([A-Z][A-Za-z\s&.,\'"-]+?)"?\)(?:\s+(?:and|AND)|,)',
                    r'undersigned,?\s+([A-Z][A-Za-z\s&.,\'"-]+?)(?:,|\s+(?:hereby|certifies|agrees))',
                    r'([A-Z][A-Za-z\s\.-]+?),?\s*(?:P\.E\.|P\.Eng\.|Ph\.D\.|Esq\.|CPA|MBA|MD|JD)',
                    r'(?:party|parties)[:\s]+([A-Z][a-zA-Z\s&.,]+?)(?:\s*(?:and|,|\n))',
                    r'(?:company|corporation|llc|inc\.?)[:\s]*([A-Z][a-zA-Z\s&.,]+?)(?:\s*(?:and|,|\n))'
                ],
                'context_required': []
            }
        }
    
    def extract_with_verification(self, document_text: str, extraction_type: str) -> List[Dict[str, Any]]:
        """Extract information only if it can be verified with high confidence"""
        
        if extraction_type not in self.extraction_patterns:
            return [{"status": "failed", "reason": f"Unknown extraction type: {extraction_type}"}]
        
        pattern_config = self.extraction_patterns[extraction_type]
        extracted_items = []
        seen_values = set()
        
        lines = document_text.split('\n')
        
        for line_num, line in enumerate(lines, 1):
            line_lower = line.lower()
            
            for pattern in pattern_config['patterns']:
                matches = re.finditer(pattern, line, re.IGNORECASE)
                
                for match in matches:
                    if len(match.groups()) > 0:
                        for group_idx in range(1, len(match.groups()) + 1):
                            if match.group(group_idx):
                                extracted_value = match.group(group_idx).strip()
                                extracted_value = re.sub(r'\s+', ' ', extracted_value)
                                extracted_value = extracted_value.strip(' ,.')
                                
                                if len(extracted_value) < 3 or extracted_value.lower() in seen_values:
                                    continue
                                
                                if extraction_type == 'party_names':
                                    if extracted_value.lower() in ['the company', 'the corporation', 'party', 'parties', 'undersigned']:
                                        continue
                                    if not any(word[0].isupper() for word in extracted_value.split() if word):
                                        continue
                                
                                seen_values.add(extracted_value.lower())
                                
                                context_verified = True
                                if pattern_config['context_required']:
                                    context_verified = any(
                                        keyword in line_lower 
                                        for keyword in pattern_config['context_required']
                                    )
                                
                                if context_verified:
                                    extracted_items.append({
                                        "value": extracted_value,
                                        "line_number": line_num,
                                        "context": line.strip(),
                                        "confidence": "high",
                                        "verified": True
                                    })
                                else:
                                    extracted_items.append({
                                        "value": extracted_value,
                                        "line_number": line_num,
                                        "context": line.strip(),
                                        "confidence": "low",
                                        "verified": False,
                                        "reason": "No supporting context found"
                                    })
                    else:
                        extracted_value = match.group(0).strip()
                        
                        if extracted_value.lower() in seen_values:
                            continue
                        
                        seen_values.add(extracted_value.lower())
                        
                        context_verified = True
                        if pattern_config['context_required']:
                            context_verified = any(
                                keyword in line_lower 
                                for keyword in pattern_config['context_required']
                            )
                        
                        if context_verified:
                            extracted_items.append({
                                "value": extracted_value,
                                "line_number": line_num,
                                "context": line.strip(),
                                "confidence": "high",
                                "verified": True
                            })
        
        verified_items = [item for item in extracted_items if item.get('verified', False)]
        verified_items.sort(key=lambda x: x['line_number'])
        
        if not verified_items:
            return [{"status": "failed_to_extract", "reason": f"No verifiable {extraction_type} found in document"}]
        
        return verified_items

class NoHallucinationAnalyzer:
    """Document analyzer that never hallucinates - only reports verifiable facts"""
    
    def __init__(self):
        self.extractor = VerifiableExtractor()
        
        self.document_types = {
            'contract': {
                'required_keywords': ['agreement', 'contract'],
                'supporting_keywords': ['parties', 'consideration', 'terms'],
                'minimum_matches': 2
            },
            'lease': {
                'required_keywords': ['lease'],
                'supporting_keywords': ['tenant', 'landlord', 'rent', 'premises'],
                'minimum_matches': 2
            },
            'employment': {
                'required_keywords': ['employment', 'employee'],
                'supporting_keywords': ['employer', 'salary', 'position', 'job'],
                'minimum_matches': 2
            },
            'policy': {
                'required_keywords': ['policy'],
                'supporting_keywords': ['procedure', 'guidelines', 'rules'],
                'minimum_matches': 1
            },
            'consent': {
                'required_keywords': ['consent'],
                'supporting_keywords': ['authorize', 'permit', 'sec', 'filing', 'qualified person', 'technical report'],
                'minimum_matches': 2
            },
            'nda': {
                'required_keywords': ['confidential', 'non-disclosure', 'nda'],
                'supporting_keywords': ['proprietary', 'secret', 'disclose'],
                'minimum_matches': 2
            },
            'license': {
                'required_keywords': ['license', 'licensing'],
                'supporting_keywords': ['grant', 'rights', 'royalty', 'software', 'intellectual property'],
                'minimum_matches': 2
            },
            'purchase': {
                'required_keywords': ['purchase', 'sale', 'acquisition'],
                'supporting_keywords': ['buyer', 'seller', 'price', 'assets', 'shares'],
                'minimum_matches': 2
            }
        }
    
    def detect_document_type_strict(self, document_text: str) -> Dict[str, Any]:
        """Detect document type only with high confidence"""
        
        text_lower = document_text.lower()
        detection_results = {}
        
        for doc_type, criteria in self.document_types.items():
            score = 0
            matched_keywords = []
            
            required_found = 0
            for keyword in criteria['required_keywords']:
                if keyword in text_lower:
                    required_found += 1
                    matched_keywords.append(keyword)
                    score += 10
            
            supporting_found = 0
            for keyword in criteria['supporting_keywords']:
                if keyword in text_lower:
                    supporting_found += 1
                    matched_keywords.append(keyword)
                    score += 3
            
            total_matches = required_found + supporting_found
            if total_matches >= criteria['minimum_matches'] and required_found > 0:
                detection_results[doc_type] = {
                    'score': score,
                    'matched_keywords': matched_keywords,
                    'required_found': required_found,
                    'supporting_found': supporting_found
                }
        
        if not detection_results:
            return {
                'type': 'unknown',
                'confidence': 0.0,
                'status': 'failed_to_detect',
                'reason': 'No clear document type indicators found',
                'matched_keywords': []
            }
        
        best_type = max(detection_results, key=lambda x: detection_results[x]['score'])
        best_result = detection_results[best_type]
        
        confidence = min(1.0, best_result['score'] / 20)
        
        return {
            'type': best_type,
            'confidence': confidence,
            'status': 'detected' if confidence > 0.7 else 'uncertain',
            'matched_keywords': best_result['matched_keywords'],
            'all_matches': detection_results
        }
    
    def extract_document_facts(self, document_text: str) -> Dict[str, Any]:
        """Extract only verifiable facts from the document"""
        
        facts = {
            'basic_stats': self._get_basic_stats(document_text),
            'dates': self.extractor.extract_with_verification(document_text, 'dates'),
            'monetary_amounts': self.extractor.extract_with_verification(document_text, 'monetary_amounts'),
            'percentages': self.extractor.extract_with_verification(document_text, 'percentages'),
            'party_names': self.extractor.extract_with_verification(document_text, 'party_names'),
            'document_structure': self._analyze_structure(document_text),
            'extraction_status': 'completed',
            'timestamp': datetime.utcnow().isoformat()
        }
        
        return facts
    
    def _get_basic_stats(self, text: str) -> Dict[str, Any]:
        """Get verifiable basic statistics"""
        
        words = text.split()
        sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        return {
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len(paragraphs),
            'character_count': len(text),
            'estimated_reading_time_minutes': round(len(words) / 200, 1)
        }
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure"""
        
        numbered_sections = len(re.findall(r'^\s*\d+\.\s', text, re.MULTILINE))
        lettered_sections = len(re.findall(r'^\s*[a-z]\.\s', text, re.MULTILINE))
        subsections = len(re.findall(r'^\s*\d+\.\d+\s', text, re.MULTILINE))
        
        potential_headers = re.findall(r'^([^:\n]{1,50}):$', text, re.MULTILINE)
        headers = len([h for h in potential_headers if len(h.split()) <= 8])
        
        return {
            'numbered_sections': numbered_sections,
            'lettered_sections': lettered_sections,
            'subsections': subsections,
            'headers': headers,
            'has_clear_structure': numbered_sections > 0 or headers > 2
        }
    
    def generate_factual_summary(self, document_text: str) -> str:
        """Generate a summary using only extracted facts"""
        
        doc_type_result = self.detect_document_type_strict(document_text)
        facts = self.extract_document_facts(document_text)
        
        summary_parts = []
        
        if doc_type_result['status'] == 'detected':
            summary_parts.append(f"**Document Type**: {doc_type_result['type'].title()} (confidence: {doc_type_result['confidence']:.1%})")
            summary_parts.append(f"**Keywords Found**: {', '.join(doc_type_result['matched_keywords'])}")
        else:
            summary_parts.append(f"**Document Type**: Failed to detect - {doc_type_result['reason']}")
        
        stats = facts['basic_stats']
        summary_parts.append(f"\n**Document Statistics**:")
        summary_parts.append(f"• Word count: {stats['word_count']:,}")
        summary_parts.append(f"• Estimated reading time: {stats['estimated_reading_time_minutes']} minutes")
        summary_parts.append(f"• Paragraphs: {stats['paragraph_count']}")
        
        structure = facts['document_structure']
        summary_parts.append(f"\n**Document Structure**:")
        if structure['has_clear_structure']:
            summary_parts.append(f"• Numbered sections: {structure['numbered_sections']}")
            summary_parts.append(f"• Headers found: {structure['headers']}")
            summary_parts.append(f"• Subsections: {structure['subsections']}")
        else:
            summary_parts.append("• No clear structural organization detected")
        
        summary_parts.append(f"\n**Extracted Information**:")
        
        dates = facts['dates']
        if dates and dates[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Dates found**: {len(dates)} verifiable dates")
            for i, date in enumerate(dates[:3], 1):
                summary_parts.append(f"  {i}. {date['value']} (Line {date['line_number']})")
            if len(dates) > 3:
                summary_parts.append(f"  ... and {len(dates) - 3} more dates")
        else:
            summary_parts.append("• **Dates**: Failed to extract any verifiable dates")
        
        amounts = facts['monetary_amounts']
        if amounts and amounts[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Financial amounts**: {len(amounts)} verifiable amounts")
            for i, amount in enumerate(amounts[:3], 1):
                summary_parts.append(f"  {i}. {amount['value']} (Line {amount['line_number']})")
            if len(amounts) > 3:
                summary_parts.append(f"  ... and {len(amounts) - 3} more amounts")
        else:
            summary_parts.append("• **Financial amounts**: Failed to extract any verifiable amounts")
        
        percentages = facts['percentages']
        if percentages and percentages[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Percentages**: {len(percentages)} found")
            for perc in percentages[:3]:
                summary_parts.append(f"  - {perc['value']} (Line {perc['line_number']})")
        else:
            summary_parts.append("• **Percentages**: Failed to extract any verifiable percentages")
        
        parties = facts['party_names']
        if parties and parties[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Parties**: {len(parties)} identified")
            for party in parties[:3]:
                summary_parts.append(f"  - {party['value']} (Line {party['line_number']})")
        else:
            summary_parts.append("• **Parties**: Failed to extract clear party names")
        
        summary_parts.append(f"\n**⚠️ IMPORTANT NOTES**:")
        summary_parts.append("• This analysis only includes information that could be verified directly from the document text")
        summary_parts.append("• All extracted items include line numbers for verification")
        summary_parts.append("• Items marked 'Failed to extract' mean the information was not clearly identifiable")
        summary_parts.append("• For legal advice, consult a qualified attorney")
        
        return '\n'.join(summary_parts)

class SafeDocumentProcessor:
    """Document processor with verification and no hallucination"""
    
    @staticmethod
    def process_document_safe(file: UploadFile) -> Tuple[str, str, Dict[str, Any]]:
        """Process document and provide processing metadata"""
        
        file_extension = file.filename.split('.')[-1].lower() if file.filename else 'unknown'
        processing_info = {
            'original_filename': file.filename,
            'file_extension': file_extension,
            'processing_method': None,
            'success': False,
            'warnings': [],
            'errors': []
        }
        
        try:
            file.file.seek(0)
            
            if file_extension == 'pdf':
                text, warnings = SafeDocumentProcessor._extract_pdf_safe(file)
                processing_info['processing_method'] = 'PDF extraction'
                processing_info['warnings'] = warnings
                
            elif file_extension in ['docx', 'doc']:
                text, warnings = SafeDocumentProcessor._extract_docx_safe(file)
                processing_info['processing_method'] = 'Word document extraction'
                processing_info['warnings'] = warnings
                
            elif file_extension == 'txt':
                content = file.file.read()
                text = content.decode('utf-8')
                processing_info['processing_method'] = 'Plain text'
                
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file_extension}. Supported: PDF, DOCX, TXT"
                )
            
            if not text or len(text.strip()) < 10:
                processing_info['errors'].append("Extracted text is too short or empty")
                raise ValueError("Failed to extract meaningful text from document")
            
            processing_info['success'] = True
            processing_info['extracted_length'] = len(text)
            processing_info['word_count'] = len(text.split())
            
            return text, file_extension, processing_info
            
        except Exception as e:
            processing_info['errors'].append(str(e))
            logger.error(f"Document processing error: {type(e).__name__}: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Document processing failed: {str(e)}")
    
    @staticmethod
    def _extract_pdf_safe(file: UploadFile) -> Tuple[str, List[str]]:
        """Extract PDF text with PyMuPDF first"""
        
        warnings = []
        
        file.file.seek(0)
        pdf_content = file.file.read()
        
        if len(pdf_content) == 0:
            raise ValueError("PDF file is empty")
        
        if PYMUPDF_AVAILABLE:
            try:
                import fitz
                
                pdf_doc = fitz.open(stream=pdf_content, filetype="pdf")
                
                try:
                    if pdf_doc.page_count == 0:
                        raise ValueError("PDF contains no pages")
                    
                    text = ""
                    pages_with_text = 0
                    tables_found = 0
                    total_pages = pdf_doc.page_count
                    
                    for page_num in range(total_pages):
                        try:
                            page = pdf_doc[page_num]
                            
                            page_text = page.get_text()
                            if page_text and page_text.strip():
                                text += f"[Page {page_num + 1}]\n{page_text}\n"
                                pages_with_text += 1
                            else:
                                warnings.append(f"Page {page_num + 1} contains no extractable text")
                            
                            try:
                                tables = page.find_tables()
                                if tables:
                                    for table_num, table in enumerate(tables):
                                        try:
                                            table_data = table.extract()
                                            if table_data:
                                                text += f"\n[TABLE {table_num + 1} FROM PAGE {page_num + 1}]\n"
                                                for row in table_data:
                                                    if row:
                                                        clean_row = [str(cell) if cell else "" for cell in row]
                                                        text += " | ".join(clean_row) + "\n"
                                                text += "[/TABLE]\n\n"
                                                tables_found += 1
                                        except Exception as table_error:
                                            logger.debug(f"Table extraction error: {table_error}")
                            except Exception as e:
                                logger.debug(f"Table detection error: {e}")
                                
                        except Exception as page_error:
                            warnings.append(f"Error processing page {page_num + 1}: {str(page_error)}")
                            logger.warning(f"Page {page_num + 1} error: {page_error}")
                    
                    if pages_with_text == 0:
                        raise ValueError("No readable text found in any PDF pages")
                    
                    if pages_with_text < total_pages:
                        warnings.append(f"Only {pages_with_text} of {total_pages} pages contained extractable text")
                    
                    if tables_found > 0:
                        warnings.append(f"PyMuPDF extracted {tables_found} tables with preserved structure")
                    else:
                        warnings.append("PyMuPDF processed PDF successfully")
                    
                    return text, warnings
                    
                finally:
                    pdf_doc.close()
                
            except Exception as e:
                warnings.append(f"PyMuPDF processing failed: {str(e)}, falling back to pdfplumber")
                logger.warning(f"PyMuPDF error: {e}")
        
        if PDFPLUMBER_AVAILABLE:
            pdf_file = io.BytesIO(pdf_content)
            
            try:
                import pdfplumber
                with pdfplumber.open(pdf_file) as pdf:
                    if len(pdf.pages) == 0:
                        raise ValueError("PDF contains no pages")
                    
                    text = ""
                    pages_with_text = 0
                    tables_found = 0
                    
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"[Page {page_num + 1}]\n{page_text}\n"
                            pages_with_text += 1
                        else:
                            warnings.append(f"Page {page_num + 1} contains no extractable text")
                        
                        try:
                            tables = page.extract_tables()
                            if tables:
                                for table_num, table in enumerate(tables):
                                    text += f"\n[TABLE {table_num + 1} FROM PAGE {page_num + 1}]\n"
                                    for row in table:
                                        if row:
                                            clean_row = [str(cell) if cell else "" for cell in row]
                                            text += " | ".join(clean_row) + "\n"
                                    text += "[/TABLE]\n\n"
                                    tables_found += 1
                        except Exception as e:
                            warnings.append(f"Could not extract tables from page {page_num + 1}: {str(e)}")
                    
                    if pages_with_text == 0:
                        raise ValueError("No readable text found in any PDF pages")
                    
                    if pages_with_text < len(pdf.pages):
                        warnings.append(f"Only {pages_with_text} of {len(pdf.pages)} pages contained extractable text")
                    
                    if tables_found > 0:
                        warnings.append(f"pdfplumber extracted {tables_found} tables with preserved structure")
                    
                    return text, warnings
                    
            except Exception as e:
                warnings.append(f"pdfplumber failed: {str(e)}, using basic PyPDF2")
                logger.warning(f"pdfplumber error: {e}")
        
        pdf_file = io.BytesIO(pdf_content)
        pdf_file.seek(0)
        
        try:
            reader = PyPDF2.PdfReader(pdf_file)
            
            if len(reader.pages) == 0:
                raise ValueError("PDF contains no pages")
            
            text = ""
            pages_with_text = 0
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += f"[Page {page_num + 1}]\n{page_text}\n\n"
                        pages_with_text += 1
                    else:
                        warnings.append(f"Page {page_num + 1} contains no extractable text")
                except Exception as e:
                    warnings.append(f"Error extracting text from page {page_num + 1}: {str(e)}")
            
            if pages_with_text == 0:
                raise ValueError("No readable text found in any PDF pages")
            
            warnings.append("Used basic PDF extraction - complex layouts may not be preserved")
            return text, warnings
            
        except Exception as e:
            logger.error(f"PyPDF2 failed: {e}")
            raise ValueError(f"All PDF extraction methods failed. Last error: {str(e)}")
    
    @staticmethod
    def _extract_docx_safe(file: UploadFile) -> Tuple[str, List[str]]:
        """Extract DOCX text with warnings about potential issues"""
        
        warnings = []
        
        file.file.seek(0)
        docx_content = file.file.read()
        
        if len(docx_content) == 0:
            raise ValueError("Word document is empty")
        
        docx_file = io.BytesIO(docx_content)
        
        try:
            doc = docx.Document(docx_file)
            
            text = ""
            paragraphs_with_text = 0
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text and paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraphs_with_text += 1
            
            tables_found = len(doc.tables)
            if tables_found > 0:
                warnings.append(f"Document contains {tables_found} tables - table content may not be fully extracted")
                
                for table_num, table in enumerate(doc.tables):
                    text += f"\n[TABLE {table_num + 1}]\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        if any(row_text):
                            text += " | ".join(row_text) + "\n"
                    text += "[/TABLE]\n\n"
            
            if paragraphs_with_text == 0:
                raise ValueError("No readable text found in Word document")
            
            if paragraphs_with_text < len(doc.paragraphs):
                warnings.append(f"Some paragraphs contained no text")
            
            return text, warnings
            
        except Exception as e:
            raise ValueError(f"Failed to process Word document: {str(e)}")

# Open-Source NLP Legal Analyzer
class OpenSourceLegalAnalyzer:
    """Legal document analyzer using open-source models"""
    
    def __init__(self):
        self.device = 0 if torch.cuda.is_available() else -1
        self.models = {}
        self.models_loaded = False
        if OPEN_SOURCE_NLP_AVAILABLE:
            self.load_models()
    
    def load_models(self):
        """Load various open-source models for different tasks"""
        
        print("Loading open-source NLP models...")
        
        try:
            # 1. Document Classification (BERT-based)
            self.models['classifier'] = pipeline(
                "text-classification",
                model="nlpaueb/legal-bert-base-uncased",
                device=self.device
            )
            
            # 2. Named Entity Recognition for legal texts
            self.models['ner'] = pipeline(
                "ner",
                model="dslim/bert-base-NER",
                aggregation_strategy="simple",
                device=self.device
            )
            
            # 3. Summarization using BART or T5
            self.models['summarizer'] = pipeline(
                "summarization",
                model="facebook/bart-large-cnn",  # or "google/flan-t5-base"
                device=self.device
            )
            
            # 4. Question Answering for fact extraction
            self.models['qa'] = pipeline(
                "question-answering",
                model="deepset/roberta-base-squad2",
                device=self.device
            )
            
            # 5. Zero-shot classification for clause detection
            self.models['zero_shot'] = pipeline(
                "zero-shot-classification",
                model="facebook/bart-large-mnli",
                device=self.device
            )
            
            self.models_loaded = True
            print("✅ All open-source models loaded successfully")
            
        except Exception as e:
            print(f"❌ Error loading models: {e}")
            self.models_loaded = False
    
    async def analyze_document(self, document_text: str, analysis_type: str) -> Tuple[str, float]:
        """Perform analysis using open-source models"""
        
        if not self.models_loaded:
            return "Open-source models not loaded. Please install transformers and torch.", 0.0
        
        if analysis_type == "summarize":
            return self._summarize_document(document_text)
        
        elif analysis_type == "extract-clauses":
            return self._extract_clauses(document_text)
        
        elif analysis_type == "risk-flagging":
            return self._flag_risks(document_text)
        
        elif analysis_type == "timeline-extraction":
            return self._extract_timeline(document_text)
        
        elif analysis_type == "obligations":
            return self._extract_obligations(document_text)
        
        elif analysis_type == "missing-clauses":
            return self._detect_missing_clauses(document_text)
        
        else:
            return "Analysis type not supported with open-source models", 0.5
    
    def _summarize_document(self, text: str) -> Tuple[str, float]:
        """Generate document summary using BART/T5"""
        
        # Split into chunks if too long
        max_chunk_length = 1024
        chunks = self._split_into_chunks(text, max_chunk_length)
        
        summaries = []
        for chunk in chunks[:5]:  # Process max 5 chunks
            try:
                summary = self.models['summarizer'](
                    chunk,
                    max_length=150,
                    min_length=50,
                    do_sample=False
                )[0]['summary_text']
                summaries.append(summary)
            except:
                continue
        
        # Combine summaries
        combined_summary = "\n\n".join(summaries)
        
        # Add document type detection
        doc_type = self._detect_document_type(text)
        
        # Extract key information using QA
        parties = self._extract_parties_qa(text)
        dates = self._extract_dates_regex(text)
        amounts = self._extract_amounts_regex(text)
        
        result = f"""## Document Summary

**Document Type**: {doc_type}

**Overview**:
{combined_summary}

**Key Information Extracted**:
- **Parties**: {', '.join(parties) if parties else 'Not clearly identified'}
- **Important Dates**: {', '.join(dates[:5]) if dates else 'No dates found'}
- **Monetary Amounts**: {', '.join(amounts[:5]) if amounts else 'No amounts found'}

**Note**: This summary was generated using open-source AI models (BART/T5). For more detailed analysis, consider using the DeepSeek-powered analysis.
"""
        
        confidence = 0.7 if summaries else 0.3
        return result, confidence
    
    def _extract_clauses(self, text: str) -> Tuple[str, float]:
        """Extract clauses using zero-shot classification"""
        
        # Define clause types to look for
        clause_types = [
            "termination clause",
            "indemnification clause",
            "liability limitation",
            "confidentiality clause",
            "payment terms",
            "governing law",
            "dispute resolution"
        ]
        
        # Split text into paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 50]
        
        found_clauses = {}
        
        for para in paragraphs[:20]:  # Analyze first 20 paragraphs
            try:
                result = self.models['zero_shot'](
                    para,
                    candidate_labels=clause_types,
                    multi_label=True
                )
                
                # If high confidence, add to found clauses
                for label, score in zip(result['labels'], result['scores']):
                    if score > 0.7:
                        if label not in found_clauses:
                            found_clauses[label] = []
                        found_clauses[label].append({
                            'text': para[:200] + '...' if len(para) > 200 else para,
                            'confidence': score
                        })
            except:
                continue
        
        # Format results
        result = "## Extracted Clauses\n\n"
        
        for clause_type, instances in found_clauses.items():
            result += f"### {clause_type.title()}\n"
            for i, instance in enumerate(instances[:2], 1):  # Show max 2 per type
                result += f"{i}. {instance['text']}\n"
                result += f"   *Confidence: {instance['confidence']:.1%}*\n\n"
        
        if not found_clauses:
            result += "No standard clauses were detected with high confidence.\n"
        
        result += "\n**Note**: Clause extraction performed using zero-shot classification. Results may vary based on document structure."
        
        confidence = 0.6 if found_clauses else 0.3
        return result, confidence
    
    def _flag_risks(self, text: str) -> Tuple[str, float]:
        """Flag potential risks using pattern matching and classification"""
        
        risk_patterns = {
            "Unlimited Liability": [
                r"unlimited liability",
                r"no limitation of liability",
                r"fully liable",
                r"without limitation"
            ],
            "Unilateral Termination": [
                r"may terminate at any time",
                r"sole discretion.*terminate",
                r"without cause.*terminate",
                r"immediate termination"
            ],
            "Broad Indemnification": [
                r"indemnify.*all claims",
                r"hold harmless.*any and all",
                r"defend.*at.*own expense",
                r"indemnification.*without limitation"
            ],
            "Unfavorable Jurisdiction": [
                r"exclusive jurisdiction.*(?!your state)",
                r"governed by.*laws of.*(?!your state)",
                r"submit to.*courts of"
            ]
        }
        
        found_risks = []
        
        for risk_type, patterns in risk_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    context_start = max(0, match.start() - 100)
                    context_end = min(len(text), match.end() + 100)
                    context = text[context_start:context_end].strip()
                    
                    found_risks.append({
                        'type': risk_type,
                        'context': context,
                        'severity': 'High' if 'unlimited' in risk_type.lower() else 'Medium'
                    })
        
        # Format results
        result = "## Risk Analysis\n\n"
        
        if found_risks:
            result += f"⚠️ **{len(found_risks)} Potential Risks Identified**\n\n"
            
            for i, risk in enumerate(found_risks[:10], 1):  # Show max 10
                result += f"### Risk {i}: {risk['type']}\n"
                result += f"**Severity**: {risk['severity']}\n"
                result += f"**Context**: ...{risk['context']}...\n\n"
        else:
            result += "✅ No high-risk patterns detected in the document.\n"
        
        result += "\n**Note**: Risk detection based on pattern matching. Always consult with a legal professional for comprehensive risk assessment."
        
        confidence = 0.6 if found_risks else 0.5
        return result, confidence
    
    def _extract_timeline(self, text: str) -> Tuple[str, float]:
        """Extract dates and deadlines using regex and NER"""
        
        # Date patterns
        date_patterns = [
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',
            r'\b\d{1,2}\s+(days?|weeks?|months?|years?)\b',
            r'\b(within|after|before|by|on)\s+\d+\s+(days?|weeks?|months?|years?)\b'
        ]
        
        timeline_items = []
        
        # Extract dates with context
        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Get surrounding context
                context_start = max(0, match.start() - 50)
                context_end = min(len(text), match.end() + 50)
                context = text[context_start:context_end].strip()
                
                # Determine type based on keywords
                deadline_keywords = ['due', 'deadline', 'expire', 'terminate', 'by', 'before']
                is_deadline = any(kw in context.lower() for kw in deadline_keywords)
                
                timeline_items.append({
                    'date': match.group(0),
                    'type': 'Deadline' if is_deadline else 'Date',
                    'context': context
                })
        
        # Sort by appearance in document
        result = "## Timeline & Important Dates\n\n"
        
        if timeline_items:
            for item in timeline_items[:15]:  # Show max 15
                result += f"• **{item['date']}** ({item['type']})\n"
                result += f"  Context: ...{item['context']}...\n\n"
        else:
            result += "No specific dates or deadlines found in the document.\n"
        
        result += "\n**Note**: Dates extracted using pattern matching. Verify context for accuracy."
        
        confidence = 0.7 if timeline_items else 0.3
        return result, confidence
    
    def _extract_obligations(self, text: str) -> Tuple[str, float]:
        """Extract obligations using linguistic patterns"""
        
        obligation_patterns = [
            r'(shall|must|will|required to|obligated to)\s+([^.]{10,100})',
            r'(agrees to|commits to|undertakes to)\s+([^.]{10,100})',
            r'(responsible for|liable for)\s+([^.]{10,100})'
        ]
        
        obligations = []
        
        for pattern in obligation_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                obligation_text = match.group(0).strip()
                modal = match.group(1)
                
                obligations.append({
                    'text': obligation_text,
                    'type': 'Mandatory' if modal in ['shall', 'must'] else 'Contractual'
                })
        
        # Format results
        result = "## Extracted Obligations\n\n"
        
        if obligations:
            # Group by type
            mandatory = [o for o in obligations if o['type'] == 'Mandatory']
            contractual = [o for o in obligations if o['type'] == 'Contractual']
            
            if mandatory:
                result += "### Mandatory Obligations (SHALL/MUST)\n"
                for i, ob in enumerate(mandatory[:10], 1):
                    result += f"{i}. {ob['text']}\n"
                result += "\n"
            
            if contractual:
                result += "### Contractual Obligations\n"
                for i, ob in enumerate(contractual[:10], 1):
                    result += f"{i}. {ob['text']}\n"
        else:
            result += "No clear obligations found in the document.\n"
        
        result += "\n**Note**: Obligations extracted using linguistic patterns. Review full document for complete understanding."
        
        confidence = 0.6 if obligations else 0.3
        return result, confidence
    
    def _detect_missing_clauses(self, text: str) -> Tuple[str, float]:
        """Detect missing standard clauses"""
        
        standard_clauses = {
            "Force Majeure": ["force majeure", "act of god", "unforeseeable"],
            "Limitation of Liability": ["limitation of liability", "liability shall not exceed", "maximum liability"],
            "Indemnification": ["indemnify", "hold harmless", "defend"],
            "Confidentiality": ["confidential", "non-disclosure", "proprietary"],
            "Termination": ["termination", "terminate", "expiration"],
            "Governing Law": ["governing law", "governed by", "laws of"],
            "Dispute Resolution": ["dispute resolution", "arbitration", "mediation"],
            "Assignment": ["assignment", "assign", "transfer rights"],
            "Severability": ["severability", "severable", "invalid provision"],
            "Entire Agreement": ["entire agreement", "whole agreement", "supersedes"]
        }
        
        text_lower = text.lower()
        missing_clauses = []
        
        for clause_name, keywords in standard_clauses.items():
            found = any(keyword in text_lower for keyword in keywords)
            if not found:
                missing_clauses.append(clause_name)
        
        # Format results
        result = "## Missing Clause Analysis\n\n"
        
        if missing_clauses:
            result += f"⚠️ **{len(missing_clauses)} Standard Clauses May Be Missing:**\n\n"
            
            for clause in missing_clauses:
                result += f"### {clause}\n"
                if clause == "Force Majeure":
                    result += "Protects parties from liability due to unforeseeable events.\n"
                elif clause == "Limitation of Liability":
                    result += "Caps potential damages and financial exposure.\n"
                elif clause == "Indemnification":
                    result += "Defines who bears risk for certain claims or damages.\n"
                elif clause == "Confidentiality":
                    result += "Protects sensitive information from disclosure.\n"
                elif clause == "Governing Law":
                    result += "Specifies which jurisdiction's laws apply.\n"
                elif clause == "Dispute Resolution":
                    result += "Defines how conflicts will be resolved.\n"
                else:
                    result += "Standard contractual provision.\n"
                result += "\n"
        else:
            result += "✅ All standard clauses appear to be present.\n"
        
        result += "\n**Note**: This is a keyword-based analysis. Some clauses may be present using different terminology."
        
        confidence = 0.5  # Lower confidence for this type of analysis
        return result, confidence
    
    # Helper methods
    def _split_into_chunks(self, text: str, max_length: int) -> List[str]:
        """Split text into chunks for processing"""
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < max_length:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _detect_document_type(self, text: str) -> str:
        """Detect document type using classification"""
        try:
            # Use zero-shot classification
            labels = ["contract", "lease agreement", "employment agreement", 
                     "NDA", "license agreement", "purchase agreement", "legal brief"]
            
            result = self.models['zero_shot'](
                text[:1000],  # Use first 1000 chars
                candidate_labels=labels
            )
            
            return result['labels'][0].title()
        except:
            return "Legal Document"
    
    def _extract_parties_qa(self, text: str) -> List[str]:
        """Extract party names using QA model"""
        parties = []
        questions = [
            "Who are the parties to this agreement?",
            "What is the name of the first party?",
            "What is the name of the second party?",
            "Who is the employer?",
            "Who is the employee?",
            "Who is the landlord?",
            "Who is the tenant?"
        ]
        
        try:
            for question in questions:
                result = self.models['qa'](question=question, context=text[:2000])
                if result['score'] > 0.5:
                    parties.append(result['answer'])
        except:
            pass
        
        return list(set(parties))  # Remove duplicates
    
    def _extract_dates_regex(self, text: str) -> List[str]:
        """Extract dates using regex"""
        date_pattern = r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b'
        dates = re.findall(date_pattern, text)
        return dates
    
    def _extract_amounts_regex(self, text: str) -> List[str]:
        """Extract monetary amounts using regex"""
        amount_pattern = r'\$[\d,]+(?:\.\d{2})?|\b\d+\s*(?:USD|dollars?)\b'
        amounts = re.findall(amount_pattern, text, re.IGNORECASE)
        return amounts

# Initialize analyzers
safe_analyzer = NoHallucinationAnalyzer()
open_source_analyzer = None
if OPEN_SOURCE_NLP_AVAILABLE:
    try:
        open_source_analyzer = OpenSourceLegalAnalyzer()
    except Exception as e:
        logger.error(f"Failed to initialize open-source analyzer: {e}")
        open_source_analyzer = None

# Modified Prompt template for improved RAG - behaving like second backend
IMPROVED_PROMPT_TEMPLATE = """You are a legal assistant providing detailed, comprehensive answers about legal documents and policies. Your role is to analyze legal documents thoroughly and provide informative, well-structured responses.

RESPONSE REQUIREMENTS:
1. **Provide Detailed Analysis**: Give comprehensive explanations, not brief summaries. Include relevant details, context, and implications.
2. **Use All Available Context**: Analyze ALL provided documents and sources. Don't limit yourself to just one source when multiple are available.
3. **Structure Your Response**: Use clear paragraphs, bullet points when appropriate, and logical organization.
4. **Answer Follow-up Questions**: If the user asks about a document you just cited, provide detailed information from that document.
5. **Cite Properly**: Use document names as shown in context (e.g., [document_name.pdf] or [document_name.pdf (Page X)]).
6. **Be Conversational and Context-Aware**: 
   - Always acknowledge and reference previous parts of our conversation when relevant
   - If the user asks about something mentioned earlier, explicitly connect it to the previous discussion
   - Build upon previous answers and maintain conversational flow
   - Reference earlier questions and answers when they provide context for the current question

RESPONSE STYLE: {response_style}
- Concise: Provide key points only
- Balanced: Structured overview with main points
- Detailed: Comprehensive analysis

CONVERSATION HISTORY (ALWAYS consider this context):
{conversation_history}

AVAILABLE LEGAL DOCUMENTS AND CONTEXT:
{context}

USER QUESTION:
{questions}

Provide a comprehensive, detailed response using the available legal documents AND the conversation history above. If this question relates to something we discussed earlier, acknowledge that connection. Include specific provisions, explanations, and relevant details from the sources. If multiple documents are relevant, synthesize information from all of them. Aim for thorough, informative responses that fully address the user's question while maintaining awareness of our ongoing conversation.

RESPONSE:"""

def process_query_improved(question: str, session_id: str, response_style: str = "balanced") -> QueryResponse:
    """Improved query processing matching second backend behavior"""
    try:
        # Load Database
        db = load_database()
        
        # Test database connectivity
        test_results = db.similarity_search("test", k=1)
        logger.info(f"Database loaded successfully with {len(test_results)} test results")
        
        # Parse Question
        questions = parse_multiple_questions(question)
        logger.info(f"Parsed {len(questions)} questions: {questions}")
        combined_query = " ".join(questions)
        
        # Get Conversation History
        conversation_history_list = conversations.get(session_id, {}).get('messages', [])
        conversation_history_context = get_conversation_context(session_id, max_messages=12)
        logger.info(f"Using conversation history with {len(conversation_history_list)} total messages")
        
        # Perform Retrieval - using modified function
        results, search_result = enhanced_retrieval_v2(db, combined_query, conversation_history_context, k=8)
        logger.info(f"Retrieved {len(results)} results")
        
        if not results:
            logger.warning("No relevant documents found")
            no_info_response = "I couldn't find any relevant information in the documents to answer your question. Please try rephrasing your question or check if the documents contain information about this topic."
            add_to_conversation(session_id, "assistant", no_info_response)
            return QueryResponse(
                response=no_info_response,
                error=None,
                context_found=False,
                sources=[],
                session_id=session_id,
                confidence_score=0.0,
                expand_available=False
            )
        
        # Create Context
        context_text, source_info = create_enhanced_context(results, search_result, questions)
        logger.info(f"Created context with {len(source_info)} sources")
        
        # Calculate confidence before generating response (keeping this feature from first backend)
        confidence_score = calculate_confidence_score(results, search_result, len(context_text))
        
        # Format Prompt
        if len(questions) > 1:
            formatted_questions = "\n".join(f"{i+1}. {q}" for i, q in enumerate(questions))
        else:
            formatted_questions = questions[0]
        
        formatted_prompt = IMPROVED_PROMPT_TEMPLATE.format(
            response_style=response_style.capitalize(),
            conversation_history=conversation_history_context if conversation_history_context else "No previous conversation in this session.",
            context=context_text,
            questions=formatted_questions
        )
        logger.info(f"Prompt length: {len(formatted_prompt)} characters")
        logger.info(f"Conversation history length: {len(conversation_history_context)} characters")
        
        # Call LLM
        api_key = os.environ.get("OPENAI_API_KEY")
        api_base = os.environ.get("OPENAI_API_BASE", "https://openrouter.ai/api/v1")
        
        if not api_key:
            error_msg = "OPENAI_API_KEY environment variable not set."
            logger.error(error_msg)
            add_to_conversation(session_id, "assistant", f"Configuration Error: {error_msg}")
            return QueryResponse(
                response=None,
                error=f"Configuration Error: {error_msg}",
                context_found=True,
                sources=source_info,
                session_id=session_id,
                confidence_score=0.0,
                expand_available=False
            )
        
        response_text = call_openrouter_api(formatted_prompt, api_key, api_base)
        if not response_text:
            response_text = "I received an empty response. Please try again."
        
        # Post-process Response (Add Sources) - matching second backend exactly
        MIN_RELEVANCE_SCORE_FOR_SOURCE_DISPLAY = 0.25
        relevant_source_info = [source for source in source_info if source['relevance'] >= MIN_RELEVANCE_SCORE_FOR_SOURCE_DISPLAY]
        
        # Add sources section only if there are relevant sources
        if relevant_source_info:
            response_text += "\n\n**SOURCES:**\n"
            for source in relevant_source_info:
                page_info = f", Page {source['page']}" if source['page'] is not None else ""
                response_text += f"- {source['file_name']}{page_info} (Relevance: {source['relevance']:.2f})\n"
        
        # Keep the style formatting from first backend (it's a nice feature)
        formatted_response, expand_available = format_response_by_style(response_text, source_info, response_style)
        
        # Update Conversation History
        add_to_conversation(session_id, "user", question)  # Add original question
        add_to_conversation(session_id, "assistant", formatted_response, source_info)  # Add full source info to history
        logger.info(f"Successfully generated response of length {len(formatted_response)}")
        
        return QueryResponse(
            response=formatted_response,
            error=None,
            context_found=True,
            sources=relevant_source_info,  # Return the filtered list for the API response
            session_id=session_id,
            confidence_score=confidence_score,
            expand_available=expand_available
        )
        
    except Exception as e:
        logger.error(f"Failed to load database or process query: {e}", exc_info=True)
        error_msg = f"Failed to process your request: {str(e)}"
        add_to_conversation(session_id, "assistant", error_msg)
        return QueryResponse(
            response=None,
            error=error_msg,
            context_found=False,
            sources=[],
            session_id=session_id,
            confidence_score=0.0,
            expand_available=False
        )

# --- API Endpoints ---

# RAG Endpoints
@app.post("/ask", response_model=QueryResponse)
async def ask_question_improved(query: Query):
    """Improved question endpoint with enhanced accuracy and user experience"""
    cleanup_expired_conversations()
    
    session_id = query.session_id or str(uuid.uuid4())
    if session_id not in conversations:
        conversations[session_id] = {
            "messages": [],
            "created_at": datetime.utcnow(),
            "last_accessed": datetime.utcnow()
        }
    else:
        conversations[session_id]["last_accessed"] = datetime.utcnow()
    
    user_question = query.question.strip()
    if not user_question:
        return QueryResponse(
            response=None,
            error="Question cannot be empty.",
            context_found=False,
            sources=[],
            session_id=session_id,
            confidence_score=0.0,
            expand_available=False
        )
    
    response = process_query_improved(user_question, session_id, query.response_style)
    return response

@app.get("/conversation/{session_id}", response_model=ConversationHistory)
def get_conversation_history(session_id: str):
    """Get conversation history for a session"""
    if session_id not in conversations:
        raise HTTPException(status_code=404, detail="Conversation not found")
    conv_data = conversations[session_id]
    messages = conv_data.get('messages', [])
    return ConversationHistory(
        session_id=session_id,
        messages=messages,
        created_at=conv_data['created_at'].isoformat() if 'created_at' in conv_data else datetime.utcnow().isoformat(),
        last_updated=conv_data['last_accessed'].isoformat() if 'last_accessed' in conv_data else datetime.utcnow().isoformat()
    )

@app.delete("/conversation/{session_id}")
def clear_conversation(session_id: str):
    """Clear conversation history for a session"""
    if session_id in conversations:
        del conversations[session_id]
        return {"message": f"Conversation {session_id} cleared"}
    else:
        raise HTTPException(status_code=404, detail="Conversation not found")

@app.get("/conversations")
def list_conversations():
    """List all active conversations"""
    return {
        "active_conversations": len(conversations),
        "sessions": [
            {
                "session_id": session_id,
                "message_count": len(data.get('messages', [])),
                "created_at": data.get('created_at', datetime.min).isoformat(),
                "last_updated": data.get('last_accessed', datetime.min).isoformat()
            }
            for session_id, data in conversations.items()
        ]
    }

# Document Analysis Endpoints
@app.post("/document-analysis", response_model=EnhancedAnalysisResponse)
async def unified_document_analysis(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    analysis_type: Optional[str] = Form("summarize")
):
    """Unified document analysis with both fact extraction and AI insights"""
    
    session_id = session_id or str(uuid.uuid4())
    
    try:
        logger.info(f"Processing document: {file.filename}, type: {analysis_type}, session: {session_id}")
        
        document_text, file_type, processing_info = SafeDocumentProcessor.process_document_safe(file)
        
        extraction_results = safe_analyzer.extract_document_facts(document_text)
        factual_summary = safe_analyzer.generate_factual_summary(document_text)
        
        # Determine which AI model to use
        ai_model_used = "fact-extraction-only"
        
        if AI_ENABLED and analysis_type != "fact-extraction-only":
            ai_analysis, ai_confidence = await perform_ai_analysis(document_text, analysis_type)
            
            # Check which model was actually used based on confidence
            if ai_confidence > 0.8:
                ai_model_used = "deepseek-chat"
            elif ai_confidence > 0.5:
                ai_model_used = "open-source-nlp"
            
            combined_summary = f"""## AI Legal Analysis: {analysis_type.replace('-', ' ').title()}

{ai_analysis}

---

## Verified Facts from Document

{factual_summary}

---

**⚠️ DISCLAIMER**: 
- The analysis above is generated by AI and should be reviewed carefully
- The verified facts section contains only information extracted directly from the document
- This analysis is for informational purposes only and does not constitute legal advice
- Always consult with a qualified attorney for legal matters
"""
        else:
            ai_analysis = None
            ai_confidence = 0.0
            combined_summary = f"""## Document Analysis: Fact Extraction Only

{f'**Note**: AI analysis not available. Showing only verified facts extracted from the document.' if analysis_type != 'fact-extraction-only' else ''}

{factual_summary}

---

**To enable AI-powered analysis**:
1. Set the OPENAI_API_KEY environment variable
2. Install aiohttp: pip install aiohttp
3. For open-source models: pip install transformers torch
4. Restart the server
"""
        
        successful_extractions = len([k for k in extraction_results.keys() 
                                    if k not in ['extraction_status', 'timestamp'] and 
                                    not (isinstance(extraction_results[k], list) and 
                                         extraction_results[k] and 
                                         extraction_results[k][0].get('status') == 'failed_to_extract')])
        
        if AI_ENABLED and ai_confidence > 0.7 and successful_extractions >= 3:
            verification_status = "high_confidence"
        elif (AI_ENABLED and ai_confidence > 0.5) or successful_extractions >= 1:
            verification_status = "medium_confidence"
        else:
            verification_status = "low_confidence"
        
        logger.info(f"Analysis completed for {file.filename}: {verification_status}")
        
        return EnhancedAnalysisResponse(
            response=combined_summary,
            summary=combined_summary,
            factual_summary=combined_summary,
            ai_analysis=ai_analysis,
            extraction_results=extraction_results,
            analysis_type=analysis_type,
            confidence_score=ai_confidence if AI_ENABLED else 0.5,
            processing_info=processing_info,
            verification_status=verification_status,
            status="completed",
            success=True,
            warnings=processing_info.get('warnings', []),
            session_id=session_id,
            timestamp=datetime.utcnow().isoformat(),
            model_used=ai_model_used
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document analysis failed: {type(e).__name__}: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        error_summary = f"""## Analysis Failed

**Error**: {type(e).__name__}: {str(e)}

The document could not be analyzed. Please check:
1. The file is a valid PDF, DOCX, or TXT document
2. The file is not corrupted
3. The file size is under 10MB

If the problem persists, please try again or contact support.
"""
        
        return EnhancedAnalysisResponse(
            response=error_summary,
            summary=error_summary,
            factual_summary=error_summary,
            ai_analysis=None,
            extraction_results=None,
            analysis_type=analysis_type,
            confidence_score=0.0,
            processing_info={"error": str(e), "error_type": type(e).__name__},
            verification_status="failed",
            status="failed",
            success=False,
            warnings=[],
            session_id=session_id,
            timestamp=datetime.utcnow().isoformat()
        )

@app.post("/verify-extraction")
async def verify_extraction(
    file: UploadFile = File(...),
    claim: str = Form(...),
    line_number: Optional[int] = Form(None)
):
    """Verify a specific claim against the document"""
    
    try:
        document_text, _, _ = SafeDocumentProcessor.process_document_safe(file)
        lines = document_text.split('\n')
        
        verification_result = {
            "claim": claim,
            "verification_status": "not_found",
            "evidence": None,
            "line_number": line_number,
            "context": None
        }
        
        if line_number and 1 <= line_number <= len(lines):
            target_line = lines[line_number - 1]
            if claim.lower() in target_line.lower():
                verification_result.update({
                    "verification_status": "verified",
                    "evidence": target_line.strip(),
                    "context": target_line.strip()
                })
            else:
                verification_result.update({
                    "verification_status": "not_found_at_line",
                    "context": target_line.strip()
                })
        else:
            for i, line in enumerate(lines, 1):
                if claim.lower() in line.lower():
                    verification_result.update({
                        "verification_status": "verified",
                        "evidence": line.strip(),
                        "line_number": i,
                        "context": line.strip()
                    })
                    break
        
        return verification_result
        
    except Exception as e:
        return {
            "claim": claim,
            "verification_status": "error",
            "error": str(e)
        }

# Health and Status Endpoints
@app.get("/health")
def unified_health_check():
    """Comprehensive health check for the unified system"""
    
    # Check ChromaDB
    db_exists = os.path.exists(CHROMA_PATH)
    db_status = "healthy" if db_exists else "not_found"
    
    try:
        if db_exists:
            db = load_database()
            test_results = db.similarity_search("test", k=1)
            db_status = "healthy" if test_results is not None else "error"
    except Exception as e:
        db_status = f"error: {str(e)}"
    
    # Return format that the frontend expects
    # The frontend checks for 'ai_enabled' and 'database_exists' to determine backend type
    return {
        "status": "healthy",
        "version": "6.0.0-Unified",
        "timestamp": datetime.utcnow().isoformat(),
        # Include both formats for compatibility
        "ai_enabled": AI_ENABLED,  # For document analysis detection
        "database_exists": db_exists,  # For RAG detection
        "database_path": CHROMA_PATH,
        "api_key_configured": bool(OPENROUTER_API_KEY),
        "active_conversations": len(conversations),
        # Unified system info
        "unified_mode": True,
        "components": {
            "rag_system": {
                "enabled": db_exists,
                "database_status": db_status,
                "database_path": CHROMA_PATH,
                "nlp_model": nlp is not None,
                "sentence_model": sentence_model is not None,
                "active_conversations": len(conversations)
            },
            "document_analysis": {
                "enabled": True,
                "ai_enabled": AI_ENABLED,
                "ai_model": "deepseek-chat" if AI_ENABLED else "not-configured",
                "open_source_nlp": {
                    "available": OPEN_SOURCE_NLP_AVAILABLE,
                    "models_loaded": open_source_analyzer.models_loaded if open_source_analyzer else False,
                    "models": ["legal-bert", "bart-summarization", "roberta-qa", "zero-shot-classification"] if open_source_analyzer and open_source_analyzer.models_loaded else []
                },
                "pdf_processors": {
                    "pymupdf": PYMUPDF_AVAILABLE,
                    "pdfplumber": PDFPLUMBER_AVAILABLE,
                    "pypdf2": True
                }
            },
            "api_configuration": {
                "openai_key_set": bool(OPENROUTER_API_KEY),
                "aiohttp_available": AIOHTTP_AVAILABLE,
                "api_base": OPENAI_API_BASE
            }
        },
        "features": [
            "✅ RAG-based Q&A with ChromaDB",
            "✅ Document upload and analysis",
            "✅ Zero-hallucination fact extraction",
            f"{'✅' if AI_ENABLED else '❌'} AI-powered legal analysis (DeepSeek)",
            f"{'✅' if OPEN_SOURCE_NLP_AVAILABLE else '❌'} Open-source NLP models (BERT/BART)",
            "✅ Conversation history management",
            "✅ Multi-format document support (PDF, DOCX, TXT)"
        ]
    }

@app.get("/debug-db")
def debug_database():
    """Debug endpoint to check database status"""
    if not os.path.exists(CHROMA_PATH):
        return {"error": "Database folder does not exist", "path": CHROMA_PATH}
    
    try:
        db_contents = os.listdir(CHROMA_PATH)
        db = load_database()
        
        test_queries = ["test", "document", "content", "information"]
        search_results = {}
        for query in test_queries:
            try:
                results = db.similarity_search(query, k=3)
                previews = [doc.page_content[:100] + "..." for doc in results]
                search_results[query] = {
                    "count": len(results),
                    "previews": previews
                }
            except Exception as e:
                search_results[query] = {"error": str(e)}
        
        status = "Database appears to be working" if any(r.get("count", 0) > 0 for r in search_results.values()) else "Database exists but no search results found"
        
        return {
            "database_exists": True,
            "database_path": CHROMA_PATH,
            "database_contents": db_contents,
            "search_tests": search_results,
            "status": status
        }
    except Exception as e:
        logger.error(f"Database debug failed: {e}", exc_info=True)
        return {
            "error": f"Database test failed: {str(e)}",
            "path": CHROMA_PATH,
            "suggestion": "Try running the ingestion script to recreate the database"
        }

@app.get("/extraction-capabilities")
def get_extraction_capabilities():
    """Get what the system can and cannot extract"""
    
    capabilities = {
        "rag_capabilities": {
            "description": "Retrieval-Augmented Generation for Q&A",
            "features": [
                "Search through ingested legal documents",
                "Multi-query search strategies",
                "Conversation history context",
                "Source citation with relevance scores",
                "Response style customization (concise/balanced/detailed)"
            ]
        },
        "fact_extraction": {
            "dates": "Only dates with clear context (due dates, deadlines, etc.)",
            "monetary_amounts": "Only amounts with payment context",
            "percentages": "Only percentages with rate/fee context",
            "party_names": "Only clearly identified parties",
            "document_structure": "Numbered sections, headers, basic organization",
            "basic_statistics": "Word count, paragraph count, reading time"
        },
        "ai_analysis": {
            "status": "enabled" if AI_ENABLED else "disabled",
            "model": "deepseek-chat" if AI_ENABLED else "not-configured",
            "capabilities": [
                "Legal document summarization",
                "Key clause extraction",
                "Missing clause detection",
                "Risk assessment and flagging",
                "Timeline and deadline extraction",
                "Party obligation analysis"
            ] if AI_ENABLED else ["AI features disabled - set OPENAI_API_KEY to enable"]
        },
        "verification_required": "All fact extractions include line numbers for manual verification",
        "fallback_behavior": "Returns 'Failed to extract' instead of guessing"
    }
    
    return capabilities

@app.get("/", response_class=HTMLResponse)
def get_unified_interface():
    """Unified interface for the combined system"""
    
    ai_status = "✅ AI Analysis Enabled" if AI_ENABLED else "❌ AI Analysis Disabled"
    db_status = "✅ Connected" if os.path.exists(CHROMA_PATH) else "❌ Not Found"
    
    ai_instructions = "" if AI_ENABLED else """
            <div class="warning">
                <h3>🤖 Enable AI Analysis</h3>
                <p>To enable AI-powered legal analysis:</p>
                <ol>
                    <li>Set environment variable: <code>export OPENAI_API_KEY="your-key"</code></li>
                    <li>Install aiohttp: <code>pip install aiohttp</code></li>
                    <li>Restart the server</li>
                </ol>
            </div>
    """
    
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Unified Legal Assistant System</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; background: #f8f9fa; }}
            .container {{ max-width: 1000px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
            h1 {{ color: #2c3e50; text-align: center; }}
            .status-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; margin: 20px 0; }}
            .status-badge {{ padding: 5px 15px; border-radius: 20px; font-size: 14px; font-weight: bold; display: inline-block; }}
            .ai-enabled {{ background: #d4edda; color: #155724; }}
            .ai-disabled {{ background: #f8d7da; color: #721c24; }}
            .db-connected {{ background: #d4edda; color: #155724; }}
            .db-disconnected {{ background: #f8d7da; color: #721c24; }}
            .feature {{ background: #ecf0f1; padding: 15px; margin: 10px 0; border-radius: 5px; }}
            .capability {{ background: #e8f5e8; padding: 10px; margin: 5px 0; border-left: 4px solid #27ae60; }}
            .ai-feature {{ background: #e3f2fd; padding: 10px; margin: 5px 0; border-left: 4px solid #2196f3; }}
            .warning {{ background: #fff3cd; padding: 15px; margin: 10px 0; border-radius: 5px; border-left: 4px solid #ffc107; }}
            code {{ background: #f1f1f1; padding: 2px 5px; border-radius: 3px; font-family: monospace; }}
            .endpoint {{ background: #f8f9fa; padding: 10px; margin: 5px 0; border-left: 4px solid #6c757d; font-family: monospace; }}
            .system-card {{ background: #fff; border: 1px solid #dee2e6; border-radius: 8px; padding: 20px; margin: 15px 0; }}
            .system-title {{ font-size: 20px; font-weight: bold; color: #495057; margin-bottom: 10px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>⚖️ Unified Legal Assistant System</h1>
            <p style="text-align: center; color: #6c757d;">Combined RAG Q&A and Document Analysis Platform</p>
            
            <div class="status-grid">
                <div style="text-align: center;">
                    <span class="status-badge {'ai-enabled' if AI_ENABLED else 'ai-disabled'}">{ai_status}</span>
                </div>
                <div style="text-align: center;">
                    <span class="status-badge {'db-connected' if os.path.exists(CHROMA_PATH) else 'db-disconnected'}">ChromaDB: {db_status}</span>
                </div>
            </div>
            
            {ai_instructions}
            
            <div class="system-card">
                <div class="system-title">📚 RAG Q&A System</div>
                <p>Ask questions about your ingested legal documents</p>
                <div class="feature">
                    <h4>Features:</h4>
                    <div class="capability">Search through ingested documents with semantic similarity</div>
                    <div class="capability">Multi-turn conversation support with history</div>
                    <div class="capability">Source citations with relevance scores</div>
                    <div class="capability">Response style customization (concise/balanced/detailed)</div>
                </div>
                <div class="feature">
                    <h4>Endpoints:</h4>
                    <div class="endpoint">POST /ask - Ask questions about documents</div>
                    <div class="endpoint">GET /conversation/{{session_id}} - Get conversation history</div>
                    <div class="endpoint">DELETE /conversation/{{session_id}} - Clear conversation</div>
                    <div class="endpoint">GET /conversations - List all conversations</div>
                </div>
            </div>
            
            <div class="system-card">
                <div class="system-title">📄 Document Analysis System</div>
                <p>Upload and analyze individual legal documents</p>
                <div class="feature">
                    <h4>Fact Extraction (Always Available):</h4>
                    <div class="capability">Extract dates with clear context</div>
                    <div class="capability">Find monetary amounts with payment context</div>
                    <div class="capability">Identify percentages with rate/fee context</div>
                    <div class="capability">Locate party names when clearly identified</div>
                    <div class="capability">Analyze document structure</div>
                    <div class="capability">Provide line numbers for all extracted information</div>
                </div>
                
                <div class="feature">
                    <h4>AI Analysis {"(Active)" if AI_ENABLED else "(Inactive)"}:</h4>
                    <div class="ai-feature">Legal document summarization</div>
                    <div class="ai-feature">Key clause extraction and analysis</div>
                    <div class="ai-feature">Missing clause detection</div>
                    <div class="ai-feature">Legal risk assessment</div>
                    <div class="ai-feature">Timeline and deadline extraction</div>
                    <div class="ai-feature">Party obligation analysis</div>
                </div>
                
                <div class="feature">
                    <h4>Analysis Types:</h4>
                    <ul>
                        <li><code>summarize</code> - Comprehensive document summary</li>
                        <li><code>extract-clauses</code> - Extract key legal clauses</li>
                        <li><code>missing-clauses</code> - Identify missing standard clauses</li>
                        <li><code>risk-flagging</code> - Flag potential legal risks</li>
                        <li><code>timeline-extraction</code> - Extract all dates and deadlines</li>
                        <li><code>obligations</code> - List party obligations</li>
                        <li><code>fact-extraction-only</code> - Only extract verifiable facts (no AI)</li>
                    </ul>
                </div>
                
                <div class="feature">
                    <h4>Endpoints:</h4>
                    <div class="endpoint">POST /document-analysis - Analyze uploaded document</div>
                    <div class="endpoint">POST /verify-extraction - Verify specific claims</div>
                </div>
            </div>
            
            <div class="feature">
                <h3>🔧 System Endpoints</h3>
                <div class="endpoint">GET /health - System health check</div>
                <div class="endpoint">GET /debug-db - Database debugging info</div>
                <div class="endpoint">GET /extraction-capabilities - View all capabilities</div>
            </div>
            
            <div class="feature">
                <h3>📋 Supported File Types</h3>
                <p>• PDF (with PyMuPDF, pdfplumber, or PyPDF2)</p>
                <p>• DOCX (Word documents)</p>
                <p>• TXT (Plain text)</p>
            </div>
            
            <p style="text-align: center; color: #7f8c8d; margin-top: 30px;">
                {"Powered by DeepSeek AI via OpenRouter 🚀" if AI_ENABLED else "Configure AI for enhanced analysis 🔧"}
                <br>Version 6.0.0-Unified
            </p>
        </div>
    </body>
    </html>
    """

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    logger.info(f"🚀 Starting Unified Legal Assistant on port {port}")
    logger.info(f"ChromaDB Path: {CHROMA_PATH}")
    logger.info(f"AI Status: {'ENABLED with DeepSeek' if AI_ENABLED else 'DISABLED - Set OPENAI_API_KEY to enable'}")
    logger.info(f"PDF processing: PyMuPDF={PYMUPDF_AVAILABLE}, pdfplumber={PDFPLUMBER_AVAILABLE}")
    uvicorn.run(app, host="0.0.0.0", port=port)
