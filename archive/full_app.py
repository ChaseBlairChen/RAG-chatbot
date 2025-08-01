# Unified Legal Assistant Backend - Enhanced RAG version
# This version uses the more creative RAG approach from the second app

from fastapi import FastAPI, Request, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
import os
import json
import requests
import re
import logging
import uuid
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Tuple, Set, Any
from dataclasses import dataclass
from enum import Enum
import io
import tempfile
import sys
import traceback

# Third-party library imports for RAG
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
import spacy
from sentence_transformers import SentenceTransformer
import numpy as np
from chromadb.config import Settings

# AI imports
try:
    import aiohttp
    AIOHTTP_AVAILABLE = True
except ImportError:
    AIOHTTP_AVAILABLE = False
    print("⚠️ aiohttp not available - AI features disabled. Install with: pip install aiohttp")

# Import open-source NLP models support
OPEN_SOURCE_NLP_AVAILABLE = False
try:
    import torch
    from transformers import (
        pipeline, 
        AutoTokenizer, 
        AutoModelForSequenceClassification,
        AutoModelForTokenClassification,
        AutoModelForQuestionAnswering,
        T5ForConditionalGeneration,
        T5Tokenizer
    )
    OPEN_SOURCE_NLP_AVAILABLE = True
    print("✅ Open-source NLP models available (transformers + torch)")
except ImportError:
    print("⚠️ Open-source NLP not available. Install with: pip install transformers torch")

# Set up logging FIRST
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Configuration ---
CHROMA_PATH = os.path.abspath(os.path.join(os.getcwd(), "chromadb-database"))
logger.info(f"Using CHROMA_PATH: {CHROMA_PATH}")

CHROMA_CLIENT_SETTINGS = Settings(
    persist_directory=CHROMA_PATH,
    anonymized_telemetry=False,
    allow_reset=True,
    is_persistent=True
)

# OpenRouter configuration
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_API_KEY = os.environ.get("OPENAI_API_KEY", "")
OPENAI_API_BASE = os.environ.get("OPENAI_API_BASE", "https://openrouter.ai/api/v1")
AI_ENABLED = AIOHTTP_AVAILABLE and bool(OPENROUTER_API_KEY)

if AI_ENABLED:
    print("✅ AI features enabled with OpenRouter/DeepSeek")
else:
    print("⚠️ AI features disabled - using fact extraction only")

# In-memory conversation storage
conversations: Dict[str, Dict] = {}

# Load NLP Models
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("spaCy model loaded successfully")
except Exception as e:
    logger.error(f"Failed to load spaCy model: {e}")
    nlp = None

try:
    sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
    logger.info("Sentence transformer loaded successfully")
except Exception as e:
    logger.error(f"Failed to load Sentence Transformer model: {e}")
    sentence_model = None

# Import document processing libraries
print("Starting document processing imports...")
PYMUPDF_AVAILABLE = False
PDFPLUMBER_AVAILABLE = False

# Import PyPDF2 (always needed as fallback)
try:
    import PyPDF2
    print("✅ PyPDF2 imported successfully")
except ImportError as e:
    print(f"❌ CRITICAL: PyPDF2 import failed: {e}")
    print("Install with: pip install PyPDF2")
    sys.exit(1)

# Import python-docx
try:
    import docx
    print("✅ python-docx imported successfully")
except ImportError as e:
    print(f"❌ CRITICAL: python-docx import failed: {e}")
    print("Install with: pip install python-docx")
    sys.exit(1)

# Try to import PyMuPDF
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    print("✅ PyMuPDF available - using high-quality PDF processing")
except ImportError as e:
    print(f"⚠️ PyMuPDF not available: {e}")
    print("Install with: pip install PyMuPDF")

# Try to import pdfplumber
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    print("✅ pdfplumber available - using enhanced PDF extraction")
except ImportError as e:
    print(f"⚠️ pdfplumber not available: {e}")
    print("Install with: pip install pdfplumber")

print(f"PDF processing status: PyMuPDF={PYMUPDF_AVAILABLE}, pdfplumber={PDFPLUMBER_AVAILABLE}")

# Create FastAPI app
app = FastAPI(
    title="Unified Legal Assistant API",
    description="Combined RAG Q&A and Document Analysis System with Enhanced Creative Responses",
    version="7.0.0-Enhanced"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Pydantic Models ---

# RAG Models
class Query(BaseModel):
    question: str
    session_id: Optional[str] = None
    response_style: Optional[str] = "balanced"  # "concise", "balanced", "detailed"

class QueryResponse(BaseModel):
    response: Optional[str] = None
    error: Optional[str] = None
    context_found: bool = False
    sources: Optional[list] = None
    session_id: str
    confidence_score: float = 0.0
    expand_available: bool = False

class ConversationHistory(BaseModel):
    session_id: str
    messages: List[Dict]
    created_at: str
    last_updated: str

# Document Analysis Models
class EnhancedAnalysisResponse(BaseModel):
    model_config = {"protected_namespaces": ()}  # Fix for Pydantic warning
    
    response: Optional[str] = None
    summary: Optional[str] = None
    factual_summary: Optional[str] = None
    ai_analysis: Optional[str] = None
    extraction_results: Optional[Dict[str, Any]] = None
    analysis_type: str
    confidence_score: float
    processing_info: Optional[Dict[str, Any]] = None
    verification_status: str
    status: str = "completed"
    success: bool = True
    warnings: List[str] = []
    session_id: str
    timestamp: str
    model_used: str = "deepseek-chat" if AI_ENABLED else "fact-extraction-only"

# --- Enhanced RAG System Functions (From Second App) ---

def cleanup_expired_conversations():
    """Remove conversations older than 1 hour"""
    now = datetime.utcnow()
    expired_sessions = [
        session_id for session_id, data in conversations.items()
        if now - data['last_accessed'] > timedelta(hours=1)
    ]
    for session_id in expired_sessions:
        del conversations[session_id]
    if expired_sessions:
        logger.info(f"Cleaned up {len(expired_sessions)} expired conversations")

def load_database():
    """Load the Chroma database"""
    try:
        embedding_function = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        db = Chroma(
            collection_name="default",
            embedding_function=embedding_function,
            persist_directory=CHROMA_PATH,
            client_settings=CHROMA_CLIENT_SETTINGS
        )
        logger.debug("Database loaded successfully")
        return db
    except Exception as e:
        logger.error(f"Failed to load database: {e}")
        raise

def get_conversation_context(session_id: str, max_messages: int = 8) -> str:
    """Get conversation context (shortened for better performance)"""
    if session_id not in conversations:
        return ""
    messages = conversations[session_id]['messages'][-max_messages:]
    context_parts = []
    for msg in messages:
        role = msg['role'].upper()
        content = msg['content'][:400] + "..." if len(msg['content']) > 400 else msg['content']
        context_parts.append(f"{role}: {content}")
    return "\n".join(context_parts) if context_parts else ""

def add_to_conversation(session_id: str, role: str, content: str, sources: Optional[List] = None):
    """Add message to conversation"""
    if session_id not in conversations:
        conversations[session_id] = {
            'messages': [],
            'created_at': datetime.utcnow(),
            'last_accessed': datetime.utcnow()
        }
    message = {
        'role': role,
        'content': content,
        'timestamp': datetime.utcnow().isoformat(),
        'sources': sources or []
    }
    conversations[session_id]['messages'].append(message)
    conversations[session_id]['last_accessed'] = datetime.utcnow()
    if len(conversations[session_id]['messages']) > 20:
        conversations[session_id]['messages'] = conversations[session_id]['messages'][-20:]

def parse_multiple_questions(query_text: str) -> List[str]:
    """Parse multiple questions from input"""
    questions = []
    query_text = query_text.strip()
    
    if '?' in query_text and not query_text.endswith('?'):
        parts = query_text.split('?')
        for part in parts:
            part = part.strip()
            if part:
                questions.append(part + '?')
    else:
        final_question = query_text
        if not final_question.endswith('?') and '?' not in final_question:
            final_question += '?'
        questions = [final_question]
    
    return questions

# --- IMPROVEMENT 1: Enhanced Retrieval System ---
def enhanced_retrieval_v2(db, query_text: str, conversation_history_context: str, k: int = 12) -> Tuple[List, Any]:
    """
    Improved retrieval with better scoring and multi-query approach
    """
    logger.info(f"[ENHANCED_RETRIEVAL] Original query: '{query_text}'")
    
    try:
        # Strategy 1: Direct query
        results_with_scores = db.similarity_search_with_relevance_scores(query_text, k=k)
        
        # Strategy 2: Query expansion with legal terms
        legal_expanded_query = expand_legal_query(query_text)
        if legal_expanded_query != query_text:
            logger.info(f"[ENHANCED_RETRIEVAL] Expanded query: '{legal_expanded_query}'")
            expanded_results = db.similarity_search_with_relevance_scores(legal_expanded_query, k=k//2)
            results_with_scores.extend(expanded_results)
        
        # Strategy 3: Break down complex questions
        sub_queries = extract_sub_queries(query_text)
        for sub_query in sub_queries[:2]:  # Limit to avoid too many queries
            sub_results = db.similarity_search_with_relevance_scores(sub_query, k=k//3)
            results_with_scores.extend(sub_results)
        
        # Remove duplicates and re-rank
        unique_results = remove_duplicate_documents(results_with_scores)
        
        # IMPROVED: Dynamic threshold based on query complexity
        min_threshold = calculate_dynamic_threshold(query_text, unique_results)
        filtered_results = [(doc, score) for doc, score in unique_results if score > min_threshold]
        
        logger.info(f"[ENHANCED_RETRIEVAL] Found {len(unique_results)} unique results, {len(filtered_results)} above threshold {min_threshold:.3f}")
        
        final_results = filtered_results if filtered_results else unique_results[:k//2]
        docs, scores = zip(*final_results) if final_results else ([], [])
        
        return list(docs), {
            "query_used": query_text,
            "scores": list(scores),
            "threshold_used": min_threshold,
            "strategies_used": ["direct", "expanded", "sub_queries"]
        }
        
    except Exception as e:
        logger.error(f"[ENHANCED_RETRIEVAL] Search failed: {e}")
        return [], {"error": str(e)}

def expand_legal_query(query: str) -> str:
    """Expand query with legal synonyms and terms"""
    legal_expansions = {
        "asylum": "asylum refugee protection persecution",
        "immigration": "immigration visa deportation removal",
        "contract": "contract agreement terms conditions breach",
        "criminal": "criminal penal prosecution sentence conviction",
        "civil": "civil tort liability damages compensation",
        "court": "court tribunal judge judicial proceeding",
        "law": "law statute regulation rule code",
        "rights": "rights protections constitutional fundamental"
    }
    
    expanded_terms = []
    query_lower = query.lower()
    
    for term, expansion in legal_expansions.items():
        if term in query_lower:
            expanded_terms.extend(expansion.split())
    
    if expanded_terms:
        return f"{query} {' '.join(set(expanded_terms))}"
    return query

def extract_sub_queries(query: str) -> List[str]:
    """Extract sub-questions from complex queries"""
    sub_queries = []
    
    # Split on common conjunctions
    conjunctions = [" and ", " or ", " also ", " plus ", " additionally "]
    current_query = query
    
    for conj in conjunctions:
        if conj in current_query.lower():
            parts = current_query.split(conj)
            sub_queries.extend([part.strip() for part in parts if len(part.strip()) > 10])
            break
    
    # Extract specific legal concepts
    legal_patterns = [
        r"what (?:is|are) ([^?]+)\??",
        r"how (?:does|do) ([^?]+)\??",
        r"when (?:is|are) ([^?]+)\??",
        r"(?:explain|describe) ([^?]+)\??"
    ]
    
    for pattern in legal_patterns:
        matches = re.findall(pattern, query, re.IGNORECASE)
        for match in matches:
            if len(match.strip()) > 5:
                sub_queries.append(f"what is {match.strip()}?")
    
    return sub_queries[:3]  # Limit to 3 sub-queries

def remove_duplicate_documents(results_with_scores: List[Tuple]) -> List[Tuple]:
    """Remove duplicate documents based on content similarity"""
    if not results_with_scores:
        return []
    
    unique_results = []
    seen_content = set()
    
    for doc, score in results_with_scores:
        # Create a hash of the first 100 characters for deduplication
        content_hash = hash(doc.page_content[:100])
        if content_hash not in seen_content:
            seen_content.add(content_hash)
            unique_results.append((doc, score))
    
    # Sort by relevance score (descending)
    return sorted(unique_results, key=lambda x: x[1], reverse=True)

def calculate_dynamic_threshold(query: str, results: List[Tuple]) -> float:
    """Calculate dynamic threshold based on query complexity and result distribution"""
    if not results:
        return 0.3
    
    scores = [score for _, score in results]
    
    # Query complexity factors
    is_complex = len(query.split()) > 8 or '?' in query[:-1]  # Multiple questions
    has_legal_terms = any(term in query.lower() for term in ['law', 'legal', 'court', 'statute', 'regulation'])
    
    # Statistical analysis of scores
    if len(scores) > 1:
        avg_score = np.mean(scores)
        std_score = np.std(scores)
        
        if is_complex:
            return max(0.25, avg_score - std_score)  # Lower threshold for complex queries
        elif has_legal_terms:
            return max(0.35, avg_score - 0.5 * std_score)  # Medium threshold
        else:
            return max(0.4, avg_score - 0.3 * std_score)  # Higher threshold
    
    return 0.3  # Default fallback

# --- IMPROVEMENT 2: Response Style Management ---
def format_response_by_style(content: str, sources: List[Dict], style: str = "balanced") -> Tuple[str, bool]:
    """Format response based on user's preferred style"""
    
    if style == "concise":
        # Extract key points and create concise response
        concise_response = create_concise_response(content, sources)
        return concise_response, True  # Expansion available
    
    elif style == "detailed":
        # Return full detailed response
        detailed_response = create_detailed_response(content, sources)
        return detailed_response, False  # No expansion needed
    
    else:  # balanced
        # Provide balanced response with clear structure
        balanced_response = create_balanced_response(content, sources)
        return balanced_response, True  # Expansion available

def create_concise_response(content: str, sources: List[Dict]) -> str:
    """Create a concise, bullet-point response"""
    # This is a simplified version - in practice, you'd use NLP to extract key points
    lines = content.split('\n')
    key_points = []
    
    for line in lines[:5]:  # Limit to first 5 lines
        if line.strip() and not line.startswith('#'):
            key_points.append(f"• {line.strip()}")
    
    concise = f"""**Quick Answer:**
{chr(10).join(key_points)}

💡 *Need more details? Ask me to expand on any point above.*"""
    
    return concise

def create_balanced_response(content: str, sources: List[Dict]) -> str:
    """Create a balanced response with clear sections"""
    # Structure the response with clear sections
    if len(content) > 800:
        preview = content[:600] + "..."
        balanced = f"""{preview}

📖 **Want the complete analysis?** Ask me to provide the full detailed response.
🔍 **Have specific questions?** Ask about any particular aspect mentioned above."""
    else:
        balanced = content
    
    return balanced

def create_detailed_response(content: str, sources: List[Dict]) -> str:
    """Return the full detailed response"""
    return content

# --- IMPROVEMENT 3: Confidence Scoring ---
def calculate_confidence_score(results: List, search_result: Dict, response_length: int) -> float:
    """Calculate confidence score based on multiple factors"""
    if not results:
        return 0.1
    
    scores = search_result.get("scores", [])
    if not scores:
        return 0.2
    
    # Factor 1: Average relevance score
    avg_relevance = np.mean(scores)
    
    # Factor 2: Number of supporting documents
    doc_factor = min(1.0, len(results) / 5.0)  # Optimal around 5 documents
    
    # Factor 3: Score distribution (consistency)
    if len(scores) > 1:
        score_std = np.std(scores)
        consistency_factor = max(0.5, 1.0 - score_std)
    else:
        consistency_factor = 0.7
    
    # Factor 4: Response completeness
    completeness_factor = min(1.0, response_length / 500.0)  # Optimal around 500 chars
    
    # Weighted combination
    confidence = (
        avg_relevance * 0.4 +
        doc_factor * 0.3 +
        consistency_factor * 0.2 +
        completeness_factor * 0.1
    )
    
    return min(1.0, max(0.0, confidence))

# --- IMPROVEMENT 4: Enhanced Prompt Template (Creative Version) ---
IMPROVED_PROMPT_TEMPLATE = """You are a legal research assistant. Your responses must be STRICTLY based on the provided legal documents.

CRITICAL REQUIREMENTS:
1. **ONLY use information from the provided context below**
2. **If the context doesn't contain sufficient information, explicitly state this**
3. **Cite specific document names for each claim**
4. **Do NOT add general legal knowledge not found in the context**

RESPONSE STYLE: {response_style}
- Concise: Provide key points only
- Balanced: Structured overview with main points
- Detailed: Comprehensive analysis

CONVERSATION HISTORY:
{conversation_history}

LEGAL DOCUMENT CONTEXT (USE ONLY THIS INFORMATION):
{context}

USER QUESTION:
{questions}

INSTRUCTIONS:
- Base your response ONLY on the provided context
- If context is insufficient, say: "Based on the available documents, I can only provide limited information..."
- Cite document names for each fact: [document_name.pdf]
- If no relevant information exists, say: "The available documents do not contain information about this topic."

RESPONSE:"""

def create_enhanced_context(results: List, search_result: Dict, questions: List[str]) -> Tuple[str, List[Dict]]:
    """Enhanced context creation with better relevance filtering"""
    if not results:
        return "", []
    
    context_parts = []
    source_info = []
    seen_sources = set()
    
    # IMPROVED: Higher minimum threshold for source inclusion
    MIN_RELEVANCE_FOR_CONTEXT = 0.4
    
    for i, (doc, score) in enumerate(zip(results, search_result.get("scores", [0.0]*len(results)))):
        # Skip low-relevance documents
        if score < MIN_RELEVANCE_FOR_CONTEXT:
            continue
            
        content = doc.page_content.strip()
        if not content:
            continue
        
        source_path = doc.metadata.get('source', 'Unknown Source')
        page = doc.metadata.get('page', None)
        
        source_id = (source_path, page)
        if source_id in seen_sources:
            continue
        seen_sources.add(source_id)
        
        display_source = os.path.basename(source_path)
        page_info = f" (Page {page})" if page is not None else ""
        
        # Truncate very long content
        if len(content) > 800:
            content = content[:800] + "... [truncated]"
        
        context_part = f"[{display_source}{page_info}] (Relevance: {score:.2f}):\n{content}"
        context_parts.append(context_part)
        
        source_info.append({
            'id': i+1,
            'file_name': display_source,
            'page': page,
            'relevance': score,
            'full_path': source_path
        })
    
    context_text = "\n\n---\n\n".join(context_parts)
    return context_text, source_info

def call_openrouter_api(prompt: str, api_key: str, api_base: str = "https://openrouter.ai/api/v1") -> str:
    """Call OpenRouter API with fallback models"""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8000",
        "X-Title": "Legal Assistant"
    }
    
    models_to_try = [
        "deepseek/deepseek-chat",
        "microsoft/phi-3-mini-128k-instruct:free",
        "meta-llama/llama-3.2-3b-instruct:free",
        "google/gemma-2-9b-it:free",
        "mistralai/mistral-7b-instruct:free",
        "openchat/openchat-7b:free"
    ]
    
    for model in models_to_try:
        try:
            payload = {
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.5,  # Lower temperature for more consistent responses
                "max_tokens": 2000,
                "top_p": 0.9
            }
            
            response = requests.post(f"{api_base}/chat/completions", headers=headers, json=payload, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and result['choices']:
                    content = result['choices'][0]['message']['content']
                    if content and content.strip():
                        return content.strip()
        except Exception as e:
            logger.error(f"Error with model {model}: {e}")
            continue
    
    return "I apologize, but I'm experiencing technical difficulties. Please try again."

# --- Updated Main Processing Function ---
def process_query_improved(question: str, session_id: str, response_style: str = "balanced") -> QueryResponse:
    """
    Improved query processing with enhanced accuracy and user experience
    """
    try:
        # Load Database
        db = load_database()
        
        # Parse Question
        questions = parse_multiple_questions(question)
        combined_query = " ".join(questions)
        
        # Get Conversation History
        conversation_history_context = get_conversation_context(session_id, max_messages=8)
        
        # IMPROVED: Enhanced Retrieval
        results, search_result = enhanced_retrieval_v2(db, combined_query, conversation_history_context, k=12)
        
        if not results:
            logger.warning("No relevant documents found")
            no_info_response = "I couldn't find any relevant information in the available legal documents to answer your question. The documents may not contain information about this specific topic."
            add_to_conversation(session_id, "assistant", no_info_response)
            return QueryResponse(
                response=no_info_response,
                error=None,
                context_found=False,
                sources=[],
                session_id=session_id,
                confidence_score=0.0,
                expand_available=False
            )
        
        # Create Context with better filtering
        context_text, source_info = create_enhanced_context(results, search_result, questions)
        
        # IMPROVED: Calculate confidence before generating response
        confidence_score = calculate_confidence_score(results, search_result, len(context_text))
        
        # Enhanced Prompt
        formatted_questions = "\n".join(f"{i+1}. {q}" for i, q in enumerate(questions)) if len(questions) > 1 else questions[0]
        
        formatted_prompt = IMPROVED_PROMPT_TEMPLATE.format(
            response_style=response_style.capitalize(),
            conversation_history=conversation_history_context if conversation_history_context else "No previous conversation.",
            context=context_text,
            questions=formatted_questions
        )
        
        # Call LLM
        api_key = os.environ.get("OPENAI_API_KEY")
        api_base = os.environ.get("OPENAI_API_BASE", "https://openrouter.ai/api/v1")
        
        if not api_key:
            error_msg = "OPENAI_API_KEY environment variable not set."
            return QueryResponse(
                response=None,
                error=f"Configuration Error: {error_msg}",
                context_found=True,
                sources=source_info,
                session_id=session_id,
                confidence_score=0.0,
                expand_available=False
            )
        
        raw_response = call_openrouter_api(formatted_prompt, api_key, api_base)
        
        # IMPROVED: Format response based on style
        formatted_response, expand_available = format_response_by_style(raw_response, source_info, response_style)
        
        # Add sources with confidence indicator
        if source_info:
            formatted_response += f"\n\n**SOURCES** (Confidence: {confidence_score:.1%}):\n"
            for source in source_info:
                page_info = f", Page {source['page']}" if source['page'] is not None else ""
                formatted_response += f"- {source['file_name']}{page_info} (Relevance: {source['relevance']:.2f})\n"
        
        # Update conversation
        add_to_conversation(session_id, "user", question)
        add_to_conversation(session_id, "assistant", formatted_response, source_info)
        
        return QueryResponse(
            response=formatted_response,
            error=None,
            context_found=True,
            sources=source_info,
            session_id=session_id,
            confidence_score=confidence_score,
            expand_available=expand_available
        )
        
    except Exception as e:
        logger.error(f"Query processing failed: {e}", exc_info=True)
        error_msg = f"Failed to process your request: {str(e)}"
        return QueryResponse(
            response=None,
            error=error_msg,
            context_found=False,
            sources=[],
            session_id=session_id,
            confidence_score=0.0,
            expand_available=False
        )

# --- Document Analysis System (from first backend) ---

# AI Analysis prompts
ANALYSIS_PROMPTS = {
    'summarize': """You are a legal document analyst. Analyze this document and provide:
1. A clear summary of the document's purpose and type
2. The main parties involved (with their roles)
3. Key terms and conditions
4. Important dates and deadlines
5. Financial obligations or amounts
6. Any notable risks or concerns

Document text:
{document_text}

Provide a structured summary in plain English while maintaining legal accuracy.""",

    'extract-clauses': """You are a legal document analyst. Extract and categorize the following types of clauses from this document:
1. Termination clauses
2. Indemnification provisions
3. Liability limitations
4. Governing law and jurisdiction
5. Confidentiality/NDA provisions
6. Payment terms
7. Dispute resolution mechanisms

For each clause found, provide:
- Clause type
- Summary of the provision
- Exact location/section reference if available
- Any unusual or concerning aspects

Document text:
{document_text}""",

    'missing-clauses': """You are a legal document analyst. Review this contract and identify commonly expected clauses that appear to be missing or inadequately addressed:

Consider standard clauses such as:
- Force majeure
- Limitation of liability
- Indemnification
- Dispute resolution/arbitration
- Confidentiality
- Termination provisions
- Assignment restrictions
- Severability
- Entire agreement
- Notice provisions
- Governing law

Document text:
{document_text}

For each missing clause, explain why it's typically important and the risks of its absence.""",

    'risk-flagging': """You are a legal risk analyst. Identify and assess legal risks in this document:

Look for:
1. Unilateral termination rights
2. Broad indemnification requirements
3. Unlimited liability exposure
4. Vague or ambiguous obligations
5. Unfavorable payment terms
6. Lack of protection clauses
7. Unusual warranty provisions
8. Problematic intellectual property terms

Document text:
{document_text}

For each risk, provide:
- Risk description
- Severity (High/Medium/Low)
- Potential impact
- Suggested mitigation""",

    'timeline-extraction': """You are a legal document analyst. Extract all time-related information:

Find and list:
1. Contract start and end dates
2. Payment deadlines
3. Notice periods
4. Renewal dates and terms
5. Termination notice requirements
6. Performance deadlines
7. Warranty periods
8. Any other time-sensitive obligations

Document text:
{document_text}

Present as a chronological timeline with clear labels.""",

    'obligations': """You are a legal document analyst. List all obligations and requirements for each party:

Identify:
1. What each party must do
2. When they must do it
3. Conditions or prerequisites
4. Consequences of non-compliance
5. Reporting or notification requirements

Document text:
{document_text}

Organize by party and priority/timeline."""
}

async def perform_ai_analysis(document_text: str, analysis_type: str) -> Tuple[str, float]:
    """Perform AI analysis using DeepSeek via OpenRouter or open-source models as fallback"""
    
    # First try DeepSeek if available
    if AI_ENABLED:
        prompt_template = ANALYSIS_PROMPTS.get(analysis_type, ANALYSIS_PROMPTS['summarize'])
        prompt = prompt_template.format(document_text=document_text[:15000])
        
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:8000",
            "X-Title": "Legal Document Analyzer"
        }
        
        payload = {
            "model": "deepseek/deepseek-chat",
            "messages": [
                {
                    "role": "system",
                    "content": "You are an expert legal document analyst. Provide thorough, accurate analysis while clearly marking any uncertainties. Always include relevant disclaimers about seeking professional legal advice."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,
            "max_tokens": 2000
        }
        
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(OPENROUTER_API_URL, json=payload, headers=headers) as response:
                    if response.status == 200:
                        data = await response.json()
                        ai_response = data['choices'][0]['message']['content']
                        confidence = 0.85
                        return ai_response, confidence
                    else:
                        error_text = await response.text()
                        logger.error(f"OpenRouter API error: {response.status} - {error_text}")
        except Exception as e:
            logger.error(f"AI analysis exception: {e}")
    
    # If DeepSeek fails or is not available, try open-source models
    if open_source_analyzer and open_source_analyzer.models_loaded:
        logger.info(f"Falling back to open-source models for {analysis_type}")
        try:
            return await open_source_analyzer.analyze_document(document_text, analysis_type)
        except Exception as e:
            logger.error(f"Open-source analysis failed: {e}")
    
    # If both fail, return error message
    return "AI analysis not available. Please set OPENAI_API_KEY environment variable or install transformers for open-source models.", 0.0

class VerifiableExtractor:
    """Extract only verifiable information from documents with source locations"""
    
    def __init__(self):
        self.extraction_patterns = {
            'dates': {
                'patterns': [
                    r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
                    r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
                    r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
                ],
                'context_required': ['due', 'deadline', 'effective', 'expires', 'terminates', 'begins', 'starts', 'ends']
            },
            'monetary_amounts': {
                'patterns': [
                    r'\$\s*[\d,]+(?:\.\d{2})?(?:\s*(?:million|Million|billion|Billion|M|B))?\b',
                    r'\$\s*[\d,]+(?:\.\d{2})?(?:\s*(?:per|each|monthly|annually|yearly))?',
                    r'\b[\d,]+(?:\.\d{2})?\s*(?:USD|usd|dollars?|Dollars?)(?:\s*(?:per|each|monthly|annually|yearly))?\b',
                    r'\b(?:USD|usd)\s*[\d,]+(?:\.\d{2})?\b',
                    r'\b[\d,]+\s*cents?\b'
                ],
                'context_required': ['pay', 'payment', 'fee', 'cost', 'price', 'rent', 'salary', 'compensation', 'amount', 'value', 'worth', 'charge', 'expense']
            },
            'percentages': {
                'patterns': [
                    r'\b\d+(?:\.\d+)?%',
                    r'\b\d+(?:\.\d+)?\s*(?:percent|per cent)\b'
                ],
                'context_required': ['rate', 'interest', 'fee', 'penalty', 'commission', 'tax', 'discount', 'margin', 'share', 'ownership', 'equity', 'royalty']
            },
            'party_names': {
                'patterns': [
                    r'([A-Z][A-Za-z\s&.,\'"-]+?(?:Corporation|Corp\.|Company|Co\.|LLC|L\.L\.C\.|Inc\.|Incorporated|Limited|Ltd\.|LLP|LP|plc|PLC))',
                    r'[Bb]y:\s*/?s?/?\s*([A-Z][A-Za-z\s\.-]+?)(?:\n|,|\s{2,})',
                    r'Name:\s*([A-Z][A-Za-z\s\.-]+?)(?:\n|,|\s{2,})',
                    r'between\s+([A-Z][A-Za-z\s&.,\'"-]+?)\s+(?:and|AND)\s+([A-Z][A-Za-z\s&.,\'"-]+?)(?:\s|,|\()',
                    r'\((?:the\s+)?"?([A-Z][A-Za-z\s&.,\'"-]+?)"?\)(?:\s+(?:and|AND)|,)',
                    r'undersigned,?\s+([A-Z][A-Za-z\s&.,\'"-]+?)(?:,|\s+(?:hereby|certifies|agrees))',
                    r'([A-Z][A-Za-z\s\.-]+?),?\s*(?:P\.E\.|P\.Eng\.|Ph\.D\.|Esq\.|CPA|MBA|MD|JD)',
                    r'(?:party|parties)[:\s]+([A-Z][a-zA-Z\s&.,]+?)(?:\s*(?:and|,|\n))',
                    r'(?:company|corporation|llc|inc\.?)[:\s]*([A-Z][a-zA-Z\s&.,]+?)(?:\s*(?:and|,|\n))'
                ],
                'context_required': []
            }
        }
    
    def extract_with_verification(self, document_text: str, extraction_type: str) -> List[Dict[str, Any]]:
        """Extract information only if it can be verified with high confidence"""
        
        if extraction_type not in self.extraction_patterns:
            return [{"status": "failed", "reason": f"Unknown extraction type: {extraction_type}"}]
        
        pattern_config = self.extraction_patterns[extraction_type]
        extracted_items = []
        seen_values = set()
        
        lines = document_text.split('\n')
        
        for line_num, line in enumerate(lines, 1):
            line_lower = line.lower()
            
            for pattern in pattern_config['patterns']:
                matches = re.finditer(pattern, line, re.IGNORECASE)
                
                for match in matches:
                    if len(match.groups()) > 0:
                        for group_idx in range(1, len(match.groups()) + 1):
                            if match.group(group_idx):
                                extracted_value = match.group(group_idx).strip()
                                extracted_value = re.sub(r'\s+', ' ', extracted_value)
                                extracted_value = extracted_value.strip(' ,.')
                                
                                if len(extracted_value) < 3 or extracted_value.lower() in seen_values:
                                    continue
                                
                                if extraction_type == 'party_names':
                                    if extracted_value.lower() in ['the company', 'the corporation', 'party', 'parties', 'undersigned']:
                                        continue
                                    if not any(word[0].isupper() for word in extracted_value.split() if word):
                                        continue
                                
                                seen_values.add(extracted_value.lower())
                                
                                context_verified = True
                                if pattern_config['context_required']:
                                    context_verified = any(
                                        keyword in line_lower 
                                        for keyword in pattern_config['context_required']
                                    )
                                
                                if context_verified:
                                    extracted_items.append({
                                        "value": extracted_value,
                                        "line_number": line_num,
                                        "context": line.strip(),
                                        "confidence": "high",
                                        "verified": True
                                    })
                                else:
                                    extracted_items.append({
                                        "value": extracted_value,
                                        "line_number": line_num,
                                        "context": line.strip(),
                                        "confidence": "low",
                                        "verified": False,
                                        "reason": "No supporting context found"
                                    })
                    else:
                        extracted_value = match.group(0).strip()
                        
                        if extracted_value.lower() in seen_values:
                            continue
                        
                        seen_values.add(extracted_value.lower())
                        
                        context_verified = True
                        if pattern_config['context_required']:
                            context_verified = any(
                                keyword in line_lower 
                                for keyword in pattern_config['context_required']
                            )
                        
                        if context_verified:
                            extracted_items.append({
                                "value": extracted_value,
                                "line_number": line_num,
                                "context": line.strip(),
                                "confidence": "high",
                                "verified": True
                            })
        
        verified_items = [item for item in extracted_items if item.get('verified', False)]
        verified_items.sort(key=lambda x: x['line_number'])
        
        if not verified_items:
            return [{"status": "failed_to_extract", "reason": f"No verifiable {extraction_type} found in document"}]
        
        return verified_items

class NoHallucinationAnalyzer:
    """Document analyzer that never hallucinates - only reports verifiable facts"""
    
    def __init__(self):
        self.extractor = VerifiableExtractor()
        
        self.document_types = {
            'contract': {
                'required_keywords': ['agreement', 'contract'],
                'supporting_keywords': ['parties', 'consideration', 'terms'],
                'minimum_matches': 2
            },
            'lease': {
                'required_keywords': ['lease'],
                'supporting_keywords': ['tenant', 'landlord', 'rent', 'premises'],
                'minimum_matches': 2
            },
            'employment': {
                'required_keywords': ['employment', 'employee'],
                'supporting_keywords': ['employer', 'salary', 'position', 'job'],
                'minimum_matches': 2
            },
            'policy': {
                'required_keywords': ['policy'],
                'supporting_keywords': ['procedure', 'guidelines', 'rules'],
                'minimum_matches': 1
            },
            'consent': {
                'required_keywords': ['consent'],
                'supporting_keywords': ['authorize', 'permit', 'sec', 'filing', 'qualified person', 'technical report'],
                'minimum_matches': 2
            },
            'nda': {
                'required_keywords': ['confidential', 'non-disclosure', 'nda'],
                'supporting_keywords': ['proprietary', 'secret', 'disclose'],
                'minimum_matches': 2
            },
            'license': {
                'required_keywords': ['license', 'licensing'],
                'supporting_keywords': ['grant', 'rights', 'royalty', 'software', 'intellectual property'],
                'minimum_matches': 2
            },
            'purchase': {
                'required_keywords': ['purchase', 'sale', 'acquisition'],
                'supporting_keywords': ['buyer', 'seller', 'price', 'assets', 'shares'],
                'minimum_matches': 2
            }
        }
    
    def detect_document_type_strict(self, document_text: str) -> Dict[str, Any]:
        """Detect document type only with high confidence"""
        
        text_lower = document_text.lower()
        detection_results = {}
        
        for doc_type, criteria in self.document_types.items():
            score = 0
            matched_keywords = []
            
            required_found = 0
            for keyword in criteria['required_keywords']:
                if keyword in text_lower:
                    required_found += 1
                    matched_keywords.append(keyword)
                    score += 10
            
            supporting_found = 0
            for keyword in criteria['supporting_keywords']:
                if keyword in text_lower:
                    supporting_found += 1
                    matched_keywords.append(keyword)
                    score += 3
            
            total_matches = required_found + supporting_found
            if total_matches >= criteria['minimum_matches'] and required_found > 0:
                detection_results[doc_type] = {
                    'score': score,
                    'matched_keywords': matched_keywords,
                    'required_found': required_found,
                    'supporting_found': supporting_found
                }
        
        if not detection_results:
            return {
                'type': 'unknown',
                'confidence': 0.0,
                'status': 'failed_to_detect',
                'reason': 'No clear document type indicators found',
                'matched_keywords': []
            }
        
        best_type = max(detection_results, key=lambda x: detection_results[x]['score'])
        best_result = detection_results[best_type]
        
        confidence = min(1.0, best_result['score'] / 20)
        
        return {
            'type': best_type,
            'confidence': confidence,
            'status': 'detected' if confidence > 0.7 else 'uncertain',
            'matched_keywords': best_result['matched_keywords'],
            'all_matches': detection_results
        }
    
    def extract_document_facts(self, document_text: str) -> Dict[str, Any]:
        """Extract only verifiable facts from the document"""
        
        facts = {
            'basic_stats': self._get_basic_stats(document_text),
            'dates': self.extractor.extract_with_verification(document_text, 'dates'),
            'monetary_amounts': self.extractor.extract_with_verification(document_text, 'monetary_amounts'),
            'percentages': self.extractor.extract_with_verification(document_text, 'percentages'),
            'party_names': self.extractor.extract_with_verification(document_text, 'party_names'),
            'document_structure': self._analyze_structure(document_text),
            'extraction_status': 'completed',
            'timestamp': datetime.utcnow().isoformat()
        }
        
        return facts
    
    def _get_basic_stats(self, text: str) -> Dict[str, Any]:
        """Get verifiable basic statistics"""
        
        words = text.split()
        sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        return {
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len(paragraphs),
            'character_count': len(text),
            'estimated_reading_time_minutes': round(len(words) / 200, 1)
        }
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure"""
        
        numbered_sections = len(re.findall(r'^\s*\d+\.\s', text, re.MULTILINE))
        lettered_sections = len(re.findall(r'^\s*[a-z]\.\s', text, re.MULTILINE))
        subsections = len(re.findall(r'^\s*\d+\.\d+\s', text, re.MULTILINE))
        
        potential_headers = re.findall(r'^([^:\n]{1,50}):$', text, re.MULTILINE)
        headers = len([h for h in potential_headers if len(h.split()) <= 8])
        
        return {
            'numbered_sections': numbered_sections,
            'lettered_sections': lettered_sections,
            'subsections': subsections,
            'headers': headers,
            'has_clear_structure': numbered_sections > 0 or headers > 2
        }
    
    def generate_factual_summary(self, document_text: str) -> str:
        """Generate a summary using only extracted facts"""
        
        doc_type_result = self.detect_document_type_strict(document_text)
        facts = self.extract_document_facts(document_text)
        
        summary_parts = []
        
        if doc_type_result['status'] == 'detected':
            summary_parts.append(f"**Document Type**: {doc_type_result['type'].title()} (confidence: {doc_type_result['confidence']:.1%})")
            summary_parts.append(f"**Keywords Found**: {', '.join(doc_type_result['matched_keywords'])}")
        else:
            summary_parts.append(f"**Document Type**: Failed to detect - {doc_type_result['reason']}")
        
        stats = facts['basic_stats']
        summary_parts.append(f"\n**Document Statistics**:")
        summary_parts.append(f"• Word count: {stats['word_count']:,}")
        summary_parts.append(f"• Estimated reading time: {stats['estimated_reading_time_minutes']} minutes")
        summary_parts.append(f"• Paragraphs: {stats['paragraph_count']}")
        
        structure = facts['document_structure']
        summary_parts.append(f"\n**Document Structure**:")
        if structure['has_clear_structure']:
            summary_parts.append(f"• Numbered sections: {structure['numbered_sections']}")
            summary_parts.append(f"• Headers found: {structure['headers']}")
            summary_parts.append(f"• Subsections: {structure['subsections']}")
        else:
            summary_parts.append("• No clear structural organization detected")
        
        summary_parts.append(f"\n**Extracted Information**:")
        
        dates = facts['dates']
        if dates and dates[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Dates found**: {len(dates)} verifiable dates")
            for i, date in enumerate(dates[:3], 1):
                summary_parts.append(f"  {i}. {date['value']} (Line {date['line_number']})")
            if len(dates) > 3:
                summary_parts.append(f"  ... and {len(dates) - 3} more dates")
        else:
            summary_parts.append("• **Dates**: Failed to extract any verifiable dates")
        
        amounts = facts['monetary_amounts']
        if amounts and amounts[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Financial amounts**: {len(amounts)} verifiable amounts")
            for i, amount in enumerate(amounts[:3], 1):
                summary_parts.append(f"  {i}. {amount['value']} (Line {amount['line_number']})")
            if len(amounts) > 3:
                summary_parts.append(f"  ... and {len(amounts) - 3} more amounts")
        else:
            summary_parts.append("• **Financial amounts**: Failed to extract any verifiable amounts")
        
        percentages = facts['percentages']
        if percentages and percentages[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Percentages**: {len(percentages)} found")
            for perc in percentages[:3]:
                summary_parts.append(f"  - {perc['value']} (Line {perc['line_number']})")
        else:
            summary_parts.append("• **Percentages**: Failed to extract any verifiable percentages")
        
        parties = facts['party_names']
        if parties and parties[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Parties**: {len(parties)} identified")
            for party in parties[:3]:
                summary_parts.append(f"  - {party['value']} (Line {party['line_number']})")
        else:
            summary_parts.append("• **Parties**: Failed to extract clear party names")
        
        summary_parts.append(f"\n**⚠️ IMPORTANT NOTES**:")
        summary_parts.append("• This analysis only includes information that could be verified directly from the document text")
        summary_parts.append("• All extracted items include line numbers for verification")
        summary_parts.append("• Items marked 'Failed to extract' mean the information was not clearly identifiable")
        summary_parts.append("• For legal advice, consult a qualified attorney")
        
        return '\n'.join(summary_parts)

class SafeDocumentProcessor:
    """Document processor with verification and no hallucination"""
    
    @staticmethod
    def process_document_safe(file: UploadFile) -> Tuple[str, str, Dict[str, Any]]:
        """Process document and provide processing metadata"""
        
        file_extension = file.filename.split('.')[-1].lower() if file.filename else 'unknown'
        processing_info = {
            'original_filename': file.filename,
            'file_extension': file_extension,
            'processing_method': None,
            'success': False,
            'warnings': [],
            'errors': []
        }
        
        try:
            file.file.seek(0)
            
            if file_extension == 'pdf':
                text, warnings = SafeDocumentProcessor._extract_pdf_safe(file)
                processing_info['processing_method'] = 'PDF extraction'
                processing_info['warnings'] = warnings
                
            elif file_extension in ['docx', 'doc']:
                text, warnings = SafeDocumentProcessor._extract_docx_safe(file)
                processing_info['processing_method'] = 'Word document extraction'
                processing_info['warnings'] = warnings
                
            elif file_extension == 'txt':
                content = file.file.read()
                text = content.decode('utf-8')
                processing_info['processing_method'] = 'Plain text'
                
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file_extension}. Supported: PDF, DOCX, TXT"
                )
            
            if not text or len(text.strip()) < 10:
                processing_info['errors'].append("Extracted text is too short or empty")
                raise ValueError("Failed to extract meaningful text from document")
            
            processing_info['success'] = True
            processing_info['extracted_length'] = len(text)
            processing_info['word_count'] = len(text.split())
            
            return text, file_extension, processing_info
            
        except Exception as e:
            processing_info['errors'].append(str(e))
            logger.error(f"Document processing error: {type(e).__name__}: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Document processing failed: {str(e)}")
    
    @staticmethod
    def _extract_pdf_safe(file: UploadFile) -> Tuple[str, List[str]]:
        """Extract PDF text with PyMuPDF first"""
        
        warnings = []
        
        file.file.seek(0)
        pdf_content = file.file.read()
        
        if len(pdf_content) == 0:
            raise ValueError("PDF file is empty")
        
        if PYMUPDF_AVAILABLE:
            try:
                import fitz
                
                pdf_doc = fitz.open(stream=pdf_content, filetype="pdf")
                
                try:
                    if pdf_doc.page_count == 0:
                        raise ValueError("PDF contains no pages")
                    
                    text = ""
                    pages_with_text = 0
                    tables_found = 0
                    total_pages = pdf_doc.page_count
                    
                    for page_num in range(total_pages):
                        try:
                            page = pdf_doc[page_num]
                            
                            page_text = page.get_text()
                            if page_text and page_text.strip():
                                text += f"[Page {page_num + 1}]\n{page_text}\n"
                                pages_with_text += 1
                            else:
                                warnings.append(f"Page {page_num + 1} contains no extractable text")
                            
                            try:
                                tables = page.find_tables()
                                if tables:
                                    for table_num, table in enumerate(tables):
                                        try:
                                            table_data = table.extract()
                                            if table_data:
                                                text += f"\n[TABLE {table_num + 1} FROM PAGE {page_num + 1}]\n"
                                                for row in table_data:
                                                    if row:
                                                        clean_row = [str(cell) if cell else "" for cell in row]
                                                        text += " | ".join(clean_row) + "\n"
                                                text += "[/TABLE]\n\n"
                                                tables_found += 1
                                        except Exception as table_error:
                                            logger.debug(f"Table extraction error: {table_error}")
                            except Exception as e:
                                logger.debug(f"Table detection error: {e}")
                                
                        except Exception as page_error:
                            warnings.append(f"Error processing page {page_num + 1}: {str(page_error)}")
                            logger.warning(f"Page {page_num + 1} error: {page_error}")
                    
                    if pages_with_text == 0:
                        raise ValueError("No readable text found in any PDF pages")
                    
                    if pages_with_text < total_pages:
                        warnings.append(f"Only {pages_with_text} of {total_pages} pages contained extractable text")
                    
                    if tables_found > 0:
                        warnings.append(f"PyMuPDF extracted {tables_found} tables with preserved structure")
                    else:
                        warnings.append("PyMuPDF processed PDF successfully")
                    
                    return text, warnings
                    
                finally:
                    pdf_doc.close()
                
            except Exception as e:
                warnings.append(f"PyMuPDF processing failed: {str(e)}, falling back to pdfplumber")
                logger.warning(f"PyMuPDF error: {e}")
        
        if PDFPLUMBER_AVAILABLE:
            pdf_file = io.BytesIO(pdf_content)
            
            try:
                import pdfplumber
                with pdfplumber.open(pdf_file) as pdf:
                    if len(pdf.pages) == 0:
                        raise ValueError("PDF contains no pages")
                    
                    text = ""
                    pages_with_text = 0
                    tables_found = 0
                    
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"[Page {page_num + 1}]\n{page_text}\n"
                            pages_with_text += 1
                        else:
                            warnings.append(f"Page {page_num + 1} contains no extractable text")
                        
                        try:
                            tables = page.extract_tables()
                            if tables:
                                for table_num, table in enumerate(tables):
                                    text += f"\n[TABLE {table_num + 1} FROM PAGE {page_num + 1}]\n"
                                    for row in table:
                                        if row:
                                            clean_row = [str(cell) if cell else "" for cell in row]
                                            text += " | ".join(clean_row) + "\n"
                                    text += "[/TABLE]\n\n"
                                    tables_found += 1
                        except Exception as e:
                            warnings.append(f"Could not extract tables from page {page_num + 1}: {str(e)}")
                    
                    if pages_with_text == 0:
                        raise ValueError("No readable text found in any PDF pages")
                    
                    if pages_with_text < len(pdf.pages):
                        warnings.append(f"Only {pages_with_text} of {len(pdf.pages)} pages contained extractable text")
                    
                    if tables_found > 0:
                        warnings.append(f"pdfplumber extracted {tables_found} tables with preserved structure")
                    
                    return text, warnings
                    
            except Exception as e:
                warnings.append(f"pdfplumber failed: {str(e)}, using basic PyPDF2")
                logger.warning(f"pdfplumber error: {e}")
        
        pdf_file = io.BytesIO(pdf_content)
        pdf_file.seek(0)
        
        try:
            reader = PyPDF2.PdfReader(pdf_file)
            
            if len(reader.pages) == 0:
                raise ValueError("PDF contains no pages")
            
            text = ""
            pages_with_text = 0
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += f"[Page {page_num + 1}]\n{page_text}\n\n"
                        pages_with_text += 1
                    else:
                        warnings.append(f"Page {page_num + 1} contains no extractable text")
                except Exception as e:
                    warnings.append(f"Error extracting text from page {page_num + 1}: {str(e)}")
            
            if pages_with_text == 0:
                raise ValueError("No readable text found in any PDF pages")
            
            warnings.append("Used basic PDF extraction - complex layouts may not be preserved")
            return text, warnings
            
        except Exception as e:
            logger.error(f"PyPDF2 failed: {e}")
            raise ValueError(f"All PDF extraction methods failed. Last error: {str(e)}")
    
    @staticmethod
    def _extract_docx_safe(file: UploadFile) -> Tuple[str, List[str]]:
        """Extract DOCX text with warnings about potential issues"""
        
        warnings = []
        
        file.file.seek(0)
        docx_content = file.file.read()
        
        if len(docx_content) == 0:
            raise ValueError("Word document is empty")
        
        docx_file = io.BytesIO(docx_content)
        
        try:
            doc = docx.Document(docx_file)
            
            text = ""
            paragraphs_with_text = 0
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text and paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraphs_with_text += 1
            
            tables_found = len(doc.tables)
            if tables_found > 0:
                warnings.append(f"Document contains {tables_found} tables - table content may not be fully extracted")
                
                for table_num, table in enumerate(doc.tables):
                    text += f"\n[TABLE {table_num + 1}]\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        if any(row_text):
                            text += " | ".join(row_text) + "\n"
                    text += "[/TABLE]\n\n"
            
            if paragraphs_with_text == 0:
                raise ValueError("No readable text found in Word document")
            
            if paragraphs_with_text < len(doc.paragraphs):
                warnings.append(f"Some paragraphs contained no text")
            
            return text, warnings
            
        except Exception as e:
            raise ValueError(f"Failed to process Word document: {str(e)}")

# Lightweight Open-Source Legal Analyzer - Single Model Version
class LightweightLegalAnalyzer:
    """Legal document analyzer using only one model at a time to save memory"""
    
    def __init__(self, model_choice="classifier"):
        """
        Initialize with only one model
        model_choice: "summarizer", "classifier", "ner", "qa", or "zero_shot"
        """
        self.device = 0 if torch.cuda.is_available() else -1
        self.models = {}
        self.models_loaded = False
        self.model_choice = model_choice
        
        if OPEN_SOURCE_NLP_AVAILABLE:
            self.load_single_model()
    
    def load_single_model(self):
        """Load only one model based on choice"""
        
        print(f"Loading single model: {self.model_choice}...")
        
        try:
            if self.model_choice == "summarizer":
                # Summarization - Uses ~1.2GB RAM
                self.models['summarizer'] = pipeline(
                    "summarization",
                    model="facebook/bart-large-cnn",
                    device=self.device
                )
                print("✅ BART summarizer loaded")
                
            elif self.model_choice == "classifier":
                # Document Classification - Uses ~400MB RAM
                self.models['classifier'] = pipeline(
                    "text-classification",
                    model="nlpaueb/legal-bert-base-uncased",
                    device=self.device
                )
                print("✅ Legal BERT classifier loaded")
                
            elif self.model_choice == "ner":
                # Named Entity Recognition - Uses ~400MB RAM
                self.models['ner'] = pipeline(
                    "ner",
                    model="dslim/bert-base-NER",
                    aggregation_strategy="simple",
                    device=self.device
                )
                print("✅ BERT NER loaded")
                
            elif self.model_choice == "qa":
                # Question Answering - Uses ~400MB RAM
                self.models['qa'] = pipeline(
                    "question-answering",
                    model="deepset/roberta-base-squad2",
                    device=self.device
                )
                print("✅ RoBERTa QA loaded")
                
            elif self.model_choice == "zero_shot":
                # Zero-shot Classification - Uses ~1.2GB RAM
                self.models['zero_shot'] = pipeline(
                    "zero-shot-classification",
                    model="facebook/bart-large-mnli",
                    device=self.device
                )
                print("✅ BART zero-shot loaded")
            
            self.models_loaded = True
            
        except Exception as e:
            print(f"❌ Error loading model: {e}")
            self.models_loaded = False
    
    async def analyze_document(self, document_text: str, analysis_type: str) -> Tuple[str, float]:
        """Perform analysis using the loaded model"""
        
        if not self.models_loaded:
            return f"Model {self.model_choice} not loaded. Please check memory availability.", 0.0
        
        # Route to appropriate analysis based on loaded model
        if self.model_choice == "summarizer" and analysis_type == "summarize":
            return self._summarize_document(document_text)
        
        elif self.model_choice == "classifier":
            return self._classify_document(document_text)
            
        elif self.model_choice == "ner":
            return self._extract_entities(document_text)
            
        elif self.model_choice == "qa" and analysis_type in ["obligations", "timeline-extraction"]:
            return self._qa_extraction(document_text, analysis_type)
            
        elif self.model_choice == "zero_shot" and analysis_type == "extract-clauses":
            return self._extract_clauses(document_text)
        
        else:
            return f"Analysis type '{analysis_type}' not supported with loaded model '{self.model_choice}'", 0.0
    
    def _summarize_document(self, text: str) -> Tuple[str, float]:
        """Generate summary using BART"""
        
        # Split into chunks if too long
        max_chunk_length = 1024
        chunks = self._split_into_chunks(text, max_chunk_length)
        
        summaries = []
        for i, chunk in enumerate(chunks[:3]):  # Process max 3 chunks to save memory
            try:
                print(f"Summarizing chunk {i+1}...")
                summary = self.models['summarizer'](
                    chunk,
                    max_length=150,
                    min_length=50,
                    do_sample=False
                )[0]['summary_text']
                summaries.append(summary)
            except Exception as e:
                print(f"Error summarizing chunk: {e}")
                continue
        
        if summaries:
            result = "## Document Summary (BART Model)\n\n" + "\n\n".join(summaries)
            return result, 0.7
        else:
            return "Failed to generate summary", 0.0
    
    def _classify_document(self, text: str) -> Tuple[str, float]:
        """Classify document type"""
        
        try:
            # Use first 512 tokens
            result = self.models['classifier'](text[:1000])
            
            classification = f"""## Document Classification (Legal-BERT)
            
**Type**: {result[0]['label']}
**Confidence**: {result[0]['score']:.1%}

Note: This is a basic classification. For detailed analysis, use the AI-powered analysis.
"""
            return classification, result[0]['score']
            
        except Exception as e:
            return f"Classification failed: {e}", 0.0
    
    def _extract_entities(self, text: str) -> Tuple[str, float]:
        """Extract named entities"""
        
        try:
            # Process first 2000 chars to save memory
            entities = self.models['ner'](text[:2000])
            
            # Group entities by type
            entity_groups = {}
            for ent in entities:
                ent_type = ent['entity_group']
                if ent_type not in entity_groups:
                    entity_groups[ent_type] = []
                entity_groups[ent_type].append(ent['word'])
            
            result = "## Named Entity Recognition (BERT-NER)\n\n"
            for ent_type, words in entity_groups.items():
                result += f"**{ent_type}**: {', '.join(set(words))}\n"
            
            return result, 0.6
            
        except Exception as e:
            return f"NER failed: {e}", 0.0
    
    def _qa_extraction(self, text: str, analysis_type: str) -> Tuple[str, float]:
        """Extract information using QA model"""
        
        questions = {
            "obligations": [
                "What must the parties do?",
                "What are the obligations?",
                "What are the requirements?"
            ],
            "timeline-extraction": [
                "When does this start?",
                "When does this end?",
                "What are the deadlines?"
            ]
        }
        
        if analysis_type not in questions:
            return "QA extraction not configured for this analysis type", 0.0
        
        answers = []
        try:
            for q in questions[analysis_type]:
                result = self.models['qa'](question=q, context=text[:1000])
                if result['score'] > 0.3:
                    answers.append(f"Q: {q}\nA: {result['answer']} (confidence: {result['score']:.1%})")
            
            if answers:
                return f"## QA Extraction Results\n\n" + "\n\n".join(answers), 0.5
            else:
                return "No confident answers found", 0.3
                
        except Exception as e:
            return f"QA extraction failed: {e}", 0.0
    
    def _extract_clauses(self, text: str) -> Tuple[str, float]:
        """Extract clauses using zero-shot classification"""
        
        clause_types = [
            "termination clause",
            "payment terms",
            "confidentiality clause",
            "liability limitation"
        ]
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 50][:10]
        
        found_clauses = []
        
        try:
            for i, para in enumerate(paragraphs):
                print(f"Analyzing paragraph {i+1}...")
                result = self.models['zero_shot'](
                    para[:500],  # Limit length
                    candidate_labels=clause_types,
                    multi_label=False  # Single label to save memory
                )
                
                if result['scores'][0] > 0.7:
                    found_clauses.append({
                        'type': result['labels'][0],
                        'text': para[:200] + '...',
                        'score': result['scores'][0]
                    })
            
            if found_clauses:
                result = "## Extracted Clauses (BART Zero-shot)\n\n"
                for clause in found_clauses:
                    result += f"**{clause['type']}** (confidence: {clause['score']:.1%})\n{clause['text']}\n\n"
                return result, 0.6
            else:
                return "No clauses found with high confidence", 0.3
                
        except Exception as e:
            return f"Clause extraction failed: {e}", 0.0
    
    def _split_into_chunks(self, text: str, max_length: int) -> List[str]:
        """Split text into chunks for processing"""
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < max_length:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks

# Open-Source NLP Legal Analyzer (Original - kept for reference but not used)
class OpenSourceLegalAnalyzer:
    """Legal document analyzer using open-source models"""
    
    def __init__(self):
        self.device = 0 if torch.cuda.is_available() else -1
        self.models = {}
        self.models_loaded = False
        if OPEN_SOURCE_NLP_AVAILABLE:
            self.load_models()
    
    def load_models(self):
        """Load various open-source models for different tasks"""
        
        print("Loading open-source NLP models...")
        
        try:
            # 1. Document Classification (BERT-based)
            self.models['classifier'] = pipeline(
                "text-classification",
                model="nlpaueb/legal-bert-base-uncased",
                device=self.device
            )
            
            # 2. Named Entity Recognition for legal texts
            self.models['ner'] = pipeline(
                "ner",
                model="dslim/bert-base-NER",
                aggregation_strategy="simple",
                device=self.device
            )
            
            # 3. Summarization using BART or T5
            self.models['summarizer'] = pipeline(
                "summarization",
                model="facebook/bart-large-cnn",  # or "google/flan-t5-base"
                device=self.device
            )
            
            # 4. Question Answering for fact extraction
            self.models['qa'] = pipeline(
                "question-answering",
                model="deepset/roberta-base-squad2",
                device=self.device
            )
            
            # 5. Zero-shot classification for clause detection
            self.models['zero_shot'] = pipeline(
                "zero-shot-classification",
                model="facebook/bart-large-mnli",
                device=self.device
            )
            
            self.models_loaded = True
            print("✅ All open-source models loaded successfully")
            
        except Exception as e:
            print(f"❌ Error loading models: {e}")
            self.models_loaded = False
    
    async def analyze_document(self, document_text: str, analysis_type: str) -> Tuple[str, float]:
        """Perform analysis using open-source models"""
        
        if not self.models_loaded:
            return "Open-source models not loaded. Please install transformers and torch.", 0.0
        
        if analysis_type == "summarize":
            return self._summarize_document(document_text)
        
        elif analysis_type == "extract-clauses":
            return self._extract_clauses(document_text)
        
        elif analysis_type == "risk-flagging":
            return self._flag_risks(document_text)
        
        elif analysis_type == "timeline-extraction":
            return self._extract_timeline(document_text)
        
        elif analysis_type == "obligations":
            return self._extract_obligations(document_text)
        
        elif analysis_type == "missing-clauses":
            return self._detect_missing_clauses(document_text)
        
        else:
            return "Analysis type not supported with open-source models", 0.5
    
    def _summarize_document(self, text: str) -> Tuple[str, float]:
        """Generate document summary using BART/T5"""
        
        # Split into chunks if too long
        max_chunk_length = 1024
        chunks = self._split_into_chunks(text, max_chunk_length)
        
        summaries = []
        for chunk in chunks[:5]:  # Process max 5 chunks
            try:
                summary = self.models['summarizer'](
                    chunk,
                    max_length=150,
                    min_length=50,
                    do_sample=False
                )[0]['summary_text']
                summaries.append(summary)
            except:
                continue
        
        # Combine summaries
        combined_summary = "\n\n".join(summaries)
        
        # Add document type detection
        doc_type = self._detect_document_type(text)
        
        # Extract key information using QA
        parties = self._extract_parties_qa(text)
        dates = self._extract_dates_regex(text)
        amounts = self._extract_amounts_regex(text)
        
        result = f"""## Document Summary

**Document Type**: {doc_type}

**Overview**:
{combined_summary}

**Key Information Extracted**:
- **Parties**: {', '.join(parties) if parties else 'Not clearly identified'}
- **Important Dates**: {', '.join(dates[:5]) if dates else 'No dates found'}
- **Monetary Amounts**: {', '.join(amounts[:5]) if amounts else 'No amounts found'}

**Note**: This summary was generated using open-source AI models (BART/T5). For more detailed analysis, consider using the DeepSeek-powered analysis.
"""
        
        confidence = 0.7 if summaries else 0.3
        return result, confidence
    
    def _extract_clauses(self, text: str) -> Tuple[str, float]:
        """Extract clauses using zero-shot classification"""
        
        # Define clause types to look for
        clause_types = [
            "termination clause",
            "indemnification clause",
            "liability limitation",
            "confidentiality clause",
            "payment terms",
            "governing law",
            "dispute resolution"
        ]
        
        # Split text into paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 50]
        
        found_clauses = {}
        
        for para in paragraphs[:20]:  # Analyze first 20 paragraphs
            try:
                result = self.models['zero_shot'](
                    para,
                    candidate_labels=clause_types,
                    multi_label=True
                )
                
                # If high confidence, add to found clauses
                for label, score in zip(result['labels'], result['scores']):
                    if score > 0.7:
                        if label not in found_clauses:
                            found_clauses[label] = []
                        found_clauses[label].append({
                            'text': para[:200] + '...' if len(para) > 200 else para,
                            'confidence': score
                        })
            except:
                continue
        
        # Format results
        result = "## Extracted Clauses\n\n"
        
        for clause_type, instances in found_clauses.items():
            result += f"### {clause_type.title()}\n"
            for i, instance in enumerate(instances[:2], 1):  # Show max 2 per type
                result += f"{i}. {instance['text']}\n"
                result += f"   *Confidence: {instance['confidence']:.1%}*\n\n"
        
        if not found_clauses:
            result += "No standard clauses were detected with high confidence.\n"
        
        result += "\n**Note**: Clause extraction performed using zero-shot classification. Results may vary based on document structure."
        
        confidence = 0.6 if found_clauses else 0.3
        return result, confidence
    
    def _flag_risks(self, text: str) -> Tuple[str, float]:
        """Flag potential risks using pattern matching and classification"""
        
        risk_patterns = {
            "Unlimited Liability": [
                r"unlimited liability",
                r"no limitation of liability",
                r"fully liable",
                r"without limitation"
            ],
            "Unilateral Termination": [
                r"may terminate at any time",
                r"sole discretion.*terminate",
                r"without cause.*terminate",
                r"immediate termination"
            ],
            "Broad Indemnification": [
                r"indemnify.*all claims",
                r"hold harmless.*any and all",
                r"defend.*at.*own expense",
                r"indemnification.*without limitation"
            ],
            "Unfavorable Jurisdiction": [
                r"exclusive jurisdiction.*(?!your state)",
                r"governed by.*laws of.*(?!your state)",
                r"submit to.*courts of"
            ]
        }
        
        found_risks = []
        
        for risk_type, patterns in risk_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    context_start = max(0, match.start() - 100)
                    context_end = min(len(text), match.end() + 100)
                    context = text[context_start:context_end].strip()
                    
                    found_risks.append({
                        'type': risk_type,
                        'context': context,
                        'severity': 'High' if 'unlimited' in risk_type.lower() else 'Medium'
                    })
        
        # Format results
        result = "## Risk Analysis\n\n"
        
        if found_risks:
            result += f"⚠️ **{len(found_risks)} Potential Risks Identified**\n\n"
            
            for i, risk in enumerate(found_risks[:10], 1):  # Show max 10
                result += f"### Risk {i}: {risk['type']}\n"
                result += f"**Severity**: {risk['severity']}\n"
                result += f"**Context**: ...{risk['context']}...\n\n"
        else:
            result += "✅ No high-risk patterns detected in the document.\n"
        
        result += "\n**Note**: Risk detection based on pattern matching. Always consult with a legal professional for comprehensive risk assessment."
        
        confidence = 0.6 if found_risks else 0.5
        return result, confidence
    
    def _extract_timeline(self, text: str) -> Tuple[str, float]:
        """Extract dates and deadlines using regex and NER"""
        
        # Date patterns
        date_patterns = [
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',
            r'\b\d{1,2}\s+(days?|weeks?|months?|years?)\b',
            r'\b(within|after|before|by|on)\s+\d+\s+(days?|weeks?|months?|years?)\b'
        ]
        
        timeline_items = []
        
        # Extract dates with context
        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Get surrounding context
                context_start = max(0, match.start() - 50)
                context_end = min(len(text), match.end() + 50)
                context = text[context_start:context_end].strip()
                
                # Determine type based on keywords
                deadline_keywords = ['due', 'deadline', 'expire', 'terminate', 'by', 'before']
                is_deadline = any(kw in context.lower() for kw in deadline_keywords)
                
                timeline_items.append({
                    'date': match.group(0),
                    'type': 'Deadline' if is_deadline else 'Date',
                    'context': context
                })
        
        # Sort by appearance in document
        result = "## Timeline & Important Dates\n\n"
        
        if timeline_items:
            for item in timeline_items[:15]:  # Show max 15
                result += f"• **{item['date']}** ({item['type']})\n"
                result += f"  Context: ...{item['context']}...\n\n"
        else:
            result += "No specific dates or deadlines found in the document.\n"
        
        result += "\n**Note**: Dates extracted using pattern matching. Verify context for accuracy."
        
        confidence = 0.7 if timeline_items else 0.3
        return result, confidence
    
    def _extract_obligations(self, text: str) -> Tuple[str, float]:
        """Extract obligations using linguistic patterns"""
        
        obligation_patterns = [
            r'(shall|must|will|required to|obligated to)\s+([^.]{10,100})',
            r'(agrees to|commits to|undertakes to)\s+([^.]{10,100})',
            r'(responsible for|liable for)\s+([^.]{10,100})'
        ]
        
        obligations = []
        
        for pattern in obligation_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                obligation_text = match.group(0).strip()
                modal = match.group(1)
                
                obligations.append({
                    'text': obligation_text,
                    'type': 'Mandatory' if modal in ['shall', 'must'] else 'Contractual'
                })
        
        # Format results
        result = "## Extracted Obligations\n\n"
        
        if obligations:
            # Group by type
            mandatory = [o for o in obligations if o['type'] == 'Mandatory']
            contractual = [o for o in obligations if o['type'] == 'Contractual']
            
            if mandatory:
                result += "### Mandatory Obligations (SHALL/MUST)\n"
                for i, ob in enumerate(mandatory[:10], 1):
                    result += f"{i}. {ob['text']}\n"
                result += "\n"
            
            if contractual:
                result += "### Contractual Obligations\n"
                for i, ob in enumerate(contractual[:10], 1):
                    result += f"{i}. {ob['text']}\n"
        else:
            result += "No clear obligations found in the document.\n"
        
        result += "\n**Note**: Obligations extracted using linguistic patterns. Review full document for complete understanding."
        
        confidence = 0.6 if obligations else 0.3
        return result, confidence
    
    def _detect_missing_clauses(self, text: str) -> Tuple[str, float]:
        """Detect missing standard clauses"""
        
        standard_clauses = {
            "Force Majeure": ["force majeure", "act of god", "unforeseeable"],
            "Limitation of Liability": ["limitation of liability", "liability shall not exceed", "maximum liability"],
            "Indemnification": ["indemnify", "hold harmless", "defend"],
            "Confidentiality": ["confidential", "non-disclosure", "proprietary"],
            "Termination": ["termination", "terminate", "expiration"],
            "Governing Law": ["governing law", "governed by", "laws of"],
            "Dispute Resolution": ["dispute resolution", "arbitration", "mediation"],
            "Assignment": ["assignment", "assign", "transfer rights"],
            "Severability": ["severability", "severable", "invalid provision"],
            "Entire Agreement": ["entire agreement", "whole agreement", "supersedes"]
        }
        
        text_lower = text.lower()
        missing_clauses = []
        
        for clause_name, keywords in standard_clauses.items():
            found = any(keyword in text_lower for keyword in keywords)
            if not found:
                missing_clauses.append(clause_name)
        
        # Format results
        result = "## Missing Clause Analysis\n\n"
        
        if missing_clauses:
            result += f"⚠️ **{len(missing_clauses)} Standard Clauses May Be Missing:**\n\n"
            
            for clause in missing_clauses:
                result += f"### {clause}\n"
                if clause == "Force Majeure":
                    result += "Protects parties from liability due to unforeseeable events.\n"
                elif clause == "Limitation of Liability":
                    result += "Caps potential damages and financial exposure.\n"
                elif clause == "Indemnification":
                    result += "Defines who bears risk for certain claims or damages.\n"
                elif clause == "Confidentiality":
                    result += "Protects sensitive information from disclosure.\n"
                elif clause == "Governing Law":
                    result += "Specifies which jurisdiction's laws apply.\n"
                elif clause == "Dispute Resolution":
                    result += "Defines how conflicts will be resolved.\n"
                else:
                    result += "Standard contractual provision.\n"
                result += "\n"
        else:
            result += "✅ All standard clauses appear to be present.\n"
        
        result += "\n**Note**: This is a keyword-based analysis. Some clauses may be present using different terminology."
        
        confidence = 0.5  # Lower confidence for this type of analysis
        return result, confidence
    
    # Helper methods
    def _split_into_chunks(self, text: str, max_length: int) -> List[str]:
        """Split text into chunks for processing"""
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < max_length:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _detect_document_type(self, text: str) -> str:
        """Detect document type using classification"""
        try:
            # Use zero-shot classification
            labels = ["contract", "lease agreement", "employment agreement", 
                     "NDA", "license agreement", "purchase agreement", "legal brief"]
            
            result = self.models['zero_shot'](
                text[:1000],  # Use first 1000 chars
                candidate_labels=labels
            )
            
            return result['labels'][0].title()
        except:
            return "Legal Document"
    
    def _extract_parties_qa(self, text: str) -> List[str]:
        """Extract party names using QA model"""
        parties = []
        questions = [
            "Who are the parties to this agreement?",
            "What is the name of the first party?",
            "What is the name of the second party?",
            "Who is the employer?",
            "Who is the employee?",
            "Who is the landlord?",
            "Who is the tenant?"
        ]
        
        try:
            for question in questions:
                result = self.models['qa'](question=question, context=text[:2000])
                if result['score'] > 0.5:
                    parties.append(result['answer'])
        except:
            pass
        
        return list(set(parties))  # Remove duplicates
    
    def _extract_dates_regex(self, text: str) -> List[str]:
        """Extract dates using regex"""
        date_pattern = r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b'
        dates = re.findall(date_pattern, text)
        return dates
    
    def _extract_amounts_regex(self, text: str) -> List[str]:
        """Extract monetary amounts using regex"""
        amount_pattern = r'\$[\d,]+(?:\.\d{2})?|\b\d+\s*(?:USD|dollars?)\b'
        amounts = re.findall(amount_pattern, text, re.IGNORECASE)
        return amounts

# Initialize analyzers
safe_analyzer = NoHallucinationAnalyzer()
open_source_analyzer = None

# Choose which BERT model to test (for low-memory systems)
# Options: "classifier" (400MB), "ner" (400MB), "qa" (400MB), 
#          "summarizer" (1.2GB), "zero_shot" (1.2GB)
# If you have 2GB RAM + 2GB swap:
BERT_MODEL_TO_TEST = "qa"  # Best balance of usefulness and memory

# If you have 2GB RAM + 4GB swap:
BERT_MODEL_TO_TEST = "zero_shot"  # Best for contract analysis

# If zero_shot is too big:
BERT_MODEL_TO_TEST = "summarizer"  # Second best option
# Set to True to use lightweight single-model loader, False to disable BERT entirely
USE_LIGHTWEIGHT_BERT = True

if OPEN_SOURCE_NLP_AVAILABLE and USE_LIGHTWEIGHT_BERT:
    try:
        print(f"\n🔧 Loading lightweight BERT analyzer with model: {BERT_MODEL_TO_TEST}")
        print("This uses less memory by loading only one model at a time.")
        
        open_source_analyzer = LightweightLegalAnalyzer(model_choice=BERT_MODEL_TO_TEST)
        
        if open_source_analyzer.models_loaded:
            print(f"✅ Successfully loaded {BERT_MODEL_TO_TEST} model")
            print(f"💡 To test a different model, change BERT_MODEL_TO_TEST in the code")
        else:
            print(f"❌ Failed to load {BERT_MODEL_TO_TEST} model - insufficient memory")
            print("💡 Try adding swap memory or using a smaller model")
            open_source_analyzer = None
            
    except Exception as e:
        logger.error(f"Failed to initialize lightweight analyzer: {e}")
        print(f"❌ Error: {e}")
        print("💡 Continuing without BERT models - DeepSeek AI is still available")
        open_source_analyzer = None
else:
    print("ℹ️ BERT models disabled or not available")

# --- API Endpoints ---

# RAG Endpoints
@app.post("/ask", response_model=QueryResponse)
async def ask_question_improved(query: Query):
    """Improved question endpoint with enhanced accuracy and user experience"""
    cleanup_expired_conversations()
    
    # Session management
    session_id = query.session_id or str(uuid.uuid4())
    if session_id not in conversations:
        conversations[session_id] = {
            "messages": [],
            "created_at": datetime.utcnow(),
            "last_accessed": datetime.utcnow()
        }
    else:
        conversations[session_id]["last_accessed"] = datetime.utcnow()
    
    # Validate input
    user_question = query.question.strip()
    if not user_question:
        return QueryResponse(
            response=None,
            error="Question cannot be empty.",
            context_found=False,
            sources=[],
            session_id=session_id,
            confidence_score=0.0,
            expand_available=False
        )
    
    # Process with improvements
    response = process_query_improved(user_question, session_id, query.response_style)
    return response

@app.get("/conversation/{session_id}", response_model=ConversationHistory)
def get_conversation_history(session_id: str):
    """Get conversation history for a session"""
    if session_id not in conversations:
        raise HTTPException(status_code=404, detail="Conversation not found")
    conv_data = conversations[session_id]
    messages = conv_data.get('messages', [])
    return ConversationHistory(
        session_id=session_id,
        messages=messages,
        created_at=conv_data['created_at'].isoformat() if 'created_at' in conv_data else datetime.utcnow().isoformat(),
        last_updated=conv_data['last_accessed'].isoformat() if 'last_accessed' in conv_data else datetime.utcnow().isoformat()
    )

@app.delete("/conversation/{session_id}")
def clear_conversation(session_id: str):
    """Clear conversation history for a session"""
    if session_id in conversations:
        del conversations[session_id]
        return {"message": f"Conversation {session_id} cleared"}
    else:
        raise HTTPException(status_code=404, detail="Conversation not found")

@app.get("/conversations")
def list_conversations():
    """List all active conversations"""
    return {
        "active_conversations": len(conversations),
        "sessions": [
            {
                "session_id": session_id,
                "message_count": len(data.get('messages', [])),
                "created_at": data.get('created_at', datetime.min).isoformat(),
                "last_updated": data.get('last_accessed', datetime.min).isoformat()
            }
            for session_id, data in conversations.items()
        ]
    }

# Document Analysis Endpoints
@app.post("/document-analysis", response_model=EnhancedAnalysisResponse)
async def unified_document_analysis(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    analysis_type: Optional[str] = Form("summarize")
):
    """Unified document analysis with both fact extraction and AI insights"""
    
    session_id = session_id or str(uuid.uuid4())
    
    try:
        logger.info(f"Processing document: {file.filename}, type: {analysis_type}, session: {session_id}")
        
        document_text, file_type, processing_info = SafeDocumentProcessor.process_document_safe(file)
        
        extraction_results = safe_analyzer.extract_document_facts(document_text)
        factual_summary = safe_analyzer.generate_factual_summary(document_text)
        
        # Determine which AI model to use
        ai_model_used = "fact-extraction-only"
        
        if AI_ENABLED and analysis_type != "fact-extraction-only":
            ai_analysis, ai_confidence = await perform_ai_analysis(document_text, analysis_type)
            
            # Check which model was actually used based on confidence
            if ai_confidence > 0.8:
                ai_model_used = "deepseek-chat"
            elif ai_confidence > 0.5:
                ai_model_used = "open-source-nlp"
            
            combined_summary = f"""## AI Legal Analysis: {analysis_type.replace('-', ' ').title()}

{ai_analysis}

---

## Verified Facts from Document

{factual_summary}

---

**⚠️ DISCLAIMER**: 
- The analysis above is generated by AI and should be reviewed carefully
- The verified facts section contains only information extracted directly from the document
- This analysis is for informational purposes only and does not constitute legal advice
- Always consult with a qualified attorney for legal matters
"""
        else:
            ai_analysis = None
            ai_confidence = 0.0
            combined_summary = f"""## Document Analysis: Fact Extraction Only

{f'**Note**: AI analysis not available. Showing only verified facts extracted from the document.' if analysis_type != 'fact-extraction-only' else ''}

{factual_summary}

---

**To enable AI-powered analysis**:
1. Set the OPENAI_API_KEY environment variable
2. Install aiohttp: pip install aiohttp
3. For open-source models: pip install transformers torch
4. Restart the server
"""
        
        successful_extractions = len([k for k in extraction_results.keys() 
                                    if k not in ['extraction_status', 'timestamp'] and 
                                    not (isinstance(extraction_results[k], list) and 
                                         extraction_results[k] and 
                                         extraction_results[k][0].get('status') == 'failed_to_extract')])
        
        if AI_ENABLED and ai_confidence > 0.7 and successful_extractions >= 3:
            verification_status = "high_confidence"
        elif (AI_ENABLED and ai_confidence > 0.5) or successful_extractions >= 1:
            verification_status = "medium_confidence"
        else:
            verification_status = "low_confidence"
        
        logger.info(f"Analysis completed for {file.filename}: {verification_status}")
        
        return EnhancedAnalysisResponse(
            response=combined_summary,
            summary=combined_summary,
            factual_summary=combined_summary,
            ai_analysis=ai_analysis,
            extraction_results=extraction_results,
            analysis_type=analysis_type,
            confidence_score=ai_confidence if AI_ENABLED else 0.5,
            processing_info=processing_info,
            verification_status=verification_status,
            status="completed",
            success=True,
            warnings=processing_info.get('warnings', []),
            session_id=session_id,
            timestamp=datetime.utcnow().isoformat(),
            model_used=ai_model_used
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document analysis failed: {type(e).__name__}: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        error_summary = f"""## Analysis Failed

**Error**: {type(e).__name__}: {str(e)}

The document could not be analyzed. Please check:
1. The file is a valid PDF, DOCX, or TXT document
2. The file is not corrupted
3. The file size is under 10MB

If the problem persists, please try again or contact support.
"""
        
        return EnhancedAnalysisResponse(
            response=error_summary,
            summary=error_summary,
            factual_summary=error_summary,
            ai_analysis=None,
            extraction_results=None,
            analysis_type=analysis_type,
            confidence_score=0.0,
            processing_info={"error": str(e), "error_type": type(e).__name__},
            verification_status="failed",
            status="failed",
            success=False,
            warnings=[],
            session_id=session_id,
            timestamp=datetime.utcnow().isoformat()
        )

@app.post("/verify-extraction")
async def verify_extraction(
    file: UploadFile = File(...),
    claim: str = Form(...),
    line_number: Optional[int] = Form(None)
):
    """Verify a specific claim against the document"""
    
    try:
        document_text, _, _ = SafeDocumentProcessor.process_document_safe(file)
        lines = document_text.split('\n')
        
        verification_result = {
            "claim": claim,
            "verification_status": "not_found",
            "evidence": None,
            "line_number": line_number,
            "context": None
        }
        
        if line_number and 1 <= line_number <= len(lines):
            target_line = lines[line_number - 1]
            if claim.lower() in target_line.lower():
                verification_result.update({
                    "verification_status": "verified",
                    "evidence": target_line.strip(),
                    "context": target_line.strip()
                })
            else:
                verification_result.update({
                    "verification_status": "not_found_at_line",
                    "context": target_line.strip()
                })
        else:
            for i, line in enumerate(lines, 1):
                if claim.lower() in line.lower():
                    verification_result.update({
                        "verification_status": "verified",
                        "evidence": line.strip(),
                        "line_number": i,
                        "context": line.strip()
                    })
                    break
        
        return verification_result
        
    except Exception as e:
        return {
            "claim": claim,
            "verification_status": "error",
            "error": str(e)
        }

# Health and Status Endpoints
@app.get("/health")
def unified_health_check():
    """Comprehensive health check for the unified system"""
    
    # Check ChromaDB
    db_exists = os.path.exists(CHROMA_PATH)
    db_status = "healthy" if db_exists else "not_found"
    
    try:
        if db_exists:
            db = load_database()
            test_results = db.similarity_search("test", k=1)
            db_status = "healthy" if test_results is not None else "error"
    except Exception as e:
        db_status = f"error: {str(e)}"
    
    return {
        "status": "healthy",
        "version": "7.0.0-Enhanced",
        "timestamp": datetime.utcnow().isoformat(),
        "ai_enabled": AI_ENABLED,
        "database_exists": db_exists,
        "database_path": CHROMA_PATH,
        "api_key_configured": bool(OPENROUTER_API_KEY),
        "active_conversations": len(conversations),
        "unified_mode": True,
        "enhanced_rag": True,  # New flag indicating enhanced creative RAG
        "components": {
            "rag_system": {
                "enabled": db_exists,
                "database_status": db_status,
                "database_path": CHROMA_PATH,
                "nlp_model": nlp is not None,
                "sentence_model": sentence_model is not None,
                "active_conversations": len(conversations),
                "features": [
                    "enhanced_retrieval_v2",
                    "confidence_scoring",
                    "response_styles",
                    "query_expansion",
                    "multi_query_strategy"
                ]
            },
            "document_analysis": {
                "enabled": True,
                "ai_enabled": AI_ENABLED,
                "ai_model": "deepseek-chat" if AI_ENABLED else "not-configured",
                "open_source_nlp": {
                    "available": OPEN_SOURCE_NLP_AVAILABLE,
                    "models_loaded": open_source_analyzer.models_loaded if open_source_analyzer else False,
                    "models": ["legal-bert", "bart-summarization", "roberta-qa", "zero-shot-classification"] if open_source_analyzer and open_source_analyzer.models_loaded else []
                },
                "pdf_processors": {
                    "pymupdf": PYMUPDF_AVAILABLE,
                    "pdfplumber": PDFPLUMBER_AVAILABLE,
                    "pypdf2": True
                }
            },
            "api_configuration": {
                "openai_key_set": bool(OPENROUTER_API_KEY),
                "aiohttp_available": AIOHTTP_AVAILABLE,
                "api_base": OPENAI_API_BASE
            }
        },
        "features": [
            "✅ Enhanced RAG with multi-query strategies",
            "✅ Dynamic confidence scoring",
            "✅ Response style customization (concise/balanced/detailed)",
            "✅ Legal query expansion and sub-query extraction",
            "✅ Document upload and analysis",
            "✅ Zero-hallucination fact extraction",
            f"{'✅' if AI_ENABLED else '❌'} AI-powered legal analysis (DeepSeek)",
            f"{'✅' if OPEN_SOURCE_NLP_AVAILABLE else '❌'} Open-source NLP models (BERT/BART)",
            "✅ Conversation history management",
            "✅ Multi-format document support (PDF, DOCX, TXT)"
        ],
        "improvements": [
            "enhanced_retrieval",
            "confidence_scoring",
            "response_styles"
        ]
    }

@app.get("/debug-db")
def debug_database():
    """Debug endpoint to check database status"""
    if not os.path.exists(CHROMA_PATH):
        return {"error": "Database folder does not exist", "path": CHROMA_PATH}
    
    try:
        db_contents = os.listdir(CHROMA_PATH)
        db = load_database()
        
        test_queries = ["test", "document", "content", "information"]
        search_results = {}
        for query in test_queries:
            try:
                results = db.similarity_search(query, k=3)
                previews = [doc.page_content[:100] + "..." for doc in results]
                search_results[query] = {
                    "count": len(results),
                    "previews": previews
                }
            except Exception as e:
                search_results[query] = {"error": str(e)}
        
        status = "Database appears to be working" if any(r.get("count", 0) > 0 for r in search_results.values()) else "Database exists but no search results found"
        
        return {
            "database_exists": True,
            "database_path": CHROMA_PATH,
            "database_contents": db_contents,
            "search_tests": search_results,
            "status": status
        }
    except Exception as e:
        logger.error(f"Database debug failed: {e}", exc_info=True)
        return {
            "error": f"Database test failed: {str(e)}",
            "path": CHROMA_PATH,
            "suggestion": "Try running the ingestion script to recreate the database"
        }

@app.get("/extraction-capabilities")
def get_extraction_capabilities():
    """Get what the system can and cannot extract"""
    
    capabilities = {
        "rag_capabilities": {
            "description": "Enhanced Retrieval-Augmented Generation for Q&A",
            "features": [
                "Multi-query search strategies",
                "Legal term expansion",
                "Sub-query extraction",
                "Dynamic relevance thresholds",
                "Confidence scoring",
                "Response style customization",
                "Conversation history context",
                "Source citation with relevance scores"
            ],
            "improvements": [
                "Better handling of complex queries",
                "More creative and comprehensive responses",
                "Improved relevance filtering",
                "Enhanced context awareness"
            ]
        },
        "fact_extraction": {
            "dates": "Only dates with clear context (due dates, deadlines, etc.)",
            "monetary_amounts": "Only amounts with payment context",
            "percentages": "Only percentages with rate/fee context",
            "party_names": "Only clearly identified parties",
            "document_structure": "Numbered sections, headers, basic organization",
            "basic_statistics": "Word count, paragraph count, reading time"
        },
        "ai_analysis": {
            "status": "enabled" if AI_ENABLED else "disabled",
            "model": "deepseek-chat" if AI_ENABLED else "not-configured",
            "capabilities": [
                "Legal document summarization",
                "Key clause extraction",
                "Missing clause detection",
                "Risk assessment and flagging",
                "Timeline and deadline extraction",
                "Party obligation analysis"
            ] if AI_ENABLED else ["AI features disabled - set OPENAI_API_KEY to enable"]
        },
        "verification_required": "All fact extractions include line numbers for manual verification",
        "fallback_behavior": "Returns 'Failed to extract' instead of guessing"
    }
    
    return capabilities

@app.get("/", response_class=HTMLResponse)
def get_unified_interface():
    """Unified interface for the combined system"""
    
    ai_status = "✅ AI Analysis Enabled" if AI_ENABLED else "❌ AI Analysis Disabled"
    db_status = "✅ Connected" if os.path.exists(CHROMA_PATH) else "❌ Not Found"
    
    ai_instructions = "" if AI_ENABLED else """
            <div class="warning">
                <h3>🤖 Enable AI Analysis</h3>
                <p>To enable AI-powered legal analysis:</p>
                <ol>
                    <li>Set environment variable: <code>export OPENAI_API_KEY="your-key"</code></li>
                    <li>Install aiohttp: <code>pip install aiohttp</code></li>
                    <li>Restart the server</li>
                </ol>
            </div>
    """
    
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Unified Legal Assistant System - Enhanced</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; background: #f8f9fa; }}
            .container {{ max-width: 1000px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
            h1 {{ color: #2c3e50; text-align: center; }}
            .status-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; margin: 20px 0; }}
            .status-badge {{ padding: 5px 15px; border-radius: 20px; font-size: 14px; font-weight: bold; display: inline-block; }}
            .ai-enabled {{ background: #d4edda; color: #155724; }}
            .ai-disabled {{ background: #f8d7da; color: #721c24; }}
            .db-connected {{ background: #d4edda; color: #155724; }}
            .db-disconnected {{ background: #f8d7da; color: #721c24; }}
            .feature {{ background: #ecf0f1; padding: 15px; margin: 10px 0; border-radius: 5px; }}
            .capability {{ background: #e8f5e8; padding: 10px; margin: 5px 0; border-left: 4px solid #27ae60; }}
            .ai-feature {{ background: #e3f2fd; padding: 10px; margin: 5px 0; border-left: 4px solid #2196f3; }}
            .enhancement {{ background: #fff8dc; padding: 10px; margin: 5px 0; border-left: 4px solid #ffd700; }}
            .warning {{ background: #fff3cd; padding: 15px; margin: 10px 0; border-radius: 5px; border-left: 4px solid #ffc107; }}
            code {{ background: #f1f1f1; padding: 2px 5px; border-radius: 3px; font-family: monospace; }}
            .endpoint {{ background: #f8f9fa; padding: 10px; margin: 5px 0; border-left: 4px solid #6c757d; font-family: monospace; }}
            .system-card {{ background: #fff; border: 1px solid #dee2e6; border-radius: 8px; padding: 20px; margin: 15px 0; }}
            .system-title {{ font-size: 20px; font-weight: bold; color: #495057; margin-bottom: 10px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>⚖️ Unified Legal Assistant System</h1>
            <p style="text-align: center; color: #6c757d;">Enhanced RAG Q&A and Document Analysis Platform</p>
            
            <div class="status-grid">
                <div style="text-align: center;">
                    <span class="status-badge {'ai-enabled' if AI_ENABLED else 'ai-disabled'}">{ai_status}</span>
                </div>
                <div style="text-align: center;">
                    <span class="status-badge {'db-connected' if os.path.exists(CHROMA_PATH) else 'db-disconnected'}">ChromaDB: {db_status}</span>
                </div>
            </div>
            
            {ai_instructions}
            
            <div class="system-card">
                <div class="system-title">📚 Enhanced RAG Q&A System</div>
                <p>Ask questions about your ingested legal documents with improved accuracy and creativity</p>
                <div class="feature">
                    <h4>Core Features:</h4>
                    <div class="capability">Multi-query search strategies for better coverage</div>
                    <div class="capability">Legal term expansion and sub-query extraction</div>
                    <div class="capability">Dynamic confidence scoring</div>
                    <div class="capability">Response style customization (concise/balanced/detailed)</div>
                    <div class="capability">Enhanced conversation history awareness</div>
                </div>
                <div class="feature">
                    <h4>⭐ Enhanced Capabilities:</h4>
                    <div class="enhancement">Better handling of complex multi-part questions</div>
                    <div class="enhancement">More creative and comprehensive responses</div>
                    <div class="enhancement">Improved relevance filtering with dynamic thresholds</div>
                    <div class="enhancement">Expansion suggestions for concise responses</div>
                </div>
                <div class="feature">
                    <h4>Endpoints:</h4>
                    <div class="endpoint">POST /ask - Ask questions with enhanced retrieval</div>
                    <div class="endpoint">GET /conversation/{{session_id}} - Get conversation history</div>
                    <div class="endpoint">DELETE /conversation/{{session_id}} - Clear conversation</div>
                    <div class="endpoint">GET /conversations - List all conversations</div>
                </div>
            </div>
            
            <div class="system-card">
                <div class="system-title">📄 Document Analysis System</div>
                <p>Upload and analyze individual legal documents</p>
                <div class="feature">
                    <h4>Fact Extraction (Always Available):</h4>
                    <div class="capability">Extract dates with clear context</div>
                    <div class="capability">Find monetary amounts with payment context</div>
                    <div class="capability">Identify percentages with rate/fee context</div>
                    <div class="capability">Locate party names when clearly identified</div>
                    <div class="capability">Analyze document structure</div>
                    <div class="capability">Provide line numbers for all extracted information</div>
                </div>
                
                <div class="feature">
                    <h4>AI Analysis {"(Active)" if AI_ENABLED else "(Inactive)"}:</h4>
                    <div class="ai-feature">Legal document summarization</div>
                    <div class="ai-feature">Key clause extraction and analysis</div>
                    <div class="ai-feature">Missing clause detection</div>
                    <div class="ai-feature">Legal risk assessment</div>
                    <div class="ai-feature">Timeline and deadline extraction</div>
                    <div class="ai-feature">Party obligation analysis</div>
                </div>
                
                <div class="feature">
                    <h4>Analysis Types:</h4>
                    <ul>
                        <li><code>summarize</code> - Comprehensive document summary</li>
                        <li><code>extract-clauses</code> - Extract key legal clauses</li>
                        <li><code>missing-clauses</code> - Identify missing standard clauses</li>
                        <li><code>risk-flagging</code> - Flag potential legal risks</li>
                        <li><code>timeline-extraction</code> - Extract all dates and deadlines</li>
                        <li><code>obligations</code> - List party obligations</li>
                        <li><code>fact-extraction-only</code> - Only extract verifiable facts (no AI)</li>
                    </ul>
                </div>
                
                <div class="feature">
                    <h4>Endpoints:</h4>
                    <div class="endpoint">POST /document-analysis - Analyze uploaded document</div>
                    <div class="endpoint">POST /verify-extraction - Verify specific claims</div>
                </div>
            </div>
            
            <div class="feature">
                <h3>🔧 System Endpoints</h3>
                <div class="endpoint">GET /health - System health check</div>
                <div class="endpoint">GET /debug-db - Database debugging info</div>
                <div class="endpoint">GET /extraction-capabilities - View all capabilities</div>
            </div>
            
            <div class="feature">
                <h3>📋 Supported File Types</h3>
                <p>• PDF (with PyMuPDF, pdfplumber, or PyPDF2)</p>
                <p>• DOCX (Word documents)</p>
                <p>• TXT (Plain text)</p>
            </div>
            
            <p style="text-align: center; color: #7f8c8d; margin-top: 30px;">
                {"Powered by DeepSeek AI via OpenRouter 🚀" if AI_ENABLED else "Configure AI for enhanced analysis 🔧"}
                <br>Version 7.0.0-Enhanced with Creative RAG
            </p>
        </div>
    </body>
    </html>
    """

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    logger.info(f"🚀 Starting Unified Legal Assistant (Enhanced) on port {port}")
    logger.info(f"ChromaDB Path: {CHROMA_PATH}")
    logger.info(f"AI Status: {'ENABLED with DeepSeek' if AI_ENABLED else 'DISABLED - Set OPENAI_API_KEY to enable'}")
    logger.info(f"PDF processing: PyMuPDF={PYMUPDF_AVAILABLE}, pdfplumber={PDFPLUMBER_AVAILABLE}")
    logger.info(f"Enhanced RAG features: confidence scoring, response styles, multi-query strategies")
    uvicorn.run(app, host="0.0.0.0", port=port)
