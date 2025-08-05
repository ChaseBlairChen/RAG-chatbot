"""
Refactored, extensible document processing service using the Strategy design pattern.
"""
import os
import io
import logging
import tempfile
from typing import Tuple, List, Dict, Callable, Optional, Any
from contextlib import contextmanager

from pydantic import BaseModel, Field

# Assuming these are in your project structure
from ..config import FeatureFlags
from ..core.exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)


# --- Pydantic Model for a Clear and Typed Return Value ---

class ProcessingResult(BaseModel):
    """Structured result of a document processing operation."""
    content: str
    page_count: int = Field(..., ge=0)
    warnings: List[str] = Field(default_factory=list)


# --- Refactored Document Processor Class ---

class DocumentProcessor:
    """
    Processes various document types using a strategy pattern for extensibility.
    
    Usage:
        processor = DocumentProcessor()
        try:
            result = processor.process(file_content, "mydoc.pdf")
            print(f"Content length: {len(result.content)}")
        except DocumentProcessingError as e:
            print(f"Failed to process document: {e}")
    """
    def __init__(self):
        # The Strategy Pattern: Map file extensions to handler methods
        self.handlers: Dict[str, Callable] = {
            '.txt': self._process_text,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
        }
        # Define the PDF processing strategies in order of preference
        self.pdf_strategies = [
            ("unstructured", self._pdf_handler_unstructured, FeatureFlags.UNSTRUCTURED_AVAILABLE),
            ("pymupdf", self._pdf_handler_pymupdf, FeatureFlags.PYMUPDF_AVAILABLE),
            ("pdfplumber", self._pdf_handler_pdfplumber, FeatureFlags.PDFPLUMBER_AVAILABLE),
            ("ocr", self._pdf_handler_ocr, FeatureFlags.OCR_AVAILABLE),
        ]

    def process(self, file_content: bytes, filename: str) -> ProcessingResult:
        """
        Processes a document from its content bytes and filename.
        This is the single public entry point.
        """
        logger.info(f"Starting processing for '{filename}', size: {len(file_content)} bytes")
        file_ext = os.path.splitext(filename)[1].lower()
        handler = self.handlers.get(file_ext, self._process_unsupported_as_text)

        try:
            result = handler(file_content, filename)
            # Perform a final validation on the extracted content
            if not self._validate_extraction(result.content):
                result.warnings.append("Low quality extraction detected - may need manual review.")
                logger.warning(f"Low quality extraction for '{filename}'")
            
            logger.info(f"Successfully processed '{filename}': {len(result.content)} chars, "
                        f"{result.page_count} pages, {len(result.warnings)} warnings.")
            return result
        except Exception as e:
            logger.error(f"Failed to process '{filename}': {e}", exc_info=True)
            # Wrap all errors in a custom exception for clean handling by the caller
            raise DocumentProcessingError(f"Handler for '{file_ext}' failed: {e}") from e

    # --- Handler Methods (Strategies) ---

    def _process_text(self, content_bytes: bytes, filename: str) -> ProcessingResult:
        content = content_bytes.decode('utf-8', errors='ignore')
        page_count = self._estimate_pages_from_text(content)
        return ProcessingResult(content=content, page_count=page_count)

    def _process_docx(self, content_bytes: bytes, filename: str) -> ProcessingResult:
        # Prefer Unstructured.io if available, as it's more powerful
        if FeatureFlags.UNSTRUCTURED_AVAILABLE:
            try:
                logger.debug("Attempting DOCX processing with Unstructured.io")
                with self._temp_file(content_bytes, ".docx") as temp_path:
                    from unstructured.partition.auto import partition
                    elements = partition(filename=temp_path)
                
                content = "\n\n".join([el.text for el in elements if hasattr(el, 'text')])
                page_count = self._estimate_pages_from_text(content)
                return ProcessingResult(content=content, page_count=page_count)
            except Exception as e:
                logger.warning(f"Unstructured.io failed for DOCX '{filename}', falling back. Error: {e}")
        
        # Fallback to python-docx
        logger.debug("Processing DOCX with python-docx fallback.")
        try:
            from docx import Document
            doc = Document(io.BytesIO(content_bytes))
            
            text_content = []
            for para in doc.paragraphs:
                text_content.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    text_content.append(row_text)
            
            content = "\n".join(text_content)
            page_count = self._estimate_pages_from_text(content)
            return ProcessingResult(content=content, page_count=page_count)
        except ImportError:
            raise DocumentProcessingError("python-docx is not installed. Cannot process .docx files.")
        
    def _process_pdf(self, content_bytes: bytes, filename: str) -> ProcessingResult:
        """Iterates through PDF processing strategies until one succeeds."""
        warnings = []
        for name, method, is_available in self.pdf_strategies:
            if is_available:
                try:
                    logger.debug(f"Attempting PDF processing with '{name}'")
                    result = method(content_bytes, filename)
                    if result and result.content.strip():
                        logger.info(f"✅ Successfully processed PDF '{filename}' with '{name}'.")
                        result.warnings.extend(warnings)
                        return result
                    else:
                        warnings.append(f"Method '{name}' produced no content.")
                except Exception as e:
                    logger.warning(f"PDF processing with '{name}' failed for '{filename}': {e}")
                    warnings.append(f"Method '{name}' failed: {str(e)}")
            else:
                logger.debug(f"Skipping PDF processing with '{name}' (not available).")
        
        raise DocumentProcessingError(f"All available PDF processing methods failed for '{filename}'.")

    def _process_unsupported_as_text(self, content_bytes: bytes, filename: str) -> ProcessingResult:
        file_ext = os.path.splitext(filename)[1].lower()
        warnings = [f"Unsupported file type '{file_ext}'. Attempting to process as plain text."]
        result = self._process_text(content_bytes, filename)
        result.warnings.extend(warnings)
        return result

    # --- PDF Strategy Implementations ---
    
    def _pdf_handler_unstructured(self, content_bytes: bytes, filename: str) -> Optional[ProcessingResult]:
        with self._temp_file(content_bytes, ".pdf") as temp_path:
            from unstructured.partition.auto import partition
            elements = partition(filename=temp_path, strategy="hi_res")
        content = "\n\n".join([el.text for el in elements if hasattr(el, 'text')])
        
        page_numbers = [el.metadata.page_number for el in elements if hasattr(el, 'metadata') and hasattr(el.metadata, 'page_number')]
        page_count = max(page_numbers) if page_numbers else self._estimate_pages_from_text(content)
        
        return ProcessingResult(content=content, page_count=page_count)

    def _pdf_handler_pymupdf(self, content_bytes: bytes, filename: str) -> Optional[ProcessingResult]:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content_bytes, filetype="pdf")
        content = "\n\n".join([page.get_text() for page in doc])
        return ProcessingResult(content=content, page_count=len(doc))

    def _pdf_handler_pdfplumber(self, content_bytes: bytes, filename: str) -> Optional[ProcessingResult]:
        import pdfplumber
        all_text = []
        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            for page in pdf.pages:
                all_text.append(page.extract_text() or "")
            page_count = len(pdf.pages)
        return ProcessingResult(content="\n".join(all_text), page_count=page_count)

    def _pdf_handler_ocr(self, content_bytes: bytes, filename: str) -> Optional[ProcessingResult]:
        from pdf2image import convert_from_bytes
        import pytesseract

        images = convert_from_bytes(content_bytes, dpi=300)
        if not images:
            return None

        all_text = []
        for i, image in enumerate(images):
            # Basic preprocessing could be added here (grayscale, contrast, etc.)
            text = pytesseract.image_to_string(image, lang='eng')
            all_text.append(f"\n--- Page {i + 1} (OCR) ---\n{text}")

        return ProcessingResult(content="\n".join(all_text), page_count=len(images))

    # --- Private Helper Methods ---

    def _validate_extraction(self, text: str) -> bool:
        """Performs basic checks on the quality of extracted text."""
        if not text or len(text.strip()) < 50:
            return False
        # Check for high ratio of replacement characters (a sign of encoding errors)
        if text.count('') / len(text) > 0.01:
            return False
        return True

    def _estimate_pages_from_text(self, text: str) -> int:
        """Estimates page count from raw text content."""
        if not text.strip():
            return 0
        # A common heuristic is ~2000-2500 characters per page.
        return max(1, len(text) // 2500)

    @contextmanager
    def _temp_file(self, content_bytes: bytes, suffix: str):
        """A context manager to safely handle temporary files."""
        # This is needed for libraries that require a file path instead of bytes.
        fd, temp_path = tempfile.mkstemp(suffix=suffix)
        os.write(fd, content_bytes)
        os.close(fd)
        try:
            yield temp_path
        finally:
            os.unlink(temp_path)
