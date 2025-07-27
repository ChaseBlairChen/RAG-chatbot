# Merged Legal Document Analysis System with AI Enhancement
# Combines fact extraction with AI-powered analysis

from fastapi import FastAPI, Request, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
import os
import json
import re
import logging
import uuid
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Tuple, Set, Any
from dataclasses import dataclass
from enum import Enum
import io
import tempfile
import sys
import traceback

# AI imports
try:
    import aiohttp
    AIOHTTP_AVAILABLE = True
except ImportError:
    AIOHTTP_AVAILABLE = False
    print("⚠️ aiohttp not available - AI features disabled. Install with: pip install aiohttp")

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# OpenRouter configuration for AI features
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_API_KEY = os.environ.get("OPENAI_API_KEY", "")
AI_ENABLED = AIOHTTP_AVAILABLE and bool(OPENROUTER_API_KEY)

if AI_ENABLED:
    print("✅ AI features enabled with OpenRouter/DeepSeek")
else:
    print("⚠️ AI features disabled - using fact extraction only")

# Import document processing libraries
print("Starting document processing imports...")
PYMUPDF_AVAILABLE = False
PDFPLUMBER_AVAILABLE = False

# Import PyPDF2 (always needed as fallback)
try:
    import PyPDF2
    print("✅ PyPDF2 imported successfully")
except ImportError as e:
    print(f"❌ CRITICAL: PyPDF2 import failed: {e}")
    print("Install with: pip install PyPDF2")
    sys.exit(1)

# Import python-docx
try:
    import docx
    print("✅ python-docx imported successfully")
except ImportError as e:
    print(f"❌ CRITICAL: python-docx import failed: {e}")
    print("Install with: pip install python-docx")
    sys.exit(1)

# Try to import PyMuPDF
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    print("✅ PyMuPDF available - using high-quality PDF processing")
except ImportError as e:
    print(f"⚠️ PyMuPDF not available: {e}")
    print("Install with: pip install PyMuPDF")

# Try to import pdfplumber
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    print("✅ pdfplumber available - using enhanced PDF extraction")
except ImportError as e:
    print(f"⚠️ pdfplumber not available: {e}")
    print("Install with: pip install pdfplumber")

print(f"PDF processing status: PyMuPDF={PYMUPDF_AVAILABLE}, pdfplumber={PDFPLUMBER_AVAILABLE}")

# AI Analysis prompts
ANALYSIS_PROMPTS = {
    'summarize': """You are a legal document analyst. Analyze this document and provide:
1. A clear summary of the document's purpose and type
2. The main parties involved (with their roles)
3. Key terms and conditions
4. Important dates and deadlines
5. Financial obligations or amounts
6. Any notable risks or concerns

Document text:
{document_text}

Provide a structured summary in plain English while maintaining legal accuracy.""",

    'extract-clauses': """You are a legal document analyst. Extract and categorize the following types of clauses from this document:
1. Termination clauses
2. Indemnification provisions
3. Liability limitations
4. Governing law and jurisdiction
5. Confidentiality/NDA provisions
6. Payment terms
7. Dispute resolution mechanisms

For each clause found, provide:
- Clause type
- Summary of the provision
- Exact location/section reference if available
- Any unusual or concerning aspects

Document text:
{document_text}""",

    'missing-clauses': """You are a legal document analyst. Review this contract and identify commonly expected clauses that appear to be missing or inadequately addressed:

Consider standard clauses such as:
- Force majeure
- Limitation of liability
- Indemnification
- Dispute resolution/arbitration
- Confidentiality
- Termination provisions
- Assignment restrictions
- Severability
- Entire agreement
- Notice provisions
- Governing law

Document text:
{document_text}

For each missing clause, explain why it's typically important and the risks of its absence.""",

    'risk-flagging': """You are a legal risk analyst. Identify and assess legal risks in this document:

Look for:
1. Unilateral termination rights
2. Broad indemnification requirements
3. Unlimited liability exposure
4. Vague or ambiguous obligations
5. Unfavorable payment terms
6. Lack of protection clauses
7. Unusual warranty provisions
8. Problematic intellectual property terms

Document text:
{document_text}

For each risk, provide:
- Risk description
- Severity (High/Medium/Low)
- Potential impact
- Suggested mitigation""",

    'timeline-extraction': """You are a legal document analyst. Extract all time-related information:

Find and list:
1. Contract start and end dates
2. Payment deadlines
3. Notice periods
4. Renewal dates and terms
5. Termination notice requirements
6. Performance deadlines
7. Warranty periods
8. Any other time-sensitive obligations

Document text:
{document_text}

Present as a chronological timeline with clear labels.""",

    'obligations': """You are a legal document analyst. List all obligations and requirements for each party:

Identify:
1. What each party must do
2. When they must do it
3. Conditions or prerequisites
4. Consequences of non-compliance
5. Reporting or notification requirements

Document text:
{document_text}

Organize by party and priority/timeline."""
}

# AI analysis function
async def perform_ai_analysis(document_text: str, analysis_type: str) -> Tuple[str, float]:
    """Perform AI analysis using DeepSeek via OpenRouter"""
    
    if not AI_ENABLED:
        return "AI analysis not available. Please set OPENAI_API_KEY environment variable.", 0.0
    
    # Get the appropriate prompt
    prompt_template = ANALYSIS_PROMPTS.get(analysis_type, ANALYSIS_PROMPTS['summarize'])
    prompt = prompt_template.format(document_text=document_text[:15000])  # Limit context length
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8000",
        "X-Title": "Legal Document Analyzer"
    }
    
    payload = {
        "model": "deepseek/deepseek-chat",
        "messages": [
            {
                "role": "system",
                "content": "You are an expert legal document analyst. Provide thorough, accurate analysis while clearly marking any uncertainties. Always include relevant disclaimers about seeking professional legal advice."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3,
        "max_tokens": 2000
    }
    
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(OPENROUTER_API_URL, json=payload, headers=headers) as response:
                if response.status == 200:
                    data = await response.json()
                    ai_response = data['choices'][0]['message']['content']
                    confidence = 0.85
                    return ai_response, confidence
                else:
                    error_text = await response.text()
                    logger.error(f"OpenRouter API error: {response.status} - {error_text}")
                    return f"AI analysis failed: {response.status}", 0.0
    except Exception as e:
        logger.error(f"AI analysis exception: {e}")
        return f"AI analysis failed: {str(e)}", 0.0

class VerifiableExtractor:
    """Extract only verifiable information from documents with source locations"""
    
    def __init__(self):
        # High-confidence extraction patterns with context requirements
        self.extraction_patterns = {
            'dates': {
                'patterns': [
                    r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
                    r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
                    r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
                ],
                'context_required': ['due', 'deadline', 'effective', 'expires', 'terminates', 'begins', 'starts', 'ends']
            },
            'monetary_amounts': {
                'patterns': [
                    r'\$\s*[\d,]+(?:\.\d{2})?(?:\s*(?:per|each|monthly|annually|yearly))?',
                    r'\b[\d,]+\s*dollars?(?:\s*(?:per|each|monthly|annually|yearly))?\b',
                    r'\b[\d,]+\s*USD\b'
                ],
                'context_required': ['pay', 'payment', 'fee', 'cost', 'price', 'rent', 'salary', 'compensation']
            },
            'percentages': {
                'patterns': [r'\b\d+(?:\.\d+)?%'],
                'context_required': ['rate', 'interest', 'fee', 'penalty', 'commission', 'tax']
            },
            'party_names': {
                'patterns': [
                    r'(?:party|parties)[:\s]+([A-Z][a-zA-Z\s&.,]+?)(?:\s*(?:and|,|\n))',
                    r'between\s+([A-Z][a-zA-Z\s&.,]+?)\s+and\s+([A-Z][a-zA-Z\s&.,]+?)(?:\s|,)',
                    r'(?:company|corporation|llc|inc\.?)[:\s]*([A-Z][a-zA-Z\s&.,]+?)(?:\s*(?:and|,|\n))'
                ],
                'context_required': []
            }
        }
    
    def extract_with_verification(self, document_text: str, extraction_type: str) -> List[Dict[str, Any]]:
        """Extract information only if it can be verified with high confidence"""
        
        if extraction_type not in self.extraction_patterns:
            return [{"status": "failed", "reason": f"Unknown extraction type: {extraction_type}"}]
        
        pattern_config = self.extraction_patterns[extraction_type]
        extracted_items = []
        
        # Split document into lines for source tracking
        lines = document_text.split('\n')
        
        for line_num, line in enumerate(lines, 1):
            line_lower = line.lower()
            
            for pattern in pattern_config['patterns']:
                matches = re.finditer(pattern, line, re.IGNORECASE)
                
                for match in matches:
                    extracted_value = match.group(0).strip()
                    
                    # Check if context requirements are met
                    context_verified = True
                    if pattern_config['context_required']:
                        context_verified = any(
                            keyword in line_lower 
                            for keyword in pattern_config['context_required']
                        )
                    
                    if context_verified:
                        extracted_items.append({
                            "value": extracted_value,
                            "line_number": line_num,
                            "context": line.strip(),
                            "confidence": "high",
                            "verified": True
                        })
                    else:
                        # Found pattern but no context - mark as unverified
                        extracted_items.append({
                            "value": extracted_value,
                            "line_number": line_num,
                            "context": line.strip(),
                            "confidence": "low",
                            "verified": False,
                            "reason": "No supporting context found"
                        })
        
        # Filter to only high-confidence, verified items
        verified_items = [item for item in extracted_items if item.get('verified', False)]
        
        if not verified_items:
            return [{"status": "failed_to_extract", "reason": f"No verifiable {extraction_type} found in document"}]
        
        return verified_items

class NoHallucinationAnalyzer:
    """Document analyzer that never hallucinates - only reports verifiable facts"""
    
    def __init__(self):
        self.extractor = VerifiableExtractor()
        
        # Strict document type detection based on explicit keywords only
        self.document_types = {
            'contract': {
                'required_keywords': ['agreement', 'contract'],
                'supporting_keywords': ['parties', 'consideration', 'terms'],
                'minimum_matches': 2
            },
            'lease': {
                'required_keywords': ['lease'],
                'supporting_keywords': ['tenant', 'landlord', 'rent', 'premises'],
                'minimum_matches': 2
            },
            'employment': {
                'required_keywords': ['employment', 'employee'],
                'supporting_keywords': ['employer', 'salary', 'position', 'job'],
                'minimum_matches': 2
            },
            'policy': {
                'required_keywords': ['policy'],
                'supporting_keywords': ['procedure', 'guidelines', 'rules'],
                'minimum_matches': 1
            }
        }
    
    def detect_document_type_strict(self, document_text: str) -> Dict[str, Any]:
        """Detect document type only with high confidence - no guessing"""
        
        text_lower = document_text.lower()
        detection_results = {}
        
        for doc_type, criteria in self.document_types.items():
            score = 0
            matched_keywords = []
            
            # Check required keywords
            required_found = 0
            for keyword in criteria['required_keywords']:
                if keyword in text_lower:
                    required_found += 1
                    matched_keywords.append(keyword)
                    score += 10
            
            # Check supporting keywords
            supporting_found = 0
            for keyword in criteria['supporting_keywords']:
                if keyword in text_lower:
                    supporting_found += 1
                    matched_keywords.append(keyword)
                    score += 3
            
            # Only consider if minimum requirements met
            total_matches = required_found + supporting_found
            if total_matches >= criteria['minimum_matches'] and required_found > 0:
                detection_results[doc_type] = {
                    'score': score,
                    'matched_keywords': matched_keywords,
                    'required_found': required_found,
                    'supporting_found': supporting_found
                }
        
        if not detection_results:
            return {
                'type': 'unknown',
                'confidence': 0.0,
                'status': 'failed_to_detect',
                'reason': 'No clear document type indicators found',
                'matched_keywords': []
            }
        
        # Get highest scoring type
        best_type = max(detection_results, key=lambda x: detection_results[x]['score'])
        best_result = detection_results[best_type]
        
        # Calculate confidence based on keyword matches
        confidence = min(1.0, best_result['score'] / 20)
        
        return {
            'type': best_type,
            'confidence': confidence,
            'status': 'detected' if confidence > 0.7 else 'uncertain',
            'matched_keywords': best_result['matched_keywords'],
            'all_matches': detection_results
        }
    
    def extract_document_facts(self, document_text: str) -> Dict[str, Any]:
        """Extract only verifiable facts from the document"""
        
        facts = {
            'basic_stats': self._get_basic_stats(document_text),
            'dates': self.extractor.extract_with_verification(document_text, 'dates'),
            'monetary_amounts': self.extractor.extract_with_verification(document_text, 'monetary_amounts'),
            'percentages': self.extractor.extract_with_verification(document_text, 'percentages'),
            'party_names': self.extractor.extract_with_verification(document_text, 'party_names'),
            'document_structure': self._analyze_structure(document_text),
            'extraction_status': 'completed',
            'timestamp': datetime.utcnow().isoformat()
        }
        
        return facts
    
    def _get_basic_stats(self, text: str) -> Dict[str, Any]:
        """Get verifiable basic statistics"""
        
        words = text.split()
        sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        return {
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len(paragraphs),
            'character_count': len(text),
            'estimated_reading_time_minutes': round(len(words) / 200, 1)
        }
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure - only count what's clearly present"""
        
        # Count numbered sections (1., 2., etc.)
        numbered_sections = len(re.findall(r'^\s*\d+\.\s', text, re.MULTILINE))
        
        # Count lettered sections (a., b., etc.)
        lettered_sections = len(re.findall(r'^\s*[a-z]\.\s', text, re.MULTILINE))
        
        # Count subsections (1.1, 1.2, etc.)
        subsections = len(re.findall(r'^\s*\d+\.\d+\s', text, re.MULTILINE))
        
        # Count headers (lines that end with ":" and are short)
        potential_headers = re.findall(r'^([^:\n]{1,50}):$', text, re.MULTILINE)
        headers = len([h for h in potential_headers if len(h.split()) <= 8])
        
        return {
            'numbered_sections': numbered_sections,
            'lettered_sections': lettered_sections,
            'subsections': subsections,
            'headers': headers,
            'has_clear_structure': numbered_sections > 0 or headers > 2
        }
    
    def generate_factual_summary(self, document_text: str) -> str:
        """Generate a summary using only extracted facts - no interpretation"""
        
        # Get document type
        doc_type_result = self.detect_document_type_strict(document_text)
        
        # Extract facts
        facts = self.extract_document_facts(document_text)
        
        # Build factual summary
        summary_parts = []
        
        # Document identification
        if doc_type_result['status'] == 'detected':
            summary_parts.append(f"**Document Type**: {doc_type_result['type'].title()} (confidence: {doc_type_result['confidence']:.1%})")
            summary_parts.append(f"**Keywords Found**: {', '.join(doc_type_result['matched_keywords'])}")
        else:
            summary_parts.append(f"**Document Type**: Failed to detect - {doc_type_result['reason']}")
        
        # Basic statistics
        stats = facts['basic_stats']
        summary_parts.append(f"\n**Document Statistics**:")
        summary_parts.append(f"• Word count: {stats['word_count']:,}")
        summary_parts.append(f"• Estimated reading time: {stats['estimated_reading_time_minutes']} minutes")
        summary_parts.append(f"• Paragraphs: {stats['paragraph_count']}")
        
        # Structure analysis
        structure = facts['document_structure']
        summary_parts.append(f"\n**Document Structure**:")
        if structure['has_clear_structure']:
            summary_parts.append(f"• Numbered sections: {structure['numbered_sections']}")
            summary_parts.append(f"• Headers found: {structure['headers']}")
            summary_parts.append(f"• Subsections: {structure['subsections']}")
        else:
            summary_parts.append("• No clear structural organization detected")
        
        # Extracted information
        summary_parts.append(f"\n**Extracted Information**:")
        
        # Dates
        dates = facts['dates']
        if dates and dates[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Dates found**: {len(dates)} verifiable dates")
            for i, date in enumerate(dates[:3], 1):  # Show first 3
                summary_parts.append(f"  {i}. {date['value']} (Line {date['line_number']})")
            if len(dates) > 3:
                summary_parts.append(f"  ... and {len(dates) - 3} more dates")
        else:
            summary_parts.append("• **Dates**: Failed to extract any verifiable dates")
        
        # Monetary amounts
        amounts = facts['monetary_amounts']
        if amounts and amounts[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Financial amounts**: {len(amounts)} verifiable amounts")
            for i, amount in enumerate(amounts[:3], 1):  # Show first 3
                summary_parts.append(f"  {i}. {amount['value']} (Line {amount['line_number']})")
            if len(amounts) > 3:
                summary_parts.append(f"  ... and {len(amounts) - 3} more amounts")
        else:
            summary_parts.append("• **Financial amounts**: Failed to extract any verifiable amounts")
        
        # Percentages
        percentages = facts['percentages']
        if percentages and percentages[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Percentages**: {len(percentages)} found")
            for perc in percentages[:3]:
                summary_parts.append(f"  - {perc['value']} (Line {perc['line_number']})")
        else:
            summary_parts.append("• **Percentages**: Failed to extract any verifiable percentages")
        
        # Party names
        parties = facts['party_names']
        if parties and parties[0].get('status') != 'failed_to_extract':
            summary_parts.append(f"• **Parties**: {len(parties)} identified")
            for party in parties[:3]:
                summary_parts.append(f"  - {party['value']} (Line {party['line_number']})")
        else:
            summary_parts.append("• **Parties**: Failed to extract clear party names")
        
        # Important disclaimer
        summary_parts.append(f"\n**⚠️ IMPORTANT NOTES**:")
        summary_parts.append("• This analysis only includes information that could be verified directly from the document text")
        summary_parts.append("• All extracted items include line numbers for verification")
        summary_parts.append("• Items marked 'Failed to extract' mean the information was not clearly identifiable")
        summary_parts.append("• For legal advice, consult a qualified attorney")
        
        return '\n'.join(summary_parts)

class SafeDocumentProcessor:
    """Document processor with verification and no hallucination"""
    
    @staticmethod
    def process_document_safe(file: UploadFile) -> Tuple[str, str, Dict[str, Any]]:
        """Process document and provide processing metadata"""
        
        file_extension = file.filename.split('.')[-1].lower() if file.filename else 'unknown'
        processing_info = {
            'original_filename': file.filename,
            'file_extension': file_extension,
            'processing_method': None,
            'success': False,
            'warnings': [],
            'errors': []
        }
        
        try:
            # Reset file pointer to beginning
            file.file.seek(0)
            
            if file_extension == 'pdf':
                text, warnings = SafeDocumentProcessor._extract_pdf_safe(file)
                processing_info['processing_method'] = 'PDF extraction'
                processing_info['warnings'] = warnings
                
            elif file_extension in ['docx', 'doc']:
                text, warnings = SafeDocumentProcessor._extract_docx_safe(file)
                processing_info['processing_method'] = 'Word document extraction'
                processing_info['warnings'] = warnings
                
            elif file_extension == 'txt':
                content = file.file.read()
                text = content.decode('utf-8')
                processing_info['processing_method'] = 'Plain text'
                
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file_extension}. Supported: PDF, DOCX, TXT"
                )
            
            # Verify text extraction quality
            if not text or len(text.strip()) < 10:
                processing_info['errors'].append("Extracted text is too short or empty")
                raise ValueError("Failed to extract meaningful text from document")
            
            processing_info['success'] = True
            processing_info['extracted_length'] = len(text)
            processing_info['word_count'] = len(text.split())
            
            return text, file_extension, processing_info
            
        except Exception as e:
            processing_info['errors'].append(str(e))
            logger.error(f"Document processing error: {type(e).__name__}: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Document processing failed: {str(e)}")
    
    @staticmethod
    def _extract_pdf_safe(file: UploadFile) -> Tuple[str, List[str]]:
        """Extract PDF text with PyMuPDF first (better than Unstructured for legal docs)"""
        
        warnings = []
        
        # Reset file pointer and read content
        file.file.seek(0)
        pdf_content = file.file.read()
        
        if len(pdf_content) == 0:
            raise ValueError("PDF file is empty")
        
        # Try PyMuPDF first (excellent for legal documents, lightweight)
        if PYMUPDF_AVAILABLE:
            try:
                import fitz  # PyMuPDF
                
                # Create document from bytes
                pdf_doc = fitz.open(stream=pdf_content, filetype="pdf")
                
                try:
                    if pdf_doc.page_count == 0:
                        raise ValueError("PDF contains no pages")
                    
                    text = ""
                    pages_with_text = 0
                    tables_found = 0
                    total_pages = pdf_doc.page_count
                    
                    for page_num in range(total_pages):
                        try:
                            page = pdf_doc[page_num]
                            
                            # Extract text with better formatting
                            page_text = page.get_text()
                            if page_text and page_text.strip():
                                text += f"[Page {page_num + 1}]\n{page_text}\n"
                                pages_with_text += 1
                            else:
                                warnings.append(f"Page {page_num + 1} contains no extractable text")
                            
                            # Extract tables (PyMuPDF can detect table-like structures)
                            try:
                                tables = page.find_tables()
                                if tables:
                                    for table_num, table in enumerate(tables):
                                        try:
                                            table_data = table.extract()
                                            if table_data:
                                                text += f"\n[TABLE {table_num + 1} FROM PAGE {page_num + 1}]\n"
                                                for row in table_data:
                                                    if row:
                                                        clean_row = [str(cell) if cell else "" for cell in row]
                                                        text += " | ".join(clean_row) + "\n"
                                                text += "[/TABLE]\n\n"
                                                tables_found += 1
                                        except Exception as table_error:
                                            logger.debug(f"Table extraction error: {table_error}")
                            except Exception as e:
                                # Table extraction is optional
                                logger.debug(f"Table detection error: {e}")
                                
                        except Exception as page_error:
                            warnings.append(f"Error processing page {page_num + 1}: {str(page_error)}")
                            logger.warning(f"Page {page_num + 1} error: {page_error}")
                    
                    if pages_with_text == 0:
                        raise ValueError("No readable text found in any PDF pages")
                    
                    if pages_with_text < total_pages:
                        warnings.append(f"Only {pages_with_text} of {total_pages} pages contained extractable text")
                    
                    if tables_found > 0:
                        warnings.append(f"PyMuPDF extracted {tables_found} tables with preserved structure")
                    else:
                        warnings.append("PyMuPDF processed PDF successfully - excellent for legal documents")
                    
                    return text, warnings
                    
                finally:
                    # Always close the document
                    pdf_doc.close()
                
            except Exception as e:
                warnings.append(f"PyMuPDF processing failed: {str(e)}, falling back to pdfplumber")
                logger.warning(f"PyMuPDF error: {e}")
        
        # Fallback to pdfplumber (good quality)
        if PDFPLUMBER_AVAILABLE:
            pdf_file = io.BytesIO(pdf_content)
            
            try:
                import pdfplumber
                with pdfplumber.open(pdf_file) as pdf:
                    if len(pdf.pages) == 0:
                        raise ValueError("PDF contains no pages")
                    
                    text = ""
                    pages_with_text = 0
                    tables_found = 0
                    
                    for page_num, page in enumerate(pdf.pages):
                        # Extract text
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"[Page {page_num + 1}]\n{page_text}\n"
                            pages_with_text += 1
                        else:
                            warnings.append(f"Page {page_num + 1} contains no extractable text")
                        
                        # Extract tables separately for better structure
                        try:
                            tables = page.extract_tables()
                            if tables:
                                for table_num, table in enumerate(tables):
                                    text += f"\n[TABLE {table_num + 1} FROM PAGE {page_num + 1}]\n"
                                    for row in table:
                                        if row:
                                            clean_row = [str(cell) if cell else "" for cell in row]
                                            text += " | ".join(clean_row) + "\n"
                                    text += "[/TABLE]\n\n"
                                    tables_found += 1
                        except Exception as e:
                            warnings.append(f"Could not extract tables from page {page_num + 1}: {str(e)}")
                    
                    if pages_with_text == 0:
                        raise ValueError("No readable text found in any PDF pages")
                    
                    if pages_with_text < len(pdf.pages):
                        warnings.append(f"Only {pages_with_text} of {len(pdf.pages)} pages contained extractable text")
                    
                    if tables_found > 0:
                        warnings.append(f"pdfplumber extracted {tables_found} tables with preserved structure")
                    
                    return text, warnings
                    
            except Exception as e:
                warnings.append(f"pdfplumber failed: {str(e)}, using basic PyPDF2")
                logger.warning(f"pdfplumber error: {e}")
        
        # Final fallback to PyPDF2 (basic quality)
        pdf_file = io.BytesIO(pdf_content)
        pdf_file.seek(0)
        
        try:
            reader = PyPDF2.PdfReader(pdf_file)
            
            if len(reader.pages) == 0:
                raise ValueError("PDF contains no pages")
            
            text = ""
            pages_with_text = 0
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += f"[Page {page_num + 1}]\n{page_text}\n\n"
                        pages_with_text += 1
                    else:
                        warnings.append(f"Page {page_num + 1} contains no extractable text")
                except Exception as e:
                    warnings.append(f"Error extracting text from page {page_num + 1}: {str(e)}")
            
            if pages_with_text == 0:
                raise ValueError("No readable text found in any PDF pages")
            
            warnings.append("Used basic PDF extraction - complex layouts may not be preserved")
            return text, warnings
            
        except Exception as e:
            logger.error(f"PyPDF2 failed: {e}")
            raise ValueError(f"All PDF extraction methods failed. Last error: {str(e)}")
    
    @staticmethod
    def _extract_docx_safe(file: UploadFile) -> Tuple[str, List[str]]:
        """Extract DOCX text with warnings about potential issues"""
        
        warnings = []
        
        # Reset file pointer and read content
        file.file.seek(0)
        docx_content = file.file.read()
        
        if len(docx_content) == 0:
            raise ValueError("Word document is empty")
        
        docx_file = io.BytesIO(docx_content)
        
        try:
            doc = docx.Document(docx_file)
            
            text = ""
            paragraphs_with_text = 0
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text and paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraphs_with_text += 1
            
            # Check tables as well
            tables_found = len(doc.tables)
            if tables_found > 0:
                warnings.append(f"Document contains {tables_found} tables - table content may not be fully extracted")
                
                # Try to extract table content
                for table_num, table in enumerate(doc.tables):
                    text += f"\n[TABLE {table_num + 1}]\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        if any(row_text):
                            text += " | ".join(row_text) + "\n"
                    text += "[/TABLE]\n\n"
            
            if paragraphs_with_text == 0:
                raise ValueError("No readable text found in Word document")
            
            if paragraphs_with_text < len(doc.paragraphs):
                warnings.append(f"Some paragraphs contained no text")
            
            return text, warnings
            
        except Exception as e:
            raise ValueError(f"Failed to process Word document: {str(e)}")

# Response models
class EnhancedAnalysisResponse(BaseModel):
    # Compatibility fields
    response: Optional[str] = None
    summary: Optional[str] = None
    factual_summary: Optional[str] = None
    
    # Analysis results
    ai_analysis: Optional[str] = None
    extraction_results: Optional[Dict[str, Any]] = None
    
    # Metadata
    analysis_type: str
    confidence_score: float
    processing_info: Optional[Dict[str, Any]] = None
    verification_status: str
    status: str = "completed"
    success: bool = True
    
    # Additional info
    warnings: List[str] = []
    session_id: str
    timestamp: str
    model_used: str = "deepseek-chat" if AI_ENABLED else "fact-extraction-only"

# Initialize the safe analyzer
safe_analyzer = NoHallucinationAnalyzer()

# FastAPI app
app = FastAPI(
    title="Legal Document Analysis with AI",
    description="Document analysis combining fact extraction with AI-powered insights",
    version="5.0.0-Unified"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/safe-document-analysis", response_model=EnhancedAnalysisResponse)
@app.post("/document-analysis")  # Main endpoint for all analysis
async def unified_document_analysis(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    analysis_type: Optional[str] = Form("summarize")
):
    """Unified document analysis with both fact extraction and AI insights"""
    
    session_id = session_id or str(uuid.uuid4())
    
    try:
        # Log incoming request
        logger.info(f"Processing document: {file.filename}, type: {analysis_type}, session: {session_id}")
        
        # Process document safely
        document_text, file_type, processing_info = SafeDocumentProcessor.process_document_safe(file)
        
        # Always extract facts first
        extraction_results = safe_analyzer.extract_document_facts(document_text)
        factual_summary = safe_analyzer.generate_factual_summary(document_text)
        
        # Perform AI analysis if enabled
        if AI_ENABLED and analysis_type != "fact-extraction-only":
            ai_analysis, ai_confidence = await perform_ai_analysis(document_text, analysis_type)
            
            # Combine AI analysis with facts
            combined_summary = f"""## AI Legal Analysis: {analysis_type.replace('-', ' ').title()}

{ai_analysis}

---

## Verified Facts from Document

{factual_summary}

---

**⚠️ DISCLAIMER**: 
- The AI analysis above is generated by DeepSeek and should be reviewed carefully
- The verified facts section contains only information extracted directly from the document
- This analysis is for informational purposes only and does not constitute legal advice
- Always consult with a qualified attorney for legal matters
"""
        else:
            # Fallback to fact extraction only
            ai_analysis = None
            ai_confidence = 0.0
            combined_summary = f"""## Document Analysis: Fact Extraction Only

{f'**Note**: AI analysis not available. Showing only verified facts extracted from the document.' if analysis_type != 'fact-extraction-only' else ''}

{factual_summary}

---

**To enable AI-powered analysis**:
1. Set the OPENAI_API_KEY environment variable
2. Install aiohttp: pip install aiohttp
3. Restart the server
"""
        
        # Determine verification status
        successful_extractions = len([k for k in extraction_results.keys() 
                                    if k not in ['extraction_status', 'timestamp'] and 
                                    not (isinstance(extraction_results[k], list) and 
                                         extraction_results[k] and 
                                         extraction_results[k][0].get('status') == 'failed_to_extract')])
        
        if AI_ENABLED and ai_confidence > 0.7 and successful_extractions >= 3:
            verification_status = "high_confidence"
        elif (AI_ENABLED and ai_confidence > 0.5) or successful_extractions >= 1:
            verification_status = "medium_confidence"
        else:
            verification_status = "low_confidence"
        
        logger.info(f"Analysis completed for {file.filename}: {verification_status}")
        
        return EnhancedAnalysisResponse(
            response=combined_summary,  # For compatibility
            summary=combined_summary,   # For compatibility
            factual_summary=combined_summary,  # For compatibility
            ai_analysis=ai_analysis,
            extraction_results=extraction_results,
            analysis_type=analysis_type,
            confidence_score=ai_confidence if AI_ENABLED else 0.5,
            processing_info=processing_info,
            verification_status=verification_status,
            status="completed",
            success=True,
            warnings=processing_info.get('warnings', []),
            session_id=session_id,
            timestamp=datetime.utcnow().isoformat(),
            model_used="deepseek-chat" if AI_ENABLED else "fact-extraction-only"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document analysis failed: {type(e).__name__}: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        error_summary = f"""## Analysis Failed

**Error**: {type(e).__name__}: {str(e)}

The document could not be analyzed. Please check:
1. The file is a valid PDF, DOCX, or TXT document
2. The file is not corrupted
3. The file size is under 10MB

If the problem persists, please try again or contact support.
"""
        
        return EnhancedAnalysisResponse(
            response=error_summary,
            summary=error_summary,
            factual_summary=error_summary,
            ai_analysis=None,
            extraction_results=None,
            analysis_type=analysis_type,
            confidence_score=0.0,
            processing_info={"error": str(e), "error_type": type(e).__name__},
            verification_status="failed",
            status="failed",
            success=False,
            warnings=[],
            session_id=session_id,
            timestamp=datetime.utcnow().isoformat()
        )

@app.post("/verify-extraction")
async def verify_extraction(
    file: UploadFile = File(...),
    claim: str = Form(...),
    line_number: Optional[int] = Form(None)
):
    """Verify a specific claim against the document"""
    
    try:
        document_text, _, _ = SafeDocumentProcessor.process_document_safe(file)
        lines = document_text.split('\n')
        
        verification_result = {
            "claim": claim,
            "verification_status": "not_found",
            "evidence": None,
            "line_number": line_number,
            "context": None
        }
        
        # If line number provided, check that specific line
        if line_number and 1 <= line_number <= len(lines):
            target_line = lines[line_number - 1]
            if claim.lower() in target_line.lower():
                verification_result.update({
                    "verification_status": "verified",
                    "evidence": target_line.strip(),
                    "context": target_line.strip()
                })
            else:
                verification_result.update({
                    "verification_status": "not_found_at_line",
                    "context": target_line.strip()
                })
        else:
            # Search entire document
            for i, line in enumerate(lines, 1):
                if claim.lower() in line.lower():
                    verification_result.update({
                        "verification_status": "verified",
                        "evidence": line.strip(),
                        "line_number": i,
                        "context": line.strip()
                    })
                    break
        
        return verification_result
        
    except Exception as e:
        return {
            "claim": claim,
            "verification_status": "error",
            "error": str(e)
        }

@app.post("/test-upload")
async def test_upload(file: UploadFile = File(...)):
    """Simple test endpoint to verify file upload"""
    try:
        file.file.seek(0)
        content = await file.read()
        return {
            "filename": file.filename,
            "size": len(content),
            "content_type": file.content_type,
            "first_100_bytes": content[:100].hex() if len(content) > 0 else "empty"
        }
    except Exception as e:
        logger.error(f"Test upload failed: {e}")
        return {"error": str(e), "type": type(e).__name__}

@app.get("/extraction-capabilities")
def get_extraction_capabilities():
    """Get what the system can and cannot extract"""
    
    capabilities = {
        "fact_extraction": {
            "dates": "Only dates with clear context (due dates, deadlines, etc.)",
            "monetary_amounts": "Only amounts with payment context",
            "percentages": "Only percentages with rate/fee context",
            "party_names": "Only clearly identified parties",
            "document_structure": "Numbered sections, headers, basic organization",
            "basic_statistics": "Word count, paragraph count, reading time"
        },
        "ai_analysis": {
            "status": "enabled" if AI_ENABLED else "disabled",
            "model": "deepseek-chat" if AI_ENABLED else "not-configured",
            "capabilities": [
                "Legal document summarization",
                "Key clause extraction",
                "Missing clause detection",
                "Risk assessment and flagging",
                "Timeline and deadline extraction",
                "Party obligation analysis"
            ] if AI_ENABLED else ["AI features disabled - set OPENAI_API_KEY to enable"]
        },
        "verification_required": "All fact extractions include line numbers for manual verification",
        "fallback_behavior": "Returns 'Failed to extract' instead of guessing"
    }
    
    return capabilities

@app.get("/health")
def health_check():
    """Health check for the unified analyzer"""
    return {
        "status": "Unified Mode Active",
        "version": "5.0.0-Unified",
        "ai_enabled": AI_ENABLED,
        "ai_model": "deepseek-chat" if AI_ENABLED else "not-configured",
        "extraction_mode": "fact-extraction + AI" if AI_ENABLED else "fact-extraction-only",
        "features": [
            "✅ Zero hallucination fact extraction",
            "📍 Source line tracking",
            "🔍 Context verification",
            f"{'✅' if AI_ENABLED else '❌'} AI-powered legal analysis",
            f"{'✅' if AI_ENABLED else '❌'} DeepSeek integration",
            "⚠️ Legal disclaimers"
        ],
        "pdf_processing": {
            "pymupdf_available": PYMUPDF_AVAILABLE,
            "pdfplumber_available": PDFPLUMBER_AVAILABLE,
            "processing_order": ["PyMuPDF (best)", "pdfplumber (good)", "PyPDF2 (basic)"]
        },
        "requirements": {
            "api_key_set": bool(OPENROUTER_API_KEY),
            "aiohttp_installed": AIOHTTP_AVAILABLE
        },
        "timestamp": datetime.utcnow().isoformat()
    }

@app.get("/", response_class=HTMLResponse)
def get_unified_interface():
    """Unified interface with AI status"""
    ai_status = "✅ AI Analysis Enabled" if AI_ENABLED else "❌ AI Analysis Disabled"
    ai_instructions = "" if AI_ENABLED else """
            <div class="warning">
                <h3>🤖 Enable AI Analysis</h3>
                <p>To enable AI-powered legal analysis:</p>
                <ol>
                    <li>Set environment variable: <code>export OPENAI_API_KEY="your-key"</code></li>
                    <li>Install aiohttp: <code>pip install aiohttp</code></li>
                    <li>Restart the server</li>
                </ol>
            </div>
    """
    
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Legal Document Analysis - Unified System</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; background: #f8f9fa; }}
            .container {{ max-width: 900px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
            h1 {{ color: #2c3e50; text-align: center; }}
            .status-badge {{ padding: 5px 15px; border-radius: 20px; font-size: 14px; font-weight: bold; }}
            .ai-enabled {{ background: #d4edda; color: #155724; }}
            .ai-disabled {{ background: #f8d7da; color: #721c24; }}
            .feature {{ background: #ecf0f1; padding: 15px; margin: 10px 0; border-radius: 5px; }}
            .capability {{ background: #e8f5e8; padding: 10px; margin: 5px 0; border-left: 4px solid #27ae60; }}
            .ai-feature {{ background: #e3f2fd; padding: 10px; margin: 5px 0; border-left: 4px solid #2196f3; }}
            .warning {{ background: #fff3cd; padding: 15px; margin: 10px 0; border-radius: 5px; border-left: 4px solid #ffc107; }}
            code {{ background: #f1f1f1; padding: 2px 5px; border-radius: 3px; font-family: monospace; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>⚖️ Legal Document Analysis System</h1>
            <p style="text-align: center;">
                <span class="status-badge {'ai-enabled' if AI_ENABLED else 'ai-disabled'}">{ai_status}</span>
            </p>
            
            {ai_instructions}
            
            <div class="feature">
                <h3>🔒 Fact Extraction (Always Available)</h3>
                <div class="capability">Extract dates with clear context</div>
                <div class="capability">Find monetary amounts with payment context</div>
                <div class="capability">Identify percentages with rate/fee context</div>
                <div class="capability">Locate party names when clearly identified</div>
                <div class="capability">Analyze document structure</div>
                <div class="capability">Provide line numbers for all extracted information</div>
            </div>
            
            <div class="feature">
                <h3>🤖 AI-Powered Analysis {"(Active)" if AI_ENABLED else "(Inactive)"}</h3>
                <div class="ai-feature">Legal document summarization</div>
                <div class="ai-feature">Key clause extraction and analysis</div>
                <div class="ai-feature">Missing clause detection</div>
                <div class="ai-feature">Legal risk assessment</div>
                <div class="ai-feature">Timeline and deadline extraction</div>
                <div class="ai-feature">Party obligation analysis</div>
            </div>
            
            <div class="feature">
                <h3>📍 Available Analysis Types</h3>
                <p>Use these values for the <code>analysis_type</code> parameter:</p>
                <ul>
                    <li><code>summarize</code> - Comprehensive document summary</li>
                    <li><code>extract-clauses</code> - Extract key legal clauses</li>
                    <li><code>missing-clauses</code> - Identify missing standard clauses</li>
                    <li><code>risk-flagging</code> - Flag potential legal risks</li>
                    <li><code>timeline-extraction</code> - Extract all dates and deadlines</li>
                    <li><code>obligations</code> - List party obligations</li>
                    <li><code>fact-extraction-only</code> - Only extract verifiable facts (no AI)</li>
                </ul>
            </div>
            
            <div class="feature">
                <h3>🔍 API Endpoints</h3>
                <p><strong>POST /document-analysis</strong> - Main analysis endpoint</p>
                <p><strong>POST /verify-extraction</strong> - Verify specific claims</p>
                <p><strong>GET /extraction-capabilities</strong> - See capabilities</p>
                <p><strong>GET /health</strong> - System status</p>
            </div>
            
            <p style="text-align: center; color: #7f8c8d; margin-top: 30px;">
                {"Powered by DeepSeek AI via OpenRouter 🚀" if AI_ENABLED else "Configure AI for enhanced analysis 🔧"}
            </p>
        </div>
    </body>
    </html>
    """

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    logger.info(f"🚀 Starting Unified Legal Document Analysis on port {port}")
    logger.info(f"AI Status: {'ENABLED with DeepSeek' if AI_ENABLED else 'DISABLED - Set OPENAI_API_KEY to enable'}")
    logger.info(f"PDF processing: PyMuPDF={PYMUPDF_AVAILABLE}, pdfplumber={PDFPLUMBER_AVAILABLE}")
    uvicorn.run(app, host="0.0.0.0", port=port)
